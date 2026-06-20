from __future__ import annotations

import io
import math
import os
import base64
import sqlite3
import hashlib
import hmac
import secrets
import json
import re
import shutil
import subprocess
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from datetime import date, datetime, time, timedelta
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple


# Dependencia opcional para despliegue productivo con PostgreSQL/Supabase.
# La app sigue funcionando en modo local con SQLite si no se configura DATABASE_URL.
try:
    import psycopg2
    from psycopg2.extras import RealDictCursor
except Exception:  # pragma: no cover - opcional en modo local
    psycopg2 = None
    RealDictCursor = None

import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

APP_DIR = Path(__file__).resolve().parent
TEMPLATE_DIR = APP_DIR / "plantilla"
TEMPLATE_GC72 = TEMPLATE_DIR / "FD-GC72-Informe_Academico.docx"
TEMPLATE_GC71 = TEMPLATE_DIR / "FD-GC71.docx"
LOGO_POLI = TEMPLATE_DIR / "logo_poli.png"
LOGO_ICONTEC = TEMPLATE_DIR / "logo_icontec.png"
DATA_DIR = APP_DIR / "app_data"
DB_PATH = DATA_DIR / "fdgc_app.sqlite3"

# -----------------------------------------------------------------------------
# Configuración productiva / Streamlit Cloud
# -----------------------------------------------------------------------------
# Modo recomendado en producción:
#   APP_ENV="production"
#   DATABASE_URL="postgresql://usuario:clave@host:5432/base"
#   INITIAL_ADMIN_USER / INITIAL_ADMIN_PASSWORD en st.secrets o variables de entorno.
# Si DATABASE_URL no existe, la app funciona con SQLite local para desarrollo o demo.

APP_VERSION = "7.0.1-hotfix-streamlit-cloud"
DEFAULT_MAX_EVIDENCE_MB = 15


def _get_secret_value(name: str, default: Optional[str] = None) -> Optional[str]:
    """Lee configuración desde Streamlit secrets o variables de entorno sin romper ejecución local."""
    try:
        import streamlit as _st
        if hasattr(_st, "secrets") and name in _st.secrets:
            return _st.secrets.get(name)
    except Exception:
        pass
    return os.getenv(name, default)


def get_app_env() -> str:
    return str(_get_secret_value("APP_ENV", os.getenv("APP_ENV", "local")) or "local").lower().strip()


def get_database_url() -> str:
    return str(_get_secret_value("DATABASE_URL", os.getenv("DATABASE_URL", "")) or "").strip()


def _database_url_es_placeholder(url: str) -> bool:
    """Evita que Streamlit Cloud intente conectarse a valores de ejemplo.

    Si el usuario pega secrets.example.toml sin reemplazar DATABASE_URL, psycopg2
    puede quedarse intentando resolver HOST y la app parece eterna en "Your app is in the oven".
    """
    if not url:
        return False
    u = url.upper()
    marcadores = ["USUARIO", "CLAVE", "HOST", "BASE", "PASSWORD", "XXXXX"]
    return any(m in u for m in marcadores)


def usar_postgres() -> bool:
    url = get_database_url().strip()
    low = url.lower()
    if _database_url_es_placeholder(url):
        return False
    return low.startswith("postgres://") or low.startswith("postgresql://")


def postgres_url_normalizada() -> str:
    # Algunos proveedores entregan postgres:// y psycopg2 espera postgresql:// en ciertos entornos.
    url = get_database_url()
    if url.startswith("postgres://"):
        return "postgresql://" + url[len("postgres://"):]
    return url


def _traducir_sql(sql: str) -> str:
    """Traduce placeholders SQLite (?) a PostgreSQL (%s) para consultas parametrizadas simples."""
    return sql.replace("?", "%s") if usar_postgres() else sql


class PgConnectionAdapter:
    """Adaptador mínimo para conservar la API usada por la app con SQLite."""

    def __init__(self):
        if psycopg2 is None:
            raise RuntimeError("psycopg2-binary no está instalado. Revise requirements.txt.")
        timeout = _safe_int_secret("DB_CONNECT_TIMEOUT", 8)
        # Timeout explícito: evita que una DATABASE_URL mala deje Streamlit cargando indefinidamente.
        self._conn = psycopg2.connect(postgres_url_normalizada(), connect_timeout=timeout)

    def execute(self, sql: str, params: Tuple[Any, ...] = tuple()):
        cur = self._conn.cursor(cursor_factory=RealDictCursor)
        cur.execute(_traducir_sql(sql), params or tuple())
        return cur

    def cursor(self):
        return self._conn.cursor(cursor_factory=RealDictCursor)

    def commit(self):
        self._conn.commit()

    def rollback(self):
        self._conn.rollback()

    def close(self):
        self._conn.close()


def conexion_db():
    """Conexión unificada. PostgreSQL si DATABASE_URL está configurado; SQLite en local/demo."""
    if usar_postgres():
        return PgConnectionAdapter()
    DATA_DIR.mkdir(exist_ok=True)
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn


def read_sql_df(sql: str, params: Tuple[Any, ...] = tuple()) -> pd.DataFrame:
    """Lectura robusta a DataFrame para SQLite/PostgreSQL."""
    conn = conexion_db()
    try:
        if usar_postgres():
            cur = conn.execute(sql, params)
            rows = cur.fetchall()
            return pd.DataFrame([dict(r) for r in rows])
        return pd.read_sql_query(sql, conn, params=params)
    finally:
        conn.close()


def db_execute(sql: str, params: Tuple[Any, ...] = tuple(), fetchone: bool = False, fetchall: bool = False):
    conn = conexion_db()
    try:
        cur = conn.execute(sql, params)
        result = None
        if fetchone:
            row = cur.fetchone()
            result = dict(row) if row is not None and usar_postgres() else row
        elif fetchall:
            rows = cur.fetchall()
            result = [dict(r) for r in rows] if usar_postgres() else rows
        conn.commit()
        return result
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        raise
    finally:
        conn.close()


def _safe_int_secret(name: str, default: int) -> int:
    try:
        return int(_get_secret_value(name, str(default)) or default)
    except Exception:
        return default


def max_evidence_bytes() -> int:
    return _safe_int_secret("MAX_EVIDENCE_MB", DEFAULT_MAX_EVIDENCE_MB) * 1024 * 1024


def initial_admin_config() -> Tuple[str, str, str, str]:
    """Credenciales de arranque parametrizables por secretos. No las deje quemadas en producción."""
    return (
        str(_get_secret_value("INITIAL_ADMIN_USER", "admin") or "admin"),
        str(_get_secret_value("INITIAL_ADMIN_PASSWORD", "Admin123*") or "Admin123*"),
        str(_get_secret_value("INITIAL_ADMIN_NAME", "Administrador del sistema") or "Administrador del sistema"),
        str(_get_secret_value("INITIAL_ADMIN_EMAIL", "") or ""),
    )


def add_column_if_missing(table: str, column: str, ddl_type: str):
    """Migración defensiva compatible con SQLite/PostgreSQL."""
    conn = conexion_db()
    try:
        if usar_postgres():
            cur = conn.execute(
                "SELECT column_name FROM information_schema.columns WHERE table_name=%s AND column_name=%s",
                (table, column),
            )
            exists = cur.fetchone() is not None
            if not exists:
                conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {ddl_type}")
        else:
            cur = conn.execute(f"PRAGMA table_info({table})")
            cols = [r[1] for r in cur.fetchall()]
            if column not in cols:
                conn.execute(f"ALTER TABLE {table} ADD COLUMN {column} {ddl_type}")
        conn.commit()
    except Exception:
        try:
            conn.rollback()
        except Exception:
            pass
        # No se detiene la app por una migración idempotente fallida; se mostrará en diagnóstico.
    finally:
        conn.close()


def health_status() -> Dict[str, Any]:
    """Diagnóstico simple para saber si el despliegue está sano."""
    status = {
        "version": APP_VERSION,
        "app_env": get_app_env(),
        "database": "PostgreSQL" if usar_postgres() else "SQLite local",
        "database_ok": False,
        "templates_ok": TEMPLATE_GC71.exists() and TEMPLATE_GC72.exists(),
        "storage": "DB + local cache" if usar_postgres() else "SQLite/local",
        "max_evidence_mb": _safe_int_secret("MAX_EVIDENCE_MB", DEFAULT_MAX_EVIDENCE_MB),
    }
    try:
        row = db_execute("SELECT COUNT(*) AS n FROM usuarios", fetchone=True)
        status["database_ok"] = row is not None
    except Exception as exc:
        status["database_error"] = str(exc)[:300]
    return status


DIAS_MAP = {
    "Lunes": 0,
    "Martes": 1,
    "Miércoles": 2,
    "Miercoles": 2,
    "Jueves": 3,
    "Viernes": 4,
    "Sábado": 5,
    "Sabado": 5,
    "Domingo": 6,
}
DIAS_INV = {v: k for k, v in DIAS_MAP.items() if "é" in k or k not in ("Miercoles", "Sabado")}

# -----------------------------------------------------------------------------
# FD-GC72 - Informe académico
# -----------------------------------------------------------------------------
COLUMNAS_GC72 = [
    "Código",
    "Grupo",
    "Asignatura",
    "% Avance en contenido",
    "% Evaluado",
    "Estudiantes matriculados",
    "Desertaron N°",
    "Desertaron %",
    "Aprueban evaluación parcial N°",
    "Aprueban evaluación parcial %",
    "Reprueban evaluación parcial N°",
    "Reprueban evaluación parcial %",
    "Aprueban a la fecha N°",
    "Aprueban a la fecha %",
    "Reprueban a la fecha N°",
    "Reprueban a la fecha %",
]
NUMERICAS_ENTERAS_GC72 = [
    "Estudiantes matriculados",
    "Desertaron N°",
    "Aprueban evaluación parcial N°",
    "Reprueban evaluación parcial N°",
    "Aprueban a la fecha N°",
    "Reprueban a la fecha N°",
]
NUMERICAS_PORCENTAJE_GC72 = [
    "% Avance en contenido",
    "% Evaluado",
    "Desertaron %",
    "Aprueban evaluación parcial %",
    "Reprueban evaluación parcial %",
    "Aprueban a la fecha %",
    "Reprueban a la fecha %",
]
MAPA_TABLA_WORD_GC72 = COLUMNAS_GC72[:]

PREFORMAS_GC72: Dict[str, Dict[str, str]] = {
    "aspectos_positivos": {
        "Participación activa": "Se evidenció participación activa y disposición del grupo para el desarrollo de las actividades académicas propuestas.",
        "Avance adecuado": "El curso presentó un avance adecuado frente a los contenidos programados, manteniendo coherencia entre las actividades de clase y los resultados esperados.",
        "Aplicación práctica": "Los estudiantes lograron relacionar los conceptos abordados con situaciones prácticas, favoreciendo la comprensión aplicada de la asignatura.",
        "Mejora progresiva": "Se observó una mejora progresiva en la apropiación de los temas, especialmente en las actividades de seguimiento y retroalimentación.",
        "Buen cumplimiento": "La mayoría de los estudiantes cumplió con las actividades evaluativas y entregables definidos para el periodo reportado.",
        "Trabajo colaborativo": "El trabajo colaborativo fortaleció la discusión académica y permitió resolver dudas de manera más efectiva durante el curso.",
    },
    "inconvenientes": {
        "Inasistencia intermitente": "Se presentaron dificultades asociadas a inasistencia intermitente de algunos estudiantes, lo que afectó la continuidad del proceso formativo.",
        "Entregas tardías": "Algunos estudiantes realizaron entregas fuera de los tiempos establecidos, situación que limitó la retroalimentación oportuna.",
        "Brechas conceptuales": "Se identificaron brechas conceptuales en temas base, por lo cual fue necesario reforzar contenidos previos para avanzar con mayor solidez.",
        "Baja participación puntual": "En algunas sesiones se presentó baja participación de parte del grupo, especialmente en actividades que requerían preparación previa.",
        "Dificultades técnicas": "Se presentaron dificultades técnicas o de acceso a herramientas requeridas para el desarrollo de algunas actividades académicas.",
        "Carga académica acumulada": "La acumulación de actividades de otros espacios académicos incidió en el ritmo de trabajo y en la oportunidad de algunas entregas.",
    },
    "propuestas": {
        "Seguimiento formativo": "Fortalecer el seguimiento formativo mediante actividades cortas de verificación, retroalimentación temprana y acompañamiento focalizado.",
        "Talleres aplicados": "Incorporar talleres aplicados por unidades temáticas para consolidar la relación entre teoría, práctica y evaluación.",
        "Nivelación inicial": "Implementar una actividad diagnóstica y espacios de nivelación para atender brechas conceptuales antes de abordar contenidos de mayor complejidad.",
        "Rúbricas claras": "Socializar rúbricas y criterios de evaluación desde el inicio de cada actividad, con el fin de mejorar la calidad de las entregas.",
        "Aprendizaje basado en problemas": "Utilizar ejercicios basados en problemas reales del contexto profesional para incrementar la pertinencia y motivación del proceso formativo.",
        "Alertas tempranas": "Aplicar alertas tempranas frente a inasistencia, bajo desempeño o entregas pendientes, articulando acciones de mejora con los estudiantes.",
    },
}

# -----------------------------------------------------------------------------
# FD-GC71 - Guía didáctica
# -----------------------------------------------------------------------------
COLUMNAS_MODULOS = [
    "Unidad",
    "Contenido / tema central",
    "Horas presenciales",
    "Sesiones",
    "Trabajo presencial",
    "Trabajo independiente",
]
COLUMNAS_HORARIOS = ["Día", "Hora inicio", "Hora fin", "Lugar / ambiente"]
COLUMNAS_EVALUACIONES = [
    "Tipo de evaluación",
    "Procedimiento de evaluación",
    "Valor (%)",
    "Fecha de realización",
    "Unidad relacionada",
    "Corte",
]
COLUMNAS_SESIONES = [
    "Unidad",
    "N° sesión",
    "Fecha",
    "Horario",
    "Contenido por desarrollar",
    "Descripción del trabajo presencial",
    "Descripción trabajo independiente",
    "Lugar / ambiente",
]

TEXTOS_PREDEFINIDOS_GC71 = {
    "justificacion": "La asignatura aporta a la formación académica y profesional mediante la articulación entre fundamentos conceptuales, aplicación práctica y análisis de situaciones propias del campo disciplinar. Su desarrollo favorece la comprensión de problemas del contexto, el uso de herramientas técnicas y la toma de decisiones sustentada en criterios académicos, éticos y profesionales.",
    "competencias": "La asignatura tributa al desarrollo de competencias asociadas con el análisis crítico, la resolución de problemas, la comunicación técnica, el trabajo colaborativo y la aplicación de conocimientos disciplinares en escenarios reales o simulados.",
    "resultados": "Al finalizar la asignatura, el estudiante estará en capacidad de reconocer, interpretar y aplicar los conceptos y procedimientos centrales del curso, integrando evidencias, criterios técnicos y estrategias de solución acordes con los resultados de aprendizaje del programa.",
    "objetivo_general": "Desarrollar en el estudiante capacidades conceptuales, metodológicas y prácticas para comprender y aplicar los contenidos de la asignatura en situaciones propias de su formación académica y profesional.",
    "objetivos_especificos": "1. Reconocer los fundamentos conceptuales de la asignatura.\n2. Aplicar procedimientos y herramientas propias del área de formación.\n3. Analizar casos, ejercicios o problemas relacionados con el contexto profesional.\n4. Comunicar resultados de manera clara, ordenada y técnicamente sustentada.",
    "metodologias": "La asignatura se desarrollará mediante clases orientadoras, talleres aplicados, análisis de casos, ejercicios prácticos, aprendizaje basado en problemas, socialización de avances y retroalimentación permanente. Se promoverá la participación activa del estudiante y la integración entre trabajo presencial e independiente.",
    "ambientes": "Aula de clase, plataforma institucional, recursos digitales de apoyo y, cuando aplique, laboratorios, salidas pedagógicas o ambientes especializados requeridos para el logro de los resultados de aprendizaje.",
    "medios": "Presentaciones, guías de clase, material bibliográfico, bases de datos académicas, plataforma virtual, software especializado cuando aplique, tablero, equipos audiovisuales y recursos institucionales necesarios para el desarrollo de las actividades.",
    "referencias": "Bibliografía básica y complementaria definida por el docente, documentos institucionales del programa, artículos académicos recientes, recursos digitales especializados y fuentes en segunda lengua cuando sean pertinentes para la asignatura.",
}

# -----------------------------------------------------------------------------
# Utilidades generales
# -----------------------------------------------------------------------------
def df_vacio(columnas: List[str], filas: int = 5) -> pd.DataFrame:
    return pd.DataFrame([{c: "" for c in columnas} for _ in range(filas)])


def limpiar_numero(valor) -> Optional[float]:
    if valor is None or pd.isna(valor):
        return None
    if isinstance(valor, (int, float)):
        return float(valor)
    texto = str(valor).strip().replace("%", "").replace(",", ".")
    if texto == "":
        return None
    try:
        return float(texto)
    except ValueError:
        return None


def formato_entero(valor) -> str:
    numero = limpiar_numero(valor)
    if numero is None:
        return ""
    return str(int(round(numero)))


def formato_porcentaje(valor) -> str:
    numero = limpiar_numero(valor)
    if numero is None:
        return ""
    return f"{int(round(numero))}%"


def porcentaje(numerador, denominador) -> str:
    n = limpiar_numero(numerador)
    d = limpiar_numero(denominador)
    if n is None or d in (None, 0):
        return ""
    return formato_porcentaje((n / d) * 100)


def limpiar_df(df: pd.DataFrame, columnas: List[str]) -> pd.DataFrame:
    df = df.copy()
    for col in columnas:
        if col not in df.columns:
            df[col] = ""
    df = df[columnas]
    mask = df.apply(lambda row: any(str(v).strip() for v in row.values if not pd.isna(v)), axis=1)
    return df[mask].reset_index(drop=True)


def normalizar_texto(texto: str) -> str:
    texto = str(texto or "").strip().upper()
    reemplazos = {
        "Á": "A",
        "É": "E",
        "Í": "I",
        "Ó": "O",
        "Ú": "U",
        "Ü": "U",
        "Ñ": "N",
    }
    for a, b in reemplazos.items():
        texto = texto.replace(a, b)
    texto = re.sub(r"\s+", " ", texto)
    return texto


def nombre_archivo_seguro(nombre: str, fecha: date | str, prefijo: str) -> str:
    base = re.sub(r"[^A-Za-zÁÉÍÓÚÜÑáéíóúüñ0-9]+", "_", str(nombre).strip()).strip("_") or "docente"
    f = fecha.strftime("%Y%m%d") if hasattr(fecha, "strftime") else re.sub(r"\D+", "", str(fecha))
    return f"{prefijo}_{base}_{f}"


def convertir_doc_si_es_posible(docx_bytes: bytes, nombre_base: str) -> Optional[bytes]:
    ejecutable = shutil.which("soffice") or shutil.which("libreoffice")
    if not ejecutable:
        return None
    with tempfile.TemporaryDirectory() as td:
        td_path = Path(td)
        docx_path = td_path / f"{nombre_base}.docx"
        doc_path = td_path / f"{nombre_base}.doc"
        docx_path.write_bytes(docx_bytes)
        subprocess.run(
            [ejecutable, "--headless", "--convert-to", "doc", "--outdir", str(td_path), str(docx_path)],
            check=False,
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            text=True,
            timeout=60,
        )
        if doc_path.exists():
            return doc_path.read_bytes()
    return None

# -----------------------------------------------------------------------------
# Utilidades Word
# -----------------------------------------------------------------------------
def set_cell_width(cell, width_inches: float):
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def fijar_anchos_tabla(tabla, anchos: List[float]):
    tabla.autofit = False
    for row in tabla.rows:
        for idx, width in enumerate(anchos):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill.replace("#", ""))


def set_cell_text(
    cell,
    texto: str,
    font_size: float = 8.0,
    bold: bool = False,
    align=WD_ALIGN_PARAGRAPH.CENTER,
    color: Optional[str] = None,
):
    cell.text = ""
    parrafo = cell.paragraphs[0]
    parrafo.alignment = align
    parrafo.paragraph_format.space_after = Pt(0)
    run = parrafo.add_run(str(texto) if texto is not None else "")
    run.bold = bold
    run.font.size = Pt(font_size)
    if color:
        color = color.replace("#", "")
        run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_paragraph_in_cell(cell, texto: str, font_size: float = 8.5, bold: bool = False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.05
    r = p.add_run(texto)
    r.bold = bold
    r.font.size = Pt(font_size)
    return p


def set_section_text(cell, titulo: str, texto: str, font_size: float = 8.5):
    cell.text = ""
    p1 = cell.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(titulo)
    r1.bold = True
    r1.font.size = Pt(9)
    if texto:
        for parte in str(texto).split("\n"):
            if parte.strip():
                add_paragraph_in_cell(cell, parte.strip(), font_size=font_size)


def set_label_value(cell, etiqueta: str, valor: str):
    cell.text = ""
    parrafo = cell.paragraphs[0]
    parrafo.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = parrafo.add_run(etiqueta)
    r1.bold = True
    r1.font.size = Pt(10)
    r2 = parrafo.add_run(f" {valor}" if valor else "")
    r2.font.size = Pt(10)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def aplicar_bordes_tabla(tabla, color="000000", size="4"):
    tbl = tabla._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def remover_parrafo(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None

# -----------------------------------------------------------------------------
# FD-GC72
# -----------------------------------------------------------------------------
def normalizar_dataframe_gc72(df: pd.DataFrame, calcular_porcentajes: bool = True) -> pd.DataFrame:
    df = limpiar_df(df, COLUMNAS_GC72)
    if calcular_porcentajes:
        for idx, row in df.iterrows():
            matriculados = row.get("Estudiantes matriculados", "")
            df.at[idx, "Desertaron %"] = porcentaje(row.get("Desertaron N°", ""), matriculados)
            df.at[idx, "Aprueban evaluación parcial %"] = porcentaje(row.get("Aprueban evaluación parcial N°", ""), matriculados)
            df.at[idx, "Reprueban evaluación parcial %"] = porcentaje(row.get("Reprueban evaluación parcial N°", ""), matriculados)
            df.at[idx, "Aprueban a la fecha %"] = porcentaje(row.get("Aprueban a la fecha N°", ""), matriculados)
            df.at[idx, "Reprueban a la fecha %"] = porcentaje(row.get("Reprueban a la fecha N°", ""), matriculados)
    return df


def texto_preformas(seleccionadas: Iterable[str], categoria: str, adicional: str = "") -> str:
    partes = [PREFORMAS_GC72[categoria][k] for k in seleccionadas if k in PREFORMAS_GC72[categoria]]
    adicional = (adicional or "").strip()
    if adicional:
        partes.append(adicional)
    return " ".join(partes).strip()


def curso_key(row, idx: int) -> str:
    codigo = re.sub(r"\W+", "_", str(row.get("Código", "")).strip())
    grupo = re.sub(r"\W+", "_", str(row.get("Grupo", "")).strip())
    asignatura = re.sub(r"\W+", "_", str(row.get("Asignatura", "")).strip())[:40]
    return f"curso_{idx}_{codigo}_{grupo}_{asignatura}"


def agregar_parrafo_antes(parrafo_referencia, texto: str = "", bold_prefix: Optional[str] = None, space_after: int = 3):
    nuevo = parrafo_referencia.insert_paragraph_before()
    nuevo.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    nuevo.paragraph_format.space_after = Pt(space_after)
    nuevo.paragraph_format.line_spacing = 1.05
    if bold_prefix and texto.startswith(bold_prefix):
        r1 = nuevo.add_run(bold_prefix)
        r1.bold = True
        r1.font.size = Pt(10)
        r2 = nuevo.add_run(texto[len(bold_prefix):])
        r2.font.size = Pt(10)
    else:
        r = nuevo.add_run(texto)
        r.font.size = Pt(10)
    return nuevo


def construir_analisis(cursos: pd.DataFrame, analisis_por_curso: Dict[str, Dict[str, str]], modo: str) -> List[Dict[str, str]]:
    bloques = []
    if modo == "Consolidado institucional":
        positivos, inconvenientes, propuestas = [], [], []
        for data in analisis_por_curso.values():
            if data.get("positivos"):
                positivos.append(data["positivos"])
            if data.get("inconvenientes"):
                inconvenientes.append(data["inconvenientes"])
            if data.get("propuestas"):
                propuestas.append(data["propuestas"])
        bloques.append({
            "titulo": "Análisis consolidado de los cursos reportados",
            "positivos": " ".join(dict.fromkeys(positivos)),
            "inconvenientes": " ".join(dict.fromkeys(inconvenientes)),
            "propuestas": " ".join(dict.fromkeys(propuestas)),
        })
        return bloques

    for idx, row in cursos.iterrows():
        key = curso_key(row, idx)
        nombre = str(row.get("Asignatura", "")).strip() or "Curso sin nombre"
        codigo = str(row.get("Código", "")).strip()
        grupo = str(row.get("Grupo", "")).strip()
        detalles = []
        if codigo:
            detalles.append(f"Código {codigo}")
        if grupo:
            detalles.append(f"Grupo {grupo}")
        encabezado = f"{nombre} ({' - '.join(detalles)})" if detalles else nombre
        data = analisis_por_curso.get(key, {})
        bloques.append({
            "titulo": encabezado,
            "positivos": data.get("positivos", ""),
            "inconvenientes": data.get("inconvenientes", ""),
            "propuestas": data.get("propuestas", ""),
        })
    return bloques


def crear_informe_gc72_docx(docente: str, periodo: str, fecha_entrega: date | str, cursos: pd.DataFrame, bloques_analisis: List[Dict[str, str]]) -> bytes:
    if not TEMPLATE_GC72.exists():
        raise FileNotFoundError(f"No se encontró la plantilla: {TEMPLATE_GC72}")
    doc = Document(str(TEMPLATE_GC72))
    cursos = normalizar_dataframe_gc72(cursos, calcular_porcentajes=False)

    datos = doc.tables[0]
    set_label_value(datos.rows[0].cells[0], "DOCENTE:", docente)
    set_label_value(datos.rows[1].cells[0], "PERÍODO ACADÉMICO:", periodo)
    fecha_texto = fecha_entrega.strftime("%d/%m/%Y") if hasattr(fecha_entrega, "strftime") else str(fecha_entrega)
    set_label_value(datos.rows[2].cells[0], "FECHA DE ENTREGA:", fecha_texto)

    tabla = doc.tables[1]
    fijar_anchos_tabla(tabla, [0.62, 0.55, 1.38, 0.72, 0.68, 0.90, 0.43, 0.38, 0.55, 0.38, 0.55, 0.40, 0.50, 0.38, 0.50, 0.38])
    fila_inicio = 3
    filas_necesarias = max(8, len(cursos))
    while len(tabla.rows) < fila_inicio + filas_necesarias:
        tabla.add_row()

    for i in range(filas_necesarias):
        row_cells = tabla.rows[fila_inicio + i].cells
        if i < len(cursos):
            registro = cursos.iloc[i]
            for j, col in enumerate(MAPA_TABLA_WORD_GC72):
                valor = registro.get(col, "")
                if col in NUMERICAS_ENTERAS_GC72:
                    texto = formato_entero(valor)
                elif col in NUMERICAS_PORCENTAJE_GC72:
                    texto = formato_porcentaje(valor)
                    if j >= 7:
                        texto = texto.replace("%", "")
                else:
                    texto = "" if pd.isna(valor) else str(valor).strip()
                set_cell_text(row_cells[j], texto, font_size=6.8 if j >= 6 else 7.2)
        else:
            for j in range(min(len(row_cells), len(MAPA_TABLA_WORD_GC72))):
                set_cell_text(row_cells[j], "", font_size=7.5)

    firma = None
    for p in doc.paragraphs:
        if p.text.strip().startswith("Firma:"):
            firma = p
            break
    if firma is None:
        firma = doc.add_paragraph("Firma: ___________________")

    agregar_parrafo_antes(firma, "")
    for bloque in bloques_analisis:
        titulo = bloque.get("titulo", "Curso")
        p_titulo = firma.insert_paragraph_before()
        p_titulo.paragraph_format.space_before = Pt(4)
        p_titulo.paragraph_format.space_after = Pt(2)
        run = p_titulo.add_run(titulo)
        run.bold = True
        run.font.size = Pt(10)

        positivos = bloque.get("positivos", "").strip() or "Sin observaciones específicas para este apartado."
        inconvenientes = bloque.get("inconvenientes", "").strip() or "No se reportan inconvenientes relevantes para el periodo informado."
        propuestas = bloque.get("propuestas", "").strip() or "Mantener seguimiento y retroalimentación permanente para sostener el avance del curso."

        agregar_parrafo_antes(firma, f"1. Aspectos positivos: {positivos}", bold_prefix="1. Aspectos positivos:")
        agregar_parrafo_antes(firma, f"2. Inconvenientes presentados: {inconvenientes}", bold_prefix="2. Inconvenientes presentados:")
        agregar_parrafo_antes(firma, f"3. Propuestas metodológicas: {propuestas}", bold_prefix="3. Propuestas metodológicas:")
        agregar_parrafo_antes(firma, "")

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# -----------------------------------------------------------------------------
# Listados, calificaciones y resumen académico
# -----------------------------------------------------------------------------
def leer_tabla_excel(uploaded_file) -> pd.DataFrame:
    if uploaded_file is None:
        return pd.DataFrame()
    nombre = uploaded_file.name.lower()
    if nombre.endswith(".csv"):
        return pd.read_csv(uploaded_file)
    if nombre.endswith((".xlsx", ".xls")):
        return pd.read_excel(uploaded_file, sheet_name=0, header=None)
    raise ValueError("Formato no soportado. Use .xls, .xlsx o .csv.")


def detectar_fila_encabezado(df: pd.DataFrame, palabras_clave: List[str]) -> int:
    claves = [normalizar_texto(p) for p in palabras_clave]
    for idx, row in df.iterrows():
        texto = " | ".join(normalizar_texto(v) for v in row.values if not pd.isna(v))
        coincidencias = sum(1 for clave in claves if clave in texto)
        if coincidencias >= max(1, min(2, len(claves))):
            return int(idx)
    return 0


def normalizar_columnas_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    return df


def leer_listado_estudiantes(uploaded_file) -> pd.DataFrame:
    if uploaded_file is None:
        return pd.DataFrame(columns=["Nombre completo", "Documento", "Correo", "Plan", "Observación", "Estado"])
    nombre = uploaded_file.name.lower()
    if nombre.endswith(".csv"):
        raw = pd.read_csv(uploaded_file, header=None)
    else:
        raw = pd.read_excel(uploaded_file, sheet_name=0, header=None)
    header_row = detectar_fila_encabezado(raw, ["NOMBRE COMPLETO", "DOCUMENTO", "CORREO"])
    df = raw.iloc[header_row + 1:].copy()
    df.columns = [str(c).strip() for c in raw.iloc[header_row].tolist()]
    df = normalizar_columnas_df(df)
    rename = {}
    for col in df.columns:
        n = normalizar_texto(col)
        if "NOMBRE" in n:
            rename[col] = "Nombre completo"
        elif "DOCUMENTO" in n or "CEDULA" in n or "CARN" in n:
            rename[col] = "Documento"
        elif "CORREO" in n or "EMAIL" in n:
            rename[col] = "Correo"
        elif n == "PLAN" or "PROGRAMA" in n:
            rename[col] = "Plan"
        elif "OBSERV" in n:
            rename[col] = "Observación"
    df = df.rename(columns=rename)
    for col in ["Nombre completo", "Documento", "Correo", "Plan", "Observación"]:
        if col not in df.columns:
            df[col] = ""
    df = df[["Nombre completo", "Documento", "Correo", "Plan", "Observación"]]
    df = df.dropna(how="all").copy()
    df = df[df["Nombre completo"].astype(str).str.strip().ne("")].reset_index(drop=True)
    df["Documento"] = df["Documento"].apply(lambda x: "" if pd.isna(x) else str(x).replace(".0", "").strip())
    df["Estado"] = df["Observación"].apply(lambda x: "Desertó" if re.search(r"DESERT|RETI|CANCEL", normalizar_texto(x)) else "Activo")
    return df


def leer_calificaciones(uploaded_file) -> Tuple[pd.DataFrame, pd.DataFrame]:
    if uploaded_file is None:
        return pd.DataFrame(), pd.DataFrame()
    nombre = uploaded_file.name.lower()
    if not nombre.endswith((".xlsx", ".xls", ".csv")):
        raise ValueError("Formato de calificaciones no soportado.")
    if nombre.endswith(".csv"):
        return pd.read_csv(uploaded_file), pd.DataFrame()
    xls = pd.ExcelFile(uploaded_file)
    hoja_cal = None
    hoja_eval = None
    for sheet in xls.sheet_names:
        n = normalizar_texto(sheet)
        if "CALIFIC" in n:
            hoja_cal = sheet
        if "EVALU" in n:
            hoja_eval = sheet
    if hoja_cal is None:
        hoja_cal = xls.sheet_names[0]
    cal = pd.read_excel(xls, sheet_name=hoja_cal)
    ev = pd.read_excel(xls, sheet_name=hoja_eval) if hoja_eval else pd.DataFrame()
    return normalizar_columnas_df(cal), normalizar_columnas_df(ev)


def encontrar_columna(df: pd.DataFrame, patrones: List[str]) -> Optional[str]:
    patrones_norm = [normalizar_texto(p) for p in patrones]
    for col in df.columns:
        n = normalizar_texto(col)
        if any(p in n for p in patrones_norm):
            return col
    return None


def contar_por_estado_y_nota(df: pd.DataFrame, columna_nota: Optional[str], corte_aprobacion: float, estado_col: Optional[str]) -> Tuple[int, int]:
    if df.empty or not columna_nota or columna_nota not in df.columns:
        return 0, 0
    notas = pd.to_numeric(df[columna_nota], errors="coerce")
    activo = pd.Series([True] * len(df), index=df.index)
    if estado_col and estado_col in df.columns:
        activo = ~df[estado_col].astype(str).apply(lambda x: bool(re.search(r"DESERT|RETI|CANCEL", normalizar_texto(x))))
    validas = notas.notna() & activo
    aprueban = int((notas[validas] >= corte_aprobacion).sum())
    reprueban = int((notas[validas] < corte_aprobacion).sum())
    return aprueban, reprueban


def resumen_gc72_desde_archivos(
    listado_df: pd.DataFrame,
    calificaciones_df: pd.DataFrame,
    evaluaciones_df: pd.DataFrame,
    codigo: str,
    grupo: str,
    asignatura: str,
    avance_contenido: float,
    porcentaje_evaluado_manual: float,
    corte_aprobacion: float = 3.0,
) -> pd.DataFrame:
    matriculados = int(len(listado_df)) if not listado_df.empty else int(len(calificaciones_df))
    estado_col_listado = "Estado" if "Estado" in listado_df.columns else None
    desertaron = 0
    if not listado_df.empty and estado_col_listado:
        desertaron = int(listado_df[estado_col_listado].astype(str).apply(lambda x: bool(re.search(r"DESERT|RETI|CANCEL", normalizar_texto(x)))).sum())
    estado_col = encontrar_columna(calificaciones_df, ["Estado", "Observación", "Observacion"])
    parcial_col = encontrar_columna(calificaciones_df, ["Nota parcial", "Parcial", "Corte 1", "Primer corte"])
    acumulada_col = encontrar_columna(calificaciones_df, ["Nota acumulada", "Acumulada", "Definitiva", "Nota final", "Final"])
    if acumulada_col is None:
        acumulada_col = parcial_col
    apr_par, rep_par = contar_por_estado_y_nota(calificaciones_df, parcial_col, corte_aprobacion, estado_col)
    apr_fecha, rep_fecha = contar_por_estado_y_nota(calificaciones_df, acumulada_col, corte_aprobacion, estado_col)

    # Si la plantilla trae hoja de evaluaciones, se usa para sugerir el % evaluado.
    porcentaje_evaluado = porcentaje_evaluado_manual
    if not evaluaciones_df.empty:
        col_valor = encontrar_columna(evaluaciones_df, ["Valor", "%"])
        if col_valor:
            valores = pd.to_numeric(evaluaciones_df[col_valor], errors="coerce").fillna(0)
            # Si aún no hay fecha o nota, el docente puede ajustar manualmente en la interfaz.
            porcentaje_evaluado = float(min(100, valores.sum())) if valores.sum() > 0 else porcentaje_evaluado_manual

    data = [{
        "Código": codigo,
        "Grupo": grupo,
        "Asignatura": asignatura,
        "% Avance en contenido": avance_contenido,
        "% Evaluado": porcentaje_evaluado,
        "Estudiantes matriculados": matriculados,
        "Desertaron N°": desertaron,
        "Desertaron %": porcentaje(desertaron, matriculados),
        "Aprueban evaluación parcial N°": apr_par,
        "Aprueban evaluación parcial %": porcentaje(apr_par, matriculados),
        "Reprueban evaluación parcial N°": rep_par,
        "Reprueban evaluación parcial %": porcentaje(rep_par, matriculados),
        "Aprueban a la fecha N°": apr_fecha,
        "Aprueban a la fecha %": porcentaje(apr_fecha, matriculados),
        "Reprueban a la fecha N°": rep_fecha,
        "Reprueban a la fecha %": porcentaje(rep_fecha, matriculados),
    }]
    return pd.DataFrame(data, columns=COLUMNAS_GC72)

# -----------------------------------------------------------------------------
# Planificación de clases y FD-GC71
# -----------------------------------------------------------------------------
def parse_time_value(value) -> Optional[time]:
    if isinstance(value, time):
        return value
    if isinstance(value, datetime):
        return value.time()
    if value is None or pd.isna(value):
        return None
    texto = str(value).strip()
    if not texto:
        return None
    for fmt in ["%H:%M", "%H:%M:%S", "%I:%M %p", "%I:%M%p"]:
        try:
            return datetime.strptime(texto.upper(), fmt).time()
        except ValueError:
            pass
    return None


def horas_entre(inicio: time, fin: time) -> float:
    di = datetime.combine(date.today(), inicio)
    df = datetime.combine(date.today(), fin)
    if df <= di:
        df += timedelta(days=1)
    return round((df - di).total_seconds() / 3600, 2)


def generar_fechas_clase(inicio: date, fin: date, horarios_df: pd.DataFrame, fechas_excluidas: Iterable[date] | None = None) -> pd.DataFrame:
    horarios = limpiar_df(horarios_df, COLUMNAS_HORARIOS)
    fechas_excluidas = set(fechas_excluidas or [])
    registros = []
    actual = inicio
    while actual <= fin:
        dia_nombre = DIAS_INV.get(actual.weekday(), "")
        for _, h in horarios.iterrows():
            dia = str(h.get("Día", "")).strip()
            if DIAS_MAP.get(dia, -1) == actual.weekday() and actual not in fechas_excluidas:
                hi = parse_time_value(h.get("Hora inicio")) or time(0, 0)
                hf = parse_time_value(h.get("Hora fin")) or time(0, 0)
                registros.append({
                    "Fecha": actual,
                    "Día": dia_nombre,
                    "Hora inicio": hi.strftime("%H:%M") if hi else "",
                    "Hora fin": hf.strftime("%H:%M") if hf else "",
                    "Horas": horas_entre(hi, hf) if hi and hf else 0,
                    "Lugar / ambiente": h.get("Lugar / ambiente", ""),
                })
        actual += timedelta(days=1)
    df = pd.DataFrame(registros)
    if df.empty:
        return pd.DataFrame(columns=["Fecha", "Día", "Hora inicio", "Hora fin", "Horas", "Lugar / ambiente"])
    df["orden"] = df["Fecha"].astype(str) + " " + df["Hora inicio"].astype(str)
    df = df.sort_values("orden").drop(columns=["orden"]).reset_index(drop=True)
    return df


def expandir_plan_sesiones(modulos_df: pd.DataFrame, fechas_clase_df: pd.DataFrame, criterio: str = "Horas presenciales") -> pd.DataFrame:
    modulos = limpiar_df(modulos_df, COLUMNAS_MODULOS)
    registros = []
    idx_fecha = 0
    n_sesion = 1
    for _, mod in modulos.iterrows():
        unidad = str(mod.get("Unidad", "")).strip() or f"Unidad {len(registros)+1}"
        contenido = str(mod.get("Contenido / tema central", "")).strip()
        trabajo_p = str(mod.get("Trabajo presencial", "")).strip()
        trabajo_i = str(mod.get("Trabajo independiente", "")).strip()
        horas_objetivo = limpiar_numero(mod.get("Horas presenciales")) or 0
        sesiones_objetivo = limpiar_numero(mod.get("Sesiones"))
        if criterio == "Sesiones" and sesiones_objetivo:
            n_bloques = int(max(1, round(sesiones_objetivo)))
        else:
            if fechas_clase_df.empty:
                n_bloques = int(max(1, math.ceil(horas_objetivo / 2))) if horas_objetivo else 1
            else:
                acumuladas = 0.0
                n_bloques = 0
                tmp_idx = idx_fecha
                while acumuladas < max(0.01, horas_objetivo) and tmp_idx < len(fechas_clase_df):
                    acumuladas += limpiar_numero(fechas_clase_df.iloc[tmp_idx].get("Horas")) or 0
                    n_bloques += 1
                    tmp_idx += 1
                n_bloques = max(1, n_bloques)
        for parte in range(1, n_bloques + 1):
            if idx_fecha < len(fechas_clase_df):
                f = fechas_clase_df.iloc[idx_fecha]
                fecha = f.get("Fecha")
                horario = f"{f.get('Hora inicio', '')} - {f.get('Hora fin', '')}".strip(" -")
                lugar = f.get("Lugar / ambiente", "")
            else:
                fecha = "Por programar"
                horario = "Por programar"
                lugar = ""
            sufijo = f" (parte {parte} de {n_bloques})" if n_bloques > 1 else ""
            registros.append({
                "Unidad": unidad,
                "N° sesión": n_sesion,
                "Fecha": fecha,
                "Horario": horario,
                "Contenido por desarrollar": f"{contenido}{sufijo}",
                "Descripción del trabajo presencial": trabajo_p,
                "Descripción trabajo independiente": trabajo_i,
                "Lugar / ambiente": lugar,
            })
            idx_fecha += 1
            n_sesion += 1
    return pd.DataFrame(registros, columns=COLUMNAS_SESIONES)


def configurar_documento_gc71(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(8.5)


def agregar_tabla_header_gc71(doc: Document):
    table = doc.add_table(rows=2, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    aplicar_bordes_tabla(table)
    widths = [2.55, 3.9, 1.75]
    for row in table.rows:
        for i, w in enumerate(widths):
            set_cell_width(row.cells[i], w)
    logo_cell = table.cell(0, 0).merge(table.cell(1, 0))
    title_cell = table.cell(0, 1).merge(table.cell(1, 1))
    if LOGO_POLI.exists():
        logo_cell.text = ""
        p = logo_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(LOGO_POLI), width=Inches(1.55))
    else:
        set_cell_text(logo_cell, "POLITÉCNICO COLOMBIANO\nJAIME ISAZA CADAVID", 8, True)
    set_cell_text(title_cell, "GUÍA DIDÁCTICA DE ASIGNATURA Y\nCONCERTACIÓN DE EVALUACIÓN", 12, True)
    set_cell_text(table.cell(0, 2), "Código: FD-GC71", 10, False)
    set_cell_text(table.cell(1, 2), "Versión: 09", 10, True)
    return table


def agregar_fila_seccion(table, titulo: str, cols: int, fill="FFF2CC"):
    row = table.add_row()
    cell = row.cells[0].merge(row.cells[cols - 1])
    shade_cell(cell, fill)
    set_cell_text(cell, titulo, 8.5, True)
    return row


def agregar_tabla_identificacion_gc71(doc: Document, datos: Dict[str, str]):
    doc.add_paragraph()
    t = doc.add_table(rows=0, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    aplicar_bordes_tabla(t)
    widths = [2.2, 2.05, 1.85, 2.0]
    for _ in range(1):
        pass
    agregar_fila_seccion(t, "1. IDENTIFICACIÓN DE LA ASIGNATURA", 4)
    campos = [
        ("PROGRAMA ACADÉMICO", datos.get("programa", "")),
        ("ASIGNATURA", datos.get("asignatura", "")),
        ("CÓDIGO", datos.get("codigo", "")),
        ("ÁREA DE FORMACIÓN", datos.get("area", "")),
        ("PRERREQUISITO(S)", datos.get("prerrequisitos", "")),
        ("CORREQUISITO(S)", datos.get("correquisitos", "")),
    ]
    for etiqueta, valor in campos:
        row = t.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        c1 = row.cells[2].merge(row.cells[3])
        set_cell_text(c0, etiqueta, 8.3, True, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(c1, valor, 8.3, False, align=WD_ALIGN_PARAGRAPH.LEFT)
    row = t.add_row()
    set_cell_text(row.cells[0], "TIPO DE ASIGNATURA", 8.3, True, align=WD_ALIGN_PARAGRAPH.LEFT)
    tipo = datos.get("tipo_asignatura", "Teórico-práctica")
    set_cell_text(row.cells[1], f"{'☒' if tipo == 'Teórica' else '☐'} Teórica", 8.3)
    set_cell_text(row.cells[2], f"{'☒' if tipo == 'Teórico-práctica' else '☐'} Teórico-práctica", 8.3)
    set_cell_text(row.cells[3], f"{'☒' if tipo == 'Práctica' else '☐'} Práctica", 8.3)
    row = t.add_row()
    set_cell_text(row.cells[0].merge(row.cells[1]), "NÚMERO DE CRÉDITOS", 8.3, True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[2].merge(row.cells[3]), datos.get("creditos", ""), 8.3, align=WD_ALIGN_PARAGRAPH.LEFT)
    row = t.add_row()
    set_cell_text(row.cells[0], "DISTRIBUCIÓN HORARIA SEMANAL", 8.3, True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[1], f"HTP: {datos.get('htp', '')}", 8.3)
    set_cell_text(row.cells[2], f"HTI: {datos.get('hti', '')}", 8.3)
    set_cell_text(row.cells[3], f"Total: {datos.get('ht_total', '')}", 8.3)
    for etiqueta, valor in [
        ("PROFESOR", datos.get("profesor", "")),
        ("CORREO ELECTRÓNICO", datos.get("correo", "")),
        ("GRUPO", datos.get("grupo", "")),
        ("PERÍODO ACADÉMICO", datos.get("periodo", "")),
    ]:
        row = t.add_row()
        c0 = row.cells[0].merge(row.cells[1])
        c1 = row.cells[2].merge(row.cells[3])
        set_cell_text(c0, etiqueta, 8.3, True, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(c1, valor, 8.3, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in t.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                set_cell_width(row.cells[i], w)
    return t


def agregar_seccion_texto_gc71(doc: Document, numero_titulo: str, texto: str):
    t = doc.add_table(rows=2, cols=1)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    aplicar_bordes_tabla(t)
    shade_cell(t.rows[0].cells[0], "FFF2CC")
    set_cell_text(t.rows[0].cells[0], numero_titulo, 8.5, True)
    set_cell_text(t.rows[1].cells[0], texto, 8.2, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_width(t.rows[0].cells[0], 8.1)
    set_cell_width(t.rows[1].cells[0], 8.1)
    return t


def agregar_contenidos_gc71(doc: Document, sesiones_df: pd.DataFrame):
    t = doc.add_table(rows=0, cols=5)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    aplicar_bordes_tabla(t)
    widths = [0.75, 1.0, 2.45, 1.95, 1.95]
    agregar_fila_seccion(t, "7. CONTENIDOS TEMÁTICOS DE LA ASIGNATURA", 5)
    sesiones = limpiar_df(sesiones_df, COLUMNAS_SESIONES)
    for unidad, dfu in sesiones.groupby("Unidad", sort=False):
        agregar_fila_seccion(t, str(unidad).upper(), 5, fill="FFF2CC")
        row = t.add_row()
        headers = ["N° sesión", "Fecha", "Contenido por desarrollar", "Descripción del trabajo presencial", "Descripción trabajo independiente"]
        for i, h in enumerate(headers):
            shade_cell(row.cells[i], "F2F2F2")
            set_cell_text(row.cells[i], h, 7.5, True)
        for _, s in dfu.iterrows():
            row = t.add_row()
            fecha = s.get("Fecha")
            if hasattr(fecha, "strftime"):
                fecha_txt = fecha.strftime("%d/%m/%Y")
            else:
                fecha_txt = str(fecha)
            values = [
                s.get("N° sesión", ""),
                fecha_txt,
                s.get("Contenido por desarrollar", ""),
                s.get("Descripción del trabajo presencial", ""),
                s.get("Descripción trabajo independiente", ""),
            ]
            for i, v in enumerate(values):
                set_cell_text(row.cells[i], v, 7.2 if i >= 2 else 7.0, align=WD_ALIGN_PARAGRAPH.LEFT if i >= 2 else WD_ALIGN_PARAGRAPH.CENTER)
    for row in t.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                set_cell_width(row.cells[i], w)
    return t


def agregar_evaluaciones_gc71(doc: Document, asignatura: str, grupo: str, evaluaciones_df: pd.DataFrame):
    t = doc.add_table(rows=0, cols=4)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    aplicar_bordes_tabla(t)
    widths = [1.45, 4.15, 0.9, 1.6]
    agregar_fila_seccion(t, "11. EVALUACIÓN DE LA ASIGNATURA", 4)
    row = t.add_row()
    c0 = row.cells[0].merge(row.cells[1])
    c1 = row.cells[2].merge(row.cells[3])
    set_cell_text(c0, f"ASIGNATURA: {asignatura}", 7.6, True, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(c1, f"GRUPO: {grupo}", 7.6, True, align=WD_ALIGN_PARAGRAPH.LEFT)
    row = t.add_row()
    headers = ["TIPO DE EVALUACIÓN°", "PROCEDIMIENTO DE EVALUACIÓN\n(Descripción de la actividad evaluativa)", "VALOR (%)", "FECHA DE\nREALIZACIÓN"]
    for i, h in enumerate(headers):
        shade_cell(row.cells[i], "FFF2CC")
        set_cell_text(row.cells[i], h, 7.2, True)
    evaluaciones = limpiar_df(evaluaciones_df, COLUMNAS_EVALUACIONES)
    for _, ev in evaluaciones.iterrows():
        row = t.add_row()
        fecha = ev.get("Fecha de realización")
        if hasattr(fecha, "strftime"):
            fecha_txt = fecha.strftime("%d/%m/%Y")
        else:
            fecha_txt = str(fecha)
        vals = [ev.get("Tipo de evaluación", ""), ev.get("Procedimiento de evaluación", ""), ev.get("Valor (%)", ""), fecha_txt]
        for i, v in enumerate(vals):
            set_cell_text(row.cells[i], v, 7.2, align=WD_ALIGN_PARAGRAPH.LEFT if i in [0, 1] else WD_ALIGN_PARAGRAPH.CENTER)
    for row in t.rows:
        for i, w in enumerate(widths):
            if i < len(row.cells):
                set_cell_width(row.cells[i], w)
    return t


def agregar_evidencia_y_control_gc71(doc: Document, datos: Dict[str, str], representantes_df: pd.DataFrame):
    t = doc.add_table(rows=0, cols=3)
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    aplicar_bordes_tabla(t)
    agregar_fila_seccion(t, "12. EVIDENCIA DE PRESENTACIÓN DE LA GUÍA Y CONCERTACIÓN DE EVALUACIÓN AL GRUPO DE ESTUDIANTES", 3)
    row = t.add_row()
    cell = row.cells[0].merge(row.cells[2])
    set_cell_text(cell, "Se deja constancia de socialización de la Guía Didáctica de Asignatura y aprobación de la concertación de evaluación según el reglamento estudiantil; para ello firman tres estudiantes en representación del grupo:", 7.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    row = t.add_row()
    for i, h in enumerate(["Nombre de los estudiantes", "N° de cédula o carné estudiantil", "Firma"]):
        shade_cell(row.cells[i], "FFF2CC")
        set_cell_text(row.cells[i], h, 7.5, True)
    reps = representantes_df.copy() if representantes_df is not None else pd.DataFrame()
    for i in range(3):
        row = t.add_row()
        nombre = reps.iloc[i].get("Nombre", "") if i < len(reps) else ""
        docu = reps.iloc[i].get("Documento", "") if i < len(reps) else ""
        set_cell_text(row.cells[0], nombre, 7.5, align=WD_ALIGN_PARAGRAPH.LEFT)
        set_cell_text(row.cells[1], docu, 7.5)
        set_cell_text(row.cells[2], "", 7.5)
    row = t.add_row()
    for i, h in enumerate(["Nombre del docente del curso", "Cédula", "Firma"]):
        shade_cell(row.cells[i], "FFF2CC")
        set_cell_text(row.cells[i], h, 7.5, True)
    row = t.add_row()
    set_cell_text(row.cells[0], datos.get("profesor", ""), 7.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(row.cells[1], datos.get("cedula_docente", ""), 7.5)
    set_cell_text(row.cells[2], "", 7.5)
    row = t.add_row()
    cell = row.cells[0].merge(row.cells[2])
    set_cell_text(cell, f"Fecha de socialización de la Guía Didáctica: {datos.get('fecha_socializacion', '')}", 7.5, True, align=WD_ALIGN_PARAGRAPH.LEFT)
    row = t.add_row()
    cell = row.cells[0].merge(row.cells[2])
    set_cell_text(cell, "Nota: El docente se compromete a devolver las evaluaciones, socializar la calificación con los estudiantes y a ingresar dicha calificación al sistema académico, correcta y oportunamente.", 7.2, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    for row in t.rows:
        set_cell_width(row.cells[0], 3.4)
        set_cell_width(row.cells[1], 2.0)
        set_cell_width(row.cells[2], 2.7)

    doc.add_paragraph()
    c = doc.add_table(rows=3, cols=2)
    c.alignment = WD_TABLE_ALIGNMENT.CENTER
    aplicar_bordes_tabla(c)
    cell = c.rows[0].cells[0].merge(c.rows[0].cells[1])
    shade_cell(cell, "FFF2CC")
    set_cell_text(cell, "CONTROL DE CAMBIOS Y VIGENCIA (DILIGENCIAR LOS DATOS ESPECÍFICOS)", 8, True)
    set_cell_text(c.rows[1].cells[0], "Fecha de Revisión por parte del Coordinador de Área:", 7.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(c.rows[1].cells[1], datos.get("fecha_revision", ""), 7.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(c.rows[2].cells[0], "Fecha de aprobación y acta de sesión del Comité de currículo del programa:", 7.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    set_cell_text(c.rows[2].cells[1], datos.get("fecha_aprobacion", ""), 7.5, align=WD_ALIGN_PARAGRAPH.LEFT)
    for row in c.rows:
        set_cell_width(row.cells[0], 4.0)
        set_cell_width(row.cells[1], 4.1)


def crear_gc71_docx(
    datos: Dict[str, str],
    sesiones_df: pd.DataFrame,
    evaluaciones_df: pd.DataFrame,
    representantes_df: Optional[pd.DataFrame] = None,
) -> bytes:
    doc = Document()
    configurar_documento_gc71(doc)
    agregar_tabla_header_gc71(doc)
    agregar_tabla_identificacion_gc71(doc, datos)
    agregar_seccion_texto_gc71(doc, "2. JUSTIFICACIÓN", datos.get("justificacion", ""))
    agregar_seccion_texto_gc71(doc, "3. COMPETENCIAS A LAS QUE LE TRIBUTA LA ASIGNATURA", datos.get("competencias", ""))
    agregar_seccion_texto_gc71(doc, "4. RESULTADOS DE APRENDIZAJE A LOS QUE LE TRIBUTA LA ASIGNATURA", datos.get("resultados", ""))
    agregar_seccion_texto_gc71(doc, "5. OBJETIVOS DE APRENDIZAJE DE LA ASIGNATURA", f"OBJETIVO(S) GENERAL(ES)\n{datos.get('objetivo_general', '')}\n\nOBJETIVOS ESPECÍFICOS\n{datos.get('objetivos_especificos', '')}")
    agregar_seccion_texto_gc71(doc, "6. METODOLOGÍAS Y ESTRATEGIAS DIDÁCTICAS DE LA ASIGNATURA", datos.get("metodologias", ""))
    agregar_contenidos_gc71(doc, sesiones_df)
    agregar_seccion_texto_gc71(doc, "8. AMBIENTES DE APRENDIZAJE DE LA ASIGNATURA", datos.get("ambientes", ""))
    agregar_seccion_texto_gc71(doc, "9. MEDIOS EDUCATIVOS PARA LA ASIGNATURA", datos.get("medios", ""))
    agregar_seccion_texto_gc71(doc, "10. REFERENCIAS BIBLIOGRÁFICAS", datos.get("referencias", ""))
    agregar_evaluaciones_gc71(doc, datos.get("asignatura", ""), datos.get("grupo", ""), evaluaciones_df)
    agregar_evidencia_y_control_gc71(doc, datos, representantes_df if representantes_df is not None else pd.DataFrame())
    # Certificación en pie visual al final.
    if LOGO_ICONTEC.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.add_run().add_picture(str(LOGO_ICONTEC), width=Inches(1.0))
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# -----------------------------------------------------------------------------
# Excel de evaluación
# -----------------------------------------------------------------------------
def crear_plantilla_evaluacion_xlsx(
    estudiantes_df: pd.DataFrame,
    evaluaciones_df: pd.DataFrame,
    datos: Dict[str, str],
    corte_aprobacion: float = 3.0,
) -> bytes:
    import xlsxwriter

    output = io.BytesIO()
    workbook = xlsxwriter.Workbook(output, {"in_memory": True})
    fmt_title = workbook.add_format({"bold": True, "font_size": 14, "align": "center", "valign": "vcenter", "bg_color": "#D9EAF7", "border": 1})
    fmt_header = workbook.add_format({"bold": True, "bg_color": "#FFF2CC", "border": 1, "align": "center", "valign": "vcenter", "text_wrap": True})
    fmt_cell = workbook.add_format({"border": 1, "valign": "vcenter"})
    fmt_num = workbook.add_format({"border": 1, "num_format": "0.00", "valign": "vcenter"})
    fmt_pct = workbook.add_format({"border": 1, "num_format": "0%", "valign": "vcenter"})
    fmt_note = workbook.add_format({"border": 1, "font_color": "#666666", "italic": True, "text_wrap": True})

    # Configuración
    ws = workbook.add_worksheet("Configuración")
    ws.merge_range("A1:D1", "Plantilla de evaluación generada desde FD-GC71", fmt_title)
    config_rows = [
        ("Programa", datos.get("programa", "")),
        ("Asignatura", datos.get("asignatura", "")),
        ("Código", datos.get("codigo", "")),
        ("Grupo", datos.get("grupo", "")),
        ("Periodo académico", datos.get("periodo", "")),
        ("Docente", datos.get("profesor", "")),
        ("Nota mínima aprobatoria", corte_aprobacion),
    ]
    ws.write_row("A3", ["Campo", "Valor"], fmt_header)
    for r, (k, v) in enumerate(config_rows, start=3):
        ws.write(r, 0, k, fmt_cell)
        ws.write(r, 1, v, fmt_num if isinstance(v, (int, float)) else fmt_cell)
    ws.set_column("A:A", 26)
    ws.set_column("B:D", 32)

    # Estudiantes
    est = estudiantes_df.copy()
    if est.empty:
        est = pd.DataFrame(columns=["Nombre completo", "Documento", "Correo", "Plan", "Observación", "Estado"])
    for col in ["Nombre completo", "Documento", "Correo", "Plan", "Observación", "Estado"]:
        if col not in est.columns:
            est[col] = ""
    est = est[["Nombre completo", "Documento", "Correo", "Plan", "Observación", "Estado"]]
    ws = workbook.add_worksheet("Estudiantes")
    ws.write_row(0, 0, est.columns.tolist(), fmt_header)
    for r, row in est.iterrows():
        for c, col in enumerate(est.columns):
            ws.write(r + 1, c, "" if pd.isna(row[col]) else row[col], fmt_cell)
    ws.autofilter(0, 0, max(1, len(est)), len(est.columns) - 1)
    ws.freeze_panes(1, 0)
    ws.set_column("A:A", 34)
    ws.set_column("B:B", 16)
    ws.set_column("C:C", 34)
    ws.set_column("D:D", 42)
    ws.set_column("E:F", 18)

    # Evaluaciones
    ev = limpiar_df(evaluaciones_df, COLUMNAS_EVALUACIONES)
    ws = workbook.add_worksheet("Evaluaciones")
    ws.write_row(0, 0, COLUMNAS_EVALUACIONES, fmt_header)
    total_valor = 0.0
    for r, row in ev.iterrows():
        for c, col in enumerate(COLUMNAS_EVALUACIONES):
            value = row.get(col, "")
            if col == "Fecha de realización" and hasattr(value, "strftime"):
                value = value.strftime("%d/%m/%Y")
            if col == "Valor (%)":
                val = limpiar_numero(value) or 0
                total_valor += val
                ws.write_number(r + 1, c, val, fmt_num)
            else:
                ws.write(r + 1, c, "" if pd.isna(value) else value, fmt_cell)
    ws.write(len(ev) + 2, 1, "Total concertado", fmt_header)
    ws.write_formula(len(ev) + 2, 2, f"=SUM(C2:C{len(ev)+1})", fmt_num)
    ws.write(len(ev) + 4, 0, "Nota", fmt_note)
    ws.merge_range(len(ev) + 4, 1, len(ev) + 4, 5, "El total de Valor (%) debe sumar 100. Las columnas de notas se generan automáticamente en la hoja Calificaciones.", fmt_note)
    ws.freeze_panes(1, 0)
    ws.set_column("A:A", 20)
    ws.set_column("B:B", 52)
    ws.set_column("C:C", 12)
    ws.set_column("D:F", 18)

    # Calificaciones
    ws = workbook.add_worksheet("Calificaciones")
    fixed_headers = ["Nombre completo", "Documento", "Correo", "Estado"]
    eval_headers = [f"E{i+1} - {str(row.get('Tipo de evaluación', 'Evaluación')).strip()[:25]}" for i, (_, row) in enumerate(ev.iterrows())]
    headers = fixed_headers + eval_headers + ["Nota parcial", "Nota acumulada", "Aprueba a la fecha", "Observación docente"]
    ws.write_row(0, 0, headers, fmt_header)
    n = max(len(est), 1)
    for r in range(n):
        if r < len(est):
            ws.write(r + 1, 0, est.iloc[r].get("Nombre completo", ""), fmt_cell)
            ws.write(r + 1, 1, est.iloc[r].get("Documento", ""), fmt_cell)
            ws.write(r + 1, 2, est.iloc[r].get("Correo", ""), fmt_cell)
            ws.write(r + 1, 3, est.iloc[r].get("Estado", "Activo"), fmt_cell)
        else:
            for c in range(4):
                ws.write(r + 1, c, "", fmt_cell)
        # notas por evaluación
        for c in range(len(eval_headers)):
            ws.write_blank(r + 1, 4 + c, None, fmt_num)
        first_eval_col = 4
        last_eval_col = 4 + len(eval_headers) - 1
        # Weighted formulas. C in Evaluaciones is Valor (%)
        excel_row = r + 2
        if eval_headers:
            weighted_terms = []
            for i in range(len(eval_headers)):
                col_letter = xlsxwriter.utility.xl_col_to_name(first_eval_col + i)
                peso_cell = f"Evaluaciones!$C${i+2}"
                weighted_terms.append(f"IF(ISNUMBER({col_letter}{excel_row}),{col_letter}{excel_row}*{peso_cell}/100,0)")
            formula = "=" + "+".join(weighted_terms)
            ws.write_formula(r + 1, last_eval_col + 1, formula, fmt_num)  # Nota parcial = acumulada sugerida
            ws.write_formula(r + 1, last_eval_col + 2, formula, fmt_num)
            ws.write_formula(r + 1, last_eval_col + 3, f'=IF({xlsxwriter.utility.xl_col_to_name(last_eval_col + 2)}{excel_row}>=Configuración!$B$10,"Sí","No")', fmt_cell)
        else:
            ws.write_blank(r + 1, 4, None, fmt_num)
    ws.freeze_panes(1, 4)
    ws.autofilter(0, 0, n, len(headers) - 1)
    ws.set_column("A:A", 34)
    ws.set_column("B:B", 16)
    ws.set_column("C:C", 32)
    ws.set_column("D:D", 14)
    if eval_headers:
        ws.set_column(4, 4 + len(eval_headers) - 1, 14)
        ws.set_column(4 + len(eval_headers), len(headers) - 1, 18)
    else:
        ws.set_column("E:H", 16)

    # Resumen para FD-GC72
    ws = workbook.add_worksheet("Resumen FD-GC72")
    resumen_headers = ["Indicador", "Valor", "Observación"]
    ws.write_row(0, 0, resumen_headers, fmt_header)
    cal_sheet = "Calificaciones"
    nota_parcial_col = xlsxwriter.utility.xl_col_to_name(4 + len(eval_headers)) if eval_headers else "E"
    nota_acum_col = xlsxwriter.utility.xl_col_to_name(5 + len(eval_headers)) if eval_headers else "F"
    estado_col = "D"
    filas = [
        ("Código", datos.get("codigo", ""), ""),
        ("Grupo", datos.get("grupo", ""), ""),
        ("Asignatura", datos.get("asignatura", ""), ""),
        ("Estudiantes matriculados", f"=COUNTA({cal_sheet}!A2:A{n+1})", "Desde listado de clase"),
        ("Desertaron N°", f'=COUNTIF({cal_sheet}!{estado_col}2:{estado_col}{n+1},"*Desert*")+COUNTIF({cal_sheet}!{estado_col}2:{estado_col}{n+1},"*Retir*")+COUNTIF({cal_sheet}!{estado_col}2:{estado_col}{n+1},"*Cancel*")', "Según columna Estado"),
        ("Aprueban evaluación parcial N°", f'=COUNTIF({cal_sheet}!{nota_parcial_col}2:{nota_parcial_col}{n+1},">="&Configuración!$B$10)', "Nota parcial >= mínima"),
        ("Reprueban evaluación parcial N°", f'=COUNTIF({cal_sheet}!{nota_parcial_col}2:{nota_parcial_col}{n+1},"<"&Configuración!$B$10)', "Nota parcial < mínima"),
        ("Aprueban a la fecha N°", f'=COUNTIF({cal_sheet}!{nota_acum_col}2:{nota_acum_col}{n+1},">="&Configuración!$B$10)', "Nota acumulada >= mínima"),
        ("Reprueban a la fecha N°", f'=COUNTIF({cal_sheet}!{nota_acum_col}2:{nota_acum_col}{n+1},"<"&Configuración!$B$10)', "Nota acumulada < mínima"),
        ("% Evaluado sugerido", "=Evaluaciones!C" + str(len(ev) + 3), "Puede ajustarse en la app"),
        ("% Avance en contenido", "", "Diligenciar en la app con base en sesiones realizadas / sesiones planificadas"),
    ]
    for r, (ind, val, obs) in enumerate(filas, start=1):
        ws.write(r, 0, ind, fmt_cell)
        if isinstance(val, str) and val.startswith("="):
            ws.write_formula(r, 1, val, fmt_num)
        else:
            ws.write(r, 1, val, fmt_cell)
        ws.write(r, 2, obs, fmt_cell)
    ws.set_column("A:A", 34)
    ws.set_column("B:B", 18)
    ws.set_column("C:C", 62)

    workbook.close()
    output.seek(0)
    return output.getvalue()

# -----------------------------------------------------------------------------
# Interfaz Streamlit
# -----------------------------------------------------------------------------
def ui_gc71(st):
    st.header("FD-GC71 - Guía didáctica y concertación de evaluación")
    st.caption("Planifica el curso desde módulos, intensidad y horario real. La app arma las fechas y genera el formato FD-GC71.")
    usuario_actual = st.session_state.get("auth_user", {})

    with st.expander("1. Identificación de la asignatura", expanded=True):
        c1, c2, c3 = st.columns([1.2, 1.2, 0.8])
        with c1:
            programa = st.text_input("Programa académico", value="")
            asignatura = st.text_input("Asignatura", value="")
            codigo = st.text_input("Código", value="")
            area = st.text_input("Área de formación", value="")
        with c2:
            profesor = st.text_input("Profesor", value=usuario_actual.get("nombre_completo", ""))
            cedula_docente = st.text_input("Cédula docente", value="")
            correo = st.text_input("Correo electrónico", value=usuario_actual.get("email", ""))
            grupo = st.text_input("Grupo", value="")
        with c3:
            periodo = st.text_input("Periodo académico", value="2026-1")
            tipo = st.selectbox("Tipo de asignatura", ["Teórica", "Teórico-práctica", "Práctica"], index=1)
            creditos = st.number_input("Número de créditos", min_value=0, step=1, value=3)
            htp = st.number_input("HTP semanal", min_value=0.0, step=0.5, value=2.0)
            hti = st.number_input("HTI semanal", min_value=0.0, step=0.5, value=4.0)
        prerrequisitos = st.text_input("Prerrequisito(s)", value="Ninguno")
        correquisitos = st.text_input("Correquisito(s)", value="Ninguno")

    with st.expander("2. Textos académicos base", expanded=False):
        col1, col2 = st.columns(2)
        with col1:
            justificacion = st.text_area("Justificación", TEXTOS_PREDEFINIDOS_GC71["justificacion"], height=140)
            competencias = st.text_area("Competencias a las que tributa", TEXTOS_PREDEFINIDOS_GC71["competencias"], height=130)
            resultados = st.text_area("Resultados de aprendizaje", TEXTOS_PREDEFINIDOS_GC71["resultados"], height=130)
            objetivo_general = st.text_area("Objetivo general", TEXTOS_PREDEFINIDOS_GC71["objetivo_general"], height=110)
        with col2:
            objetivos_especificos = st.text_area("Objetivos específicos", TEXTOS_PREDEFINIDOS_GC71["objetivos_especificos"], height=140)
            metodologias = st.text_area("Metodologías y estrategias didácticas", TEXTOS_PREDEFINIDOS_GC71["metodologias"], height=130)
            ambientes = st.text_area("Ambientes de aprendizaje", TEXTOS_PREDEFINIDOS_GC71["ambientes"], height=100)
            medios = st.text_area("Medios educativos", TEXTOS_PREDEFINIDOS_GC71["medios"], height=100)
            referencias = st.text_area("Referencias bibliográficas", TEXTOS_PREDEFINIDOS_GC71["referencias"], height=100)

    st.subheader("3. Módulos / unidades e intensidad")
    st.write("Define la intensidad por unidad. Puedes trabajar por horas presenciales o por número exacto de sesiones.")
    criterio = st.radio("Criterio de expansión del cronograma", ["Horas presenciales", "Sesiones"], horizontal=True)
    default_modulos = pd.DataFrame([
        {"Unidad": "UNIDAD 1. Fundamentos", "Contenido / tema central": "Introducción, conceptos base y alcance de la asignatura", "Horas presenciales": 4, "Sesiones": 2, "Trabajo presencial": "Clase orientadora, explicación conceptual y taller diagnóstico.", "Trabajo independiente": "Lectura de apoyo y preparación de preguntas orientadoras."},
        {"Unidad": "UNIDAD 2. Aplicación", "Contenido / tema central": "Desarrollo de procedimientos, ejercicios y análisis de casos", "Horas presenciales": 8, "Sesiones": 4, "Trabajo presencial": "Taller aplicado, solución de ejercicios y discusión guiada.", "Trabajo independiente": "Desarrollo de actividad práctica y revisión bibliográfica."},
        {"Unidad": "UNIDAD 3. Integración", "Contenido / tema central": "Proyecto, socialización y retroalimentación", "Horas presenciales": 4, "Sesiones": 2, "Trabajo presencial": "Acompañamiento al proyecto y socialización de resultados.", "Trabajo independiente": "Ajuste de entregables y preparación de sustentación."},
    ])
    modulos_df = st.data_editor(
        default_modulos,
        num_rows="dynamic",
        hide_index=True,
        use_container_width=True,
        column_config={
            "Horas presenciales": st.column_config.NumberColumn("Horas presenciales", min_value=0.0, step=0.5),
            "Sesiones": st.column_config.NumberColumn("Sesiones", min_value=1, step=1),
        },
        key="modulos_gc71",
    )

    st.subheader("4. Horario de clase")
    c1, c2, c3 = st.columns(3)
    with c1:
        fecha_inicio = st.date_input("Fecha de inicio", value=date.today(), format="DD/MM/YYYY", key="gc71_inicio")
    with c2:
        fecha_fin = st.date_input("Fecha de finalización", value=date.today() + timedelta(days=112), format="DD/MM/YYYY", key="gc71_fin")
    with c3:
        fechas_no_clase_txt = st.text_area("Fechas sin clase (dd/mm/aaaa, una por línea)", value="", height=100)
    default_horarios = pd.DataFrame([
        {"Día": "Lunes", "Hora inicio": "18:00", "Hora fin": "20:00", "Lugar / ambiente": "Aula de clase"},
    ])
    horarios_df = st.data_editor(
        default_horarios,
        num_rows="dynamic",
        hide_index=True,
        use_container_width=True,
        column_config={"Día": st.column_config.SelectboxColumn("Día", options=["Lunes", "Martes", "Miércoles", "Jueves", "Viernes", "Sábado", "Domingo"]),},
        key="horarios_gc71",
    )
    fechas_excluidas = []
    for linea in fechas_no_clase_txt.splitlines():
        linea = linea.strip()
        if not linea:
            continue
        try:
            fechas_excluidas.append(datetime.strptime(linea, "%d/%m/%Y").date())
        except ValueError:
            st.warning(f"No pude interpretar esta fecha sin clase: {linea}. Use dd/mm/aaaa.")
    fechas_clase = generar_fechas_clase(fecha_inicio, fecha_fin, horarios_df, fechas_excluidas)
    sesiones_df = expandir_plan_sesiones(modulos_df, fechas_clase, criterio=criterio)

    with st.expander("Vista previa del cronograma generado", expanded=True):
        st.dataframe(sesiones_df, use_container_width=True)
        if len(sesiones_df) > len(fechas_clase):
            st.warning("Hay más sesiones requeridas que fechas disponibles. Las últimas sesiones quedan como 'Por programar'.")

    st.subheader("5. Evaluación concertada")
    default_evals = pd.DataFrame([
        {"Tipo de evaluación": "Parcial", "Procedimiento de evaluación": "Prueba escrita o práctica sobre los contenidos de las unidades desarrolladas.", "Valor (%)": 30, "Fecha de realización": sesiones_df.iloc[min(2, len(sesiones_df)-1)]["Fecha"] if not sesiones_df.empty else "", "Unidad relacionada": "UNIDAD 1", "Corte": "Parcial"},
        {"Tipo de evaluación": "Taller / Proyecto", "Procedimiento de evaluación": "Actividad aplicada con entrega de evidencias y criterios definidos en rúbrica.", "Valor (%)": 40, "Fecha de realización": sesiones_df.iloc[min(5, len(sesiones_df)-1)]["Fecha"] if not sesiones_df.empty else "", "Unidad relacionada": "UNIDAD 2", "Corte": "Seguimiento"},
        {"Tipo de evaluación": "Final", "Procedimiento de evaluación": "Socialización, sustentación o prueba final integradora.", "Valor (%)": 30, "Fecha de realización": sesiones_df.iloc[-1]["Fecha"] if not sesiones_df.empty else "", "Unidad relacionada": "UNIDAD 3", "Corte": "Final"},
    ])
    evaluaciones_df = st.data_editor(
        default_evals,
        num_rows="dynamic",
        hide_index=True,
        use_container_width=True,
        column_config={"Valor (%)": st.column_config.NumberColumn("Valor (%)", min_value=0, max_value=100, step=1)},
        key="evaluaciones_gc71",
    )
    total_eval = sum(limpiar_numero(v) or 0 for v in evaluaciones_df.get("Valor (%)", []))
    if round(total_eval, 2) != 100:
        st.warning(f"La evaluación suma {total_eval:.1f}%. El formato puede generarse, pero lo sano es que sume 100%. Aquí no maquillamos cadáveres.")
    else:
        st.success("La evaluación concertada suma 100%.")

    st.subheader("6. Listado tradicional de clase y plantilla de evaluación")
    listado_file = st.file_uploader("Cargar listado tradicional de clase (.xls, .xlsx o .csv)", type=["xls", "xlsx", "csv"], key="listado_gc71")
    estudiantes_df = pd.DataFrame(columns=["Nombre completo", "Documento", "Correo", "Plan", "Observación", "Estado"])
    if listado_file is not None:
        try:
            estudiantes_df = leer_listado_estudiantes(listado_file)
            st.success(f"Listado leído: {len(estudiantes_df)} estudiantes.")
            st.dataframe(estudiantes_df.head(20), use_container_width=True)
        except Exception as e:
            st.error(f"No se pudo leer el listado: {e}")
    representantes_df = pd.DataFrame(columns=["Nombre", "Documento"])
    if not estudiantes_df.empty:
        reps = estudiantes_df.head(3).copy()
        representantes_df = pd.DataFrame({"Nombre": reps["Nombre completo"], "Documento": reps["Documento"]})

    st.subheader("7. Fechas administrativas")
    c1, c2, c3 = st.columns(3)
    with c1:
        fecha_socializacion = st.date_input("Fecha de socialización", value=fecha_inicio, format="DD/MM/YYYY")
    with c2:
        fecha_revision = st.text_input("Fecha revisión coordinador", value="")
    with c3:
        fecha_aprobacion = st.text_input("Fecha aprobación comité / acta", value="")

    datos = {
        "programa": programa,
        "asignatura": asignatura,
        "codigo": codigo,
        "area": area,
        "prerrequisitos": prerrequisitos,
        "correquisitos": correquisitos,
        "tipo_asignatura": tipo,
        "creditos": str(creditos),
        "htp": str(htp),
        "hti": str(hti),
        "ht_total": str(htp + hti),
        "profesor": profesor,
        "cedula_docente": cedula_docente,
        "correo": correo,
        "grupo": grupo,
        "periodo": periodo,
        "justificacion": justificacion,
        "competencias": competencias,
        "resultados": resultados,
        "objetivo_general": objetivo_general,
        "objetivos_especificos": objetivos_especificos,
        "metodologias": metodologias,
        "ambientes": ambientes,
        "medios": medios,
        "referencias": referencias,
        "fecha_socializacion": fecha_socializacion.strftime("%d/%m/%Y"),
        "fecha_revision": fecha_revision,
        "fecha_aprobacion": fecha_aprobacion,
    }

    st.subheader("8. Descargar")
    col_doc, col_xlsx = st.columns(2)
    with col_doc:
        if not profesor.strip() or not asignatura.strip():
            st.info("Para descargar FD-GC71 ingrese como mínimo profesor y asignatura.")
        else:
            try:
                docx_bytes = crear_gc71_docx(datos, sesiones_df, evaluaciones_df, representantes_df)
                nombre_base = nombre_archivo_seguro(profesor, fecha_socializacion, "FD_GC71_Guia_Didactica")
                st.download_button(
                    "Descargar FD-GC71 en Word (.docx)",
                    data=docx_bytes,
                    file_name=f"{nombre_base}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            except Exception as e:
                st.error(f"No se pudo generar FD-GC71: {e}")
    with col_xlsx:
        try:
            plantilla_bytes = crear_plantilla_evaluacion_xlsx(estudiantes_df, evaluaciones_df, datos)
            nombre_xlsx = nombre_archivo_seguro(profesor or "docente", fecha_socializacion, "Plantilla_Evaluacion")
            st.download_button(
                "Descargar plantilla Excel de evaluación",
                data=plantilla_bytes,
                file_name=f"{nombre_xlsx}.xlsx",
                mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                use_container_width=True,
            )
        except Exception as e:
            st.error(f"No se pudo generar la plantilla de evaluación: {e}")

    st.session_state["ultimo_datos_gc71"] = datos
    st.session_state["ultimas_sesiones_gc71"] = sesiones_df
    st.session_state["ultimas_evaluaciones_gc71"] = evaluaciones_df
    st.session_state["ultimo_listado_estudiantes"] = estudiantes_df


def ui_gc72(st):
    st.header("FD-GC72 - Informe académico")
    st.caption("Puede diligenciarlo manualmente o alimentarlo con el listado de clase y las plantillas de evaluación cargadas a mitad/final del curso.")

    with st.sidebar:
        st.markdown("### Opciones FD-GC72")
        calcular = st.checkbox("Calcular porcentajes automáticamente", value=True)
        modo_analisis = st.radio("Tipo de análisis descriptivo", ["Por cada curso", "Consolidado institucional"], index=0)
        exportar_doc = st.checkbox("Intentar generar .doc además de .docx", value=False)

    st.subheader("1. Datos generales")
    c1, c2, c3 = st.columns([1.4, 1, 1])
    with c1:
        docente = st.text_input("Docente", value=st.session_state.get("ultimo_datos_gc71", {}).get("profesor", ""), key="gc72_docente")
    with c2:
        periodo = st.text_input("Período académico", value=st.session_state.get("ultimo_datos_gc71", {}).get("periodo", "2026-1"), key="gc72_periodo")
    with c3:
        fecha_entrega = st.date_input("Fecha de entrega", value=date.today(), format="DD/MM/YYYY", key="gc72_fecha")

    st.subheader("2. Alimentación automática desde listado y calificaciones")
    c1, c2 = st.columns(2)
    with c1:
        listado_file = st.file_uploader("Listado tradicional de clase", type=["xls", "xlsx", "csv"], key="listado_gc72")
    with c2:
        calificaciones_file = st.file_uploader("Plantilla/calificaciones de mitad o final del curso", type=["xlsx", "xls", "csv"], key="calificaciones_gc72")
    datos_gc71 = st.session_state.get("ultimo_datos_gc71", {})
    c1, c2, c3, c4 = st.columns(4)
    with c1:
        codigo_auto = st.text_input("Código asignatura", value=datos_gc71.get("codigo", ""), key="gc72_codigo_auto")
    with c2:
        grupo_auto = st.text_input("Grupo", value=datos_gc71.get("grupo", ""), key="gc72_grupo_auto")
    with c3:
        asignatura_auto = st.text_input("Asignatura", value=datos_gc71.get("asignatura", ""), key="gc72_asignatura_auto")
    with c4:
        corte_aprobacion = st.number_input("Nota mínima aprobatoria", min_value=0.0, max_value=5.0, value=3.0, step=0.1)
    sesiones_previas = st.session_state.get("ultimas_sesiones_gc71", pd.DataFrame())
    avance_sugerido = 0.0
    if isinstance(sesiones_previas, pd.DataFrame) and not sesiones_previas.empty:
        sesiones_realizadas = st.number_input("Sesiones realizadas para calcular avance", min_value=0, max_value=len(sesiones_previas), value=min(len(sesiones_previas), 0), step=1)
        avance_sugerido = round((sesiones_realizadas / len(sesiones_previas)) * 100, 1)
    else:
        avance_sugerido = st.number_input("% Avance en contenido", min_value=0.0, max_value=100.0, value=0.0, step=1.0)
    porcentaje_evaluado_manual = st.number_input("% Evaluado sugerido/manual", min_value=0.0, max_value=100.0, value=0.0, step=1.0)

    df_auto = pd.DataFrame(columns=COLUMNAS_GC72)
    if listado_file is not None or calificaciones_file is not None:
        try:
            listado_df = leer_listado_estudiantes(listado_file) if listado_file is not None else st.session_state.get("ultimo_listado_estudiantes", pd.DataFrame())
            cal_df, ev_df = leer_calificaciones(calificaciones_file) if calificaciones_file is not None else (pd.DataFrame(), pd.DataFrame())
            df_auto = resumen_gc72_desde_archivos(listado_df, cal_df, ev_df, codigo_auto, grupo_auto, asignatura_auto, avance_sugerido, porcentaje_evaluado_manual, corte_aprobacion)
            st.success("Resumen generado desde los archivos cargados. Revísalo antes de descargar: Excel no reemplaza criterio docente, pero sí evita trabajo de mula.")
            st.dataframe(df_auto, use_container_width=True)
        except Exception as e:
            st.error(f"No se pudo generar el resumen automático: {e}")

    st.subheader("3. Cursos")
    archivo = st.file_uploader("Opcional: cargar cursos desde Excel o CSV con columnas FD-GC72", type=["xlsx", "xls", "csv"], key="cursos_gc72")
    base = df_auto if not df_auto.empty else df_vacio(COLUMNAS_GC72, filas=3)
    if archivo is not None:
        try:
            if archivo.name.lower().endswith(".csv"):
                base = pd.read_csv(archivo)
            else:
                base = pd.read_excel(archivo)
        except Exception as e:
            st.error(f"No se pudo leer archivo de cursos: {e}")
    for col in COLUMNAS_GC72:
        if col not in base.columns:
            base[col] = ""
    base = base[COLUMNAS_GC72]

    config_columnas = {
        "Código": st.column_config.TextColumn("Código", width="small"),
        "Grupo": st.column_config.TextColumn("Grupo", width="small"),
        "Asignatura": st.column_config.TextColumn("Asignatura", width="medium"),
        "% Avance en contenido": st.column_config.NumberColumn("% Avance", min_value=0, max_value=100, step=1),
        "% Evaluado": st.column_config.NumberColumn("% Evaluado", min_value=0, max_value=100, step=1),
        "Estudiantes matriculados": st.column_config.NumberColumn("Matriculados", min_value=0, step=1),
        "Desertaron N°": st.column_config.NumberColumn("Desertaron N°", min_value=0, step=1),
        "Desertaron %": st.column_config.TextColumn("Desertaron %", disabled=calcular),
        "Aprueban evaluación parcial N°": st.column_config.NumberColumn("Aprueban parcial N°", min_value=0, step=1),
        "Aprueban evaluación parcial %": st.column_config.TextColumn("Aprueban parcial %", disabled=calcular),
        "Reprueban evaluación parcial N°": st.column_config.NumberColumn("Reprueban parcial N°", min_value=0, step=1),
        "Reprueban evaluación parcial %": st.column_config.TextColumn("Reprueban parcial %", disabled=calcular),
        "Aprueban a la fecha N°": st.column_config.NumberColumn("Aprueban fecha N°", min_value=0, step=1),
        "Aprueban a la fecha %": st.column_config.TextColumn("Aprueban fecha %", disabled=calcular),
        "Reprueban a la fecha N°": st.column_config.NumberColumn("Reprueban fecha N°", min_value=0, step=1),
        "Reprueban a la fecha %": st.column_config.TextColumn("Reprueban fecha %", disabled=calcular),
    }
    cursos_editados = st.data_editor(base, column_config=config_columnas, hide_index=True, num_rows="dynamic", use_container_width=True, key="tabla_gc72")
    cursos = normalizar_dataframe_gc72(cursos_editados, calcular_porcentajes=calcular)
    if cursos.empty:
        st.info("Agregue al menos un curso para activar el análisis y la descarga.")
        return

    st.subheader("4. Análisis descriptivo con preformas")
    analisis_por_curso: Dict[str, Dict[str, str]] = {}
    cursos_iteracion = cursos if modo_analisis == "Por cada curso" else cursos.head(1).copy()
    if modo_analisis == "Consolidado institucional":
        cursos_iteracion.loc[:, "Asignatura"] = "Consolidado institucional"
        cursos_iteracion.loc[:, "Código"] = ""
        cursos_iteracion.loc[:, "Grupo"] = ""
    for idx, row in cursos_iteracion.iterrows():
        key = curso_key(row, idx)
        titulo = str(row.get("Asignatura", "")).strip() or f"Curso {idx + 1}"
        subtitulo = f"{titulo} | Código: {row.get('Código', '')} | Grupo: {row.get('Grupo', '')}" if modo_analisis == "Por cada curso" else "Análisis consolidado para todos los cursos"
        with st.expander(subtitulo, expanded=(idx == 0)):
            col_a, col_b, col_c = st.columns(3)
            with col_a:
                sel_pos = st.multiselect("Aspectos positivos", list(PREFORMAS_GC72["aspectos_positivos"].keys()), default=["Avance adecuado", "Aplicación práctica"], key=f"pos_{key}")
                add_pos = st.text_area("Texto adicional positivo", key=f"pos_add_{key}", height=90)
            with col_b:
                sel_inc = st.multiselect("Inconvenientes", list(PREFORMAS_GC72["inconvenientes"].keys()), default=[], key=f"inc_{key}")
                add_inc = st.text_area("Texto adicional de inconvenientes", key=f"inc_add_{key}", height=90)
            with col_c:
                sel_prop = st.multiselect("Propuestas metodológicas", list(PREFORMAS_GC72["propuestas"].keys()), default=["Seguimiento formativo", "Talleres aplicados"], key=f"prop_{key}")
                add_prop = st.text_area("Texto adicional de propuestas", key=f"prop_add_{key}", height=90)
            analisis_por_curso[key] = {
                "positivos": texto_preformas(sel_pos, "aspectos_positivos", add_pos),
                "inconvenientes": texto_preformas(sel_inc, "inconvenientes", add_inc),
                "propuestas": texto_preformas(sel_prop, "propuestas", add_prop),
            }
    bloques = construir_analisis(cursos if modo_analisis == "Por cada curso" else cursos_iteracion, analisis_por_curso, modo_analisis)

    st.subheader("5. Descargar")
    if not docente.strip():
        st.warning("Ingrese el nombre del docente antes de descargar.")
        return
    try:
        docx_bytes = crear_informe_gc72_docx(docente, periodo, fecha_entrega, cursos, bloques)
        nombre_base = nombre_archivo_seguro(docente, fecha_entrega, "FD_GC72_Informe_Academico")
        st.download_button("Descargar FD-GC72 en Word (.docx)", data=docx_bytes, file_name=f"{nombre_base}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
        if exportar_doc:
            doc_bytes = convertir_doc_si_es_posible(docx_bytes, nombre_base)
            if doc_bytes:
                st.download_button("Descargar FD-GC72 en Word 97-2003 (.doc)", data=doc_bytes, file_name=f"{nombre_base}.doc", mime="application/msword", use_container_width=True)
            else:
                st.info("No se pudo generar .doc. Instale LibreOffice o use .docx.")
    except Exception as e:
        st.error(f"No se pudo generar FD-GC72: {e}")


def ui_ayuda(st):
    st.header("Flujo recomendado")
    st.markdown(
        """
1. **Inicio del curso:** diligencie FD-GC71, defina módulos, intensidad, horario y evaluación concertada. Descargue la guía y la plantilla de evaluación.
2. **Durante el curso:** cargue notas en la plantilla Excel. Mantenga actualizada la columna Estado si hay retiros o deserciones.
3. **Mitad/final del curso:** suba el listado tradicional y la plantilla de evaluación al módulo FD-GC72. La app calcula matriculados, desertores, aprobados, reprobados y porcentajes.
4. **Cierre:** revise la tabla, ajuste el análisis descriptivo con preformas y descargue el informe académico.

Columnas esperadas del listado tradicional: **NOMBRE COMPLETO, DOCUMENTO, PLAN, OBSERVACIÓN, CORREO**. Si el archivo trae más columnas, no pasa nada; la app toma lo que necesita.
        """
    )


# -----------------------------------------------------------------------------
# Seguridad, login, perfiles y auditoría local
# -----------------------------------------------------------------------------
ROLES_PERMISOS = {
    "Administrador": {
        "descripcion": "Control total: usuarios, perfiles, auditoría y generación de formatos.",
        "modulos": ["Inicio", "FD-GC71 - Planeación", "FD-GC72 - Informe académico", "Usuarios y perfiles", "Auditoría", "Mi cuenta", "Ayuda / flujo recomendado"],
    },
    "Coordinador": {
        "descripcion": "Revisión académica: puede usar formatos, ver auditoría y acompañar cierres.",
        "modulos": ["Inicio", "FD-GC71 - Planeación", "FD-GC72 - Informe académico", "Auditoría", "Mi cuenta", "Ayuda / flujo recomendado"],
    },
    "Docente": {
        "descripcion": "Operación docente: planeación, evaluación, informes y descarga de soportes.",
        "modulos": ["Inicio", "FD-GC71 - Planeación", "FD-GC72 - Informe académico", "Mi cuenta", "Ayuda / flujo recomendado"],
    },
    "Consulta": {
        "descripcion": "Acceso de lectura al flujo y orientación operativa.",
        "modulos": ["Inicio", "Mi cuenta", "Ayuda / flujo recomendado"],
    },
}


def ahora_iso() -> str:
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


# La conexión productiva fue definida al inicio del archivo.
# Se conserva este punto solo por compatibilidad con la versión anterior.
# No redefinir conexion_db aquí.


def hash_password(password: str, salt: Optional[str] = None) -> Tuple[str, str]:
    salt = salt or secrets.token_hex(16)
    digest = hashlib.pbkdf2_hmac("sha256", password.encode("utf-8"), salt.encode("utf-8"), 240_000)
    return salt, digest.hex()


def verificar_password(password: str, salt: str, password_hash: str) -> bool:
    _, digest = hash_password(password, salt)
    return hmac.compare_digest(digest, password_hash)


def validar_password_seguro(password: str):
    """Política de contraseña. En producción exige mayor rigor; en local conserva usabilidad."""
    min_len = 12 if get_app_env() == "production" else 8
    if len(password or "") < min_len:
        raise ValueError(f"La contraseña debe tener mínimo {min_len} caracteres.")
    if get_app_env() == "production":
        reglas = [
            (re.search(r"[A-ZÁÉÍÓÚÑ]", password or ""), "una mayúscula"),
            (re.search(r"[a-záéíóúñ]", password or ""), "una minúscula"),
            (re.search(r"\d", password or ""), "un número"),
            (re.search(r"[^A-Za-z0-9ÁÉÍÓÚáéíóúÑñ]", password or ""), "un carácter especial"),
        ]
        faltantes = [nombre for ok, nombre in reglas if not ok]
        if faltantes:
            raise ValueError("La contraseña debe incluir " + ", ".join(faltantes) + ".")



def init_db():
    conn = conexion_db()
    cur = conn.cursor()
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS usuarios (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario TEXT NOT NULL UNIQUE,
            nombre_completo TEXT NOT NULL,
            email TEXT,
            rol TEXT NOT NULL,
            salt TEXT NOT NULL,
            password_hash TEXT NOT NULL,
            activo INTEGER NOT NULL DEFAULT 1,
            debe_cambiar_clave INTEGER NOT NULL DEFAULT 0,
            creado_en TEXT NOT NULL,
            actualizado_en TEXT NOT NULL,
            ultimo_login TEXT
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS auditoria (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            fecha TEXT NOT NULL,
            usuario TEXT,
            rol TEXT,
            accion TEXT NOT NULL,
            detalle TEXT
        )
        """
    )
    cur.execute(
        """
        CREATE TABLE IF NOT EXISTS proyectos_guardados (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            usuario_id INTEGER NOT NULL,
            tipo TEXT NOT NULL,
            titulo TEXT NOT NULL,
            payload_json TEXT NOT NULL,
            creado_en TEXT NOT NULL,
            actualizado_en TEXT NOT NULL,
            FOREIGN KEY(usuario_id) REFERENCES usuarios(id)
        )
        """
    )
    conn.commit()

    total = cur.execute("SELECT COUNT(*) AS n FROM usuarios").fetchone()["n"]
    if total == 0:
        salt, digest = hash_password("Admin123*")
        cur.execute(
            """
            INSERT INTO usuarios(usuario, nombre_completo, email, rol, salt, password_hash, activo, debe_cambiar_clave, creado_en, actualizado_en)
            VALUES (?, ?, ?, ?, ?, ?, 1, 1, ?, ?)
            """,
            ("admin", "Administrador del sistema", "", "Administrador", salt, digest, ahora_iso(), ahora_iso()),
        )
        conn.commit()
    conn.close()


def registrar_auditoria(accion: str, detalle: str = ""):
    try:
        usuario = st.session_state.get("auth_user", {}) if "st" in globals() else {}
    except Exception:
        usuario = {}
    conn = conexion_db()
    conn.execute(
        "INSERT INTO auditoria(fecha, usuario, rol, accion, detalle) VALUES (?, ?, ?, ?, ?)",
        (ahora_iso(), usuario.get("usuario", "sistema"), usuario.get("rol", ""), accion, detalle[:1500] if detalle else ""),
    )
    conn.commit()
    conn.close()


def autenticar_usuario(usuario: str, password: str) -> Optional[Dict[str, str]]:
    """Autenticación con bloqueo temporal por intentos fallidos."""
    usuario_clean = usuario.strip()
    conn = conexion_db()
    row = conn.execute("SELECT * FROM usuarios WHERE lower(usuario)=lower(?)", (usuario_clean,)).fetchone()
    if not row:
        conn.close()
        return None
    if not row["activo"]:
        conn.close()
        return {"error": "Usuario inactivo. Solicite habilitación al administrador."}

    bloqueado_hasta = None
    try:
        bloqueado_hasta = row["bloqueado_hasta"]
    except Exception:
        bloqueado_hasta = None
    if bloqueado_hasta:
        try:
            until = datetime.fromisoformat(str(bloqueado_hasta))
            if datetime.now() < until:
                conn.close()
                return {"error": f"Usuario bloqueado temporalmente hasta {until.strftime('%H:%M:%S')}."}
        except Exception:
            pass

    if not verificar_password(password, row["salt"], row["password_hash"]):
        try:
            intentos = int(row["intentos_fallidos"] or 0) + 1
        except Exception:
            intentos = 1
        bloqueado = None
        if intentos >= 5:
            bloqueado = (datetime.now() + timedelta(minutes=15)).isoformat(timespec="seconds")
            intentos = 0
        try:
            conn.execute("UPDATE usuarios SET intentos_fallidos=?, bloqueado_hasta=?, actualizado_en=? WHERE id=?", (intentos, bloqueado, ahora_iso(), row["id"]))
            conn.commit()
        except Exception:
            pass
        conn.close()
        return None

    conn.execute("UPDATE usuarios SET ultimo_login=?, intentos_fallidos=0, bloqueado_hasta=NULL WHERE id=?", (ahora_iso(), row["id"]))
    conn.commit()
    conn.close()
    return {
        "id": row["id"],
        "usuario": row["usuario"],
        "nombre_completo": row["nombre_completo"],
        "email": row["email"] or "",
        "rol": row["rol"],
        "debe_cambiar_clave": bool(row["debe_cambiar_clave"]),
    }


def crear_usuario(usuario: str, nombre: str, email: str, rol: str, password: str, activo: bool = True, debe_cambiar: bool = True):
    if rol not in ROLES_PERMISOS:
        raise ValueError("Rol no válido.")
    validar_password_seguro(password)
    salt, digest = hash_password(password)
    conn = conexion_db()
    conn.execute(
        """
        INSERT INTO usuarios(usuario, nombre_completo, email, rol, salt, password_hash, activo, debe_cambiar_clave, creado_en, actualizado_en)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (usuario.strip(), nombre.strip(), email.strip(), rol, salt, digest, int(activo), int(debe_cambiar), ahora_iso(), ahora_iso()),
    )
    conn.commit()
    conn.close()
    registrar_auditoria("Crear usuario", f"Usuario={usuario.strip()} | Rol={rol}")


def actualizar_usuario(user_id: int, nombre: str, email: str, rol: str, activo: bool):
    if rol not in ROLES_PERMISOS:
        raise ValueError("Rol no válido.")
    conn = conexion_db()
    conn.execute(
        "UPDATE usuarios SET nombre_completo=?, email=?, rol=?, activo=?, actualizado_en=? WHERE id=?",
        (nombre.strip(), email.strip(), rol, int(activo), ahora_iso(), user_id),
    )
    conn.commit()
    conn.close()
    registrar_auditoria("Actualizar usuario", f"ID={user_id} | Rol={rol} | Activo={activo}")


def cambiar_password(user_id: int, password: str, forzar_cambio: bool = False):
    validar_password_seguro(password)
    salt, digest = hash_password(password)
    conn = conexion_db()
    conn.execute(
        "UPDATE usuarios SET salt=?, password_hash=?, debe_cambiar_clave=?, actualizado_en=? WHERE id=?",
        (salt, digest, int(forzar_cambio), ahora_iso(), user_id),
    )
    conn.commit()
    conn.close()
    registrar_auditoria("Cambiar contraseña", f"ID={user_id}")


def listar_usuarios() -> pd.DataFrame:
    df = read_sql_df(
        "SELECT id, usuario, nombre_completo, email, rol, activo, debe_cambiar_clave, creado_en, actualizado_en, ultimo_login FROM usuarios ORDER BY rol, usuario"
    )
    if not df.empty:
        df["activo"] = df["activo"].map({1: "Sí", 0: "No"})
        df["debe_cambiar_clave"] = df["debe_cambiar_clave"].map({1: "Sí", 0: "No"})
    return df


def tiene_permiso(modulo: str) -> bool:
    user = st.session_state.get("auth_user")
    if not user:
        return False
    return modulo in ROLES_PERMISOS.get(user.get("rol"), {}).get("modulos", [])


def pantalla_login(st):
    st.set_page_config(page_title="Login | Gestor FD-GC71 / FD-GC72", layout="centered")
    st.title("Ingreso al gestor académico")
    st.caption("FD-GC71 / FD-GC72 con control de acceso por perfiles.")
    with st.form("login_form", clear_on_submit=False):
        usuario = st.text_input("Usuario", value="admin")
        password = st.text_input("Contraseña", type="password")
        entrar = st.form_submit_button("Entrar", use_container_width=True)
    if entrar:
        user = autenticar_usuario(usuario, password)
        if user and not user.get("error"):
            st.session_state["auth_user"] = user
            registrar_auditoria("Login", "Ingreso correcto")
            st.rerun()
        elif user and user.get("error"):
            st.error(user["error"])
        else:
            st.error("Usuario o contraseña incorrectos.")
    with st.expander("Primera instalación"):
        if get_app_env() == "production":
            st.info("Las credenciales iniciales se leen desde los secretos INITIAL_ADMIN_USER e INITIAL_ADMIN_PASSWORD.")
        else:
            admin_user, _, _, _ = initial_admin_config()
            st.write(f"Usuario inicial: `{admin_user}`")
            st.write("Contraseña inicial: definida en secrets o, en local, `Admin123*`.")
        st.warning("Cambie la contraseña apenas ingrese y cree usuarios nominales. La puerta abierta también sirve para que entre el incendio.")


def ui_inicio(st):
    user = st.session_state.get("auth_user", {})
    st.header("Panel de inicio")
    c1, c2, c3 = st.columns(3)
    c1.metric("Usuario", user.get("usuario", ""))
    c2.metric("Perfil", user.get("rol", ""))
    c3.metric("Estado", "Activo")
    st.info(ROLES_PERMISOS.get(user.get("rol"), {}).get("descripcion", ""))
    if get_app_env() == "production" and not usar_postgres():
        st.error("APP_ENV está en production pero no hay DATABASE_URL. Configure PostgreSQL antes de operar con datos reales.")
    elif usar_postgres():
        st.success("Persistencia productiva activa: PostgreSQL externo configurado.")
    st.subheader("Permisos habilitados")
    permisos = ROLES_PERMISOS.get(user.get("rol"), {}).get("modulos", [])
    st.write(", ".join([p for p in permisos if p not in ["Inicio", "Mi cuenta"]]))
    if user.get("debe_cambiar_clave"):
        st.warning("Debe cambiar la contraseña inicial desde el módulo Mi cuenta antes de operar formalmente.")

    st.subheader("Ruta operativa recomendada")
    st.markdown(
        """
- **Inicio del curso:** planee la asignatura en FD-GC71, defina unidades, intensidad, horario y evaluación concertada.
- **Durante el curso:** use la plantilla de evaluación descargada para registrar notas y estados.
- **Mitad y final:** cargue listado y calificaciones en FD-GC72 para calcular avance, evaluación, aprobados, reprobados y desertores.
- **Cierre:** descargue el informe académico con análisis descriptivo por curso.
        """
    )


def ui_mi_cuenta(st):
    user = st.session_state.get("auth_user", {})
    st.header("Mi cuenta")
    st.write(f"**Nombre:** {user.get('nombre_completo', '')}")
    st.write(f"**Usuario:** {user.get('usuario', '')}")
    st.write(f"**Correo:** {user.get('email', '')}")
    st.write(f"**Perfil:** {user.get('rol', '')}")
    with st.form("form_cambiar_clave"):
        nueva = st.text_input("Nueva contraseña", type="password")
        confirmar = st.text_input("Confirmar contraseña", type="password")
        enviar = st.form_submit_button("Actualizar contraseña", use_container_width=True)
    if enviar:
        if nueva != confirmar:
            st.error("Las contraseñas no coinciden.")
        else:
            try:
                cambiar_password(int(user["id"]), nueva, forzar_cambio=False)
                st.session_state["auth_user"]["debe_cambiar_clave"] = False
                st.success("Contraseña actualizada.")
            except Exception as e:
                st.error(str(e))


def ui_admin_usuarios(st):
    st.header("Usuarios y perfiles")
    st.caption("Administración local de usuarios. La base queda en app_data/fdgc_app.sqlite3.")
    usuarios = listar_usuarios()
    st.dataframe(usuarios, use_container_width=True, hide_index=True)

    tab_crear, tab_editar, tab_clave = st.tabs(["Crear usuario", "Editar perfil", "Resetear contraseña"])
    with tab_crear:
        with st.form("crear_usuario"):
            c1, c2 = st.columns(2)
            with c1:
                usuario = st.text_input("Usuario nuevo")
                nombre = st.text_input("Nombre completo")
                email = st.text_input("Correo")
            with c2:
                rol = st.selectbox("Perfil", list(ROLES_PERMISOS.keys()), index=2)
                password = st.text_input("Contraseña inicial", type="password", value="Cambio123*")
                activo = st.checkbox("Activo", value=True)
                debe_cambiar = st.checkbox("Forzar cambio de contraseña", value=True)
            enviar = st.form_submit_button("Crear usuario", use_container_width=True)
        if enviar:
            try:
                crear_usuario(usuario, nombre, email, rol, password, activo, debe_cambiar)
                st.success("Usuario creado.")
                st.rerun()
            except Exception as e:
                st.error(f"No se pudo crear: {e}")

    with tab_editar:
        if usuarios.empty:
            st.info("No hay usuarios.")
        else:
            ids = usuarios["id"].tolist()
            seleccionado = st.selectbox("Usuario a editar", ids, format_func=lambda x: usuarios.loc[usuarios["id"] == x, "usuario"].iloc[0], key="editar_usuario_id")
            fila = usuarios.loc[usuarios["id"] == seleccionado].iloc[0]
            with st.form("editar_usuario"):
                nombre = st.text_input("Nombre completo", value=str(fila["nombre_completo"]))
                email = st.text_input("Correo", value=str(fila["email"]))
                rol = st.selectbox("Perfil", list(ROLES_PERMISOS.keys()), index=list(ROLES_PERMISOS.keys()).index(str(fila["rol"])))
                activo = st.checkbox("Activo", value=(str(fila["activo"]) == "Sí"))
                enviar = st.form_submit_button("Guardar cambios", use_container_width=True)
            if enviar:
                try:
                    actualizar_usuario(int(seleccionado), nombre, email, rol, activo)
                    st.success("Usuario actualizado.")
                    st.rerun()
                except Exception as e:
                    st.error(f"No se pudo actualizar: {e}")

    with tab_clave:
        if usuarios.empty:
            st.info("No hay usuarios.")
        else:
            seleccionado = st.selectbox("Usuario", usuarios["id"].tolist(), format_func=lambda x: usuarios.loc[usuarios["id"] == x, "usuario"].iloc[0], key="reset_usuario_id")
            with st.form("reset_clave"):
                nueva = st.text_input("Nueva contraseña", type="password", value="Cambio123*")
                forzar = st.checkbox("Forzar cambio al iniciar", value=True)
                enviar = st.form_submit_button("Resetear contraseña", use_container_width=True)
            if enviar:
                try:
                    cambiar_password(int(seleccionado), nueva, forzar_cambio=forzar)
                    st.success("Contraseña reseteada.")
                except Exception as e:
                    st.error(str(e))

    st.subheader("Matriz de perfiles")
    matriz = []
    for rol, cfg in ROLES_PERMISOS.items():
        matriz.append({"Perfil": rol, "Descripción": cfg["descripcion"], "Módulos": ", ".join(cfg["modulos"])})
    st.dataframe(pd.DataFrame(matriz), use_container_width=True, hide_index=True)


def ui_auditoria(st):
    st.header("Auditoría")
    st.caption("Registro básico de ingresos y cambios de seguridad.")
    df = read_sql_df("SELECT fecha, usuario, rol, accion, detalle FROM auditoria ORDER BY id DESC LIMIT 500")
    st.dataframe(df, use_container_width=True, hide_index=True)
    csv = df.to_csv(index=False).encode("utf-8-sig")
    st.download_button("Descargar auditoría CSV", data=csv, file_name="auditoria_fdgc.csv", mime="text/csv", use_container_width=True)


def main():
    import streamlit as st
    globals()["st"] = st
    init_db()

    if "auth_user" not in st.session_state:
        pantalla_login(st)
        return

    user = st.session_state.get("auth_user", {})
    st.set_page_config(page_title="Gestor FD-GC71 / FD-GC72", layout="wide")
    st.title("Gestor académico FD-GC71 / FD-GC72")
    st.caption("Planeación de clases, concertación de evaluación, plantilla de notas, informe académico y control de acceso por perfiles.")

    with st.sidebar:
        st.markdown(f"**{user.get('nombre_completo', '')}**")
        st.caption(f"Perfil: {user.get('rol', '')}")
        if st.button("Cerrar sesión", use_container_width=True):
            registrar_auditoria("Logout", "Cierre de sesión")
            st.session_state.pop("auth_user", None)
            st.rerun()
        st.divider()
        modulos = ROLES_PERMISOS.get(user.get("rol"), {}).get("modulos", ["Inicio", "Mi cuenta", "Ayuda / flujo recomendado"])
        pagina = st.radio("Módulo", modulos)

    if not tiene_permiso(pagina):
        st.error("Este perfil no tiene permisos para abrir este módulo.")
        return

    if pagina == "Inicio":
        ui_inicio(st)
    elif pagina.startswith("FD-GC71"):
        ui_gc71(st)
    elif pagina.startswith("FD-GC72"):
        ui_gc72(st)
    elif pagina == "Usuarios y perfiles":
        ui_admin_usuarios(st)
    elif pagina == "Auditoría":
        ui_auditoria(st)
    elif pagina == "Mi cuenta":
        ui_mi_cuenta(st)
    else:
        ui_ayuda(st)


# =============================================================================
# VERSIÓN ENTERPRISE / ESTADIO: expediente, persistencia, control, evidencias
# =============================================================================
# Esta sección se define al final para sobreescribir/elevar algunas funciones
# anteriores sin romper la compatibilidad del prototipo inicial.

import zipfile
import uuid
from html import escape

EVIDENCE_DIR = DATA_DIR / "evidencias"
EXPORT_DIR = DATA_DIR / "exportaciones"
BACKUP_DIR = DATA_DIR / "backups"

MODULO_CENTRO = "Centro de control"
MODULO_EXPEDIENTE = "Expediente académico"
MODULO_PLANEADOR = "Planeador superior"
MODULO_EVIDENCIAS = "Evidencias y soportes"
MODULO_VALIDADOR = "Validador institucional"
MODULO_BACKUP = "Copias y restauración"
MODULO_DIAGNOSTICO = "Diagnóstico productivo"

ROLES_PERMISOS = {
    "Administrador": {
        "descripcion": "Gobierno total del sistema: usuarios, cursos, formatos, evidencias, auditoría, respaldos y parametrización.",
        "modulos": ["Inicio", MODULO_CENTRO, MODULO_EXPEDIENTE, MODULO_PLANEADOR, "FD-GC71 - Planeación", "FD-GC72 - Informe académico", MODULO_EVIDENCIAS, MODULO_VALIDADOR, MODULO_BACKUP, MODULO_DIAGNOSTICO, "Usuarios y perfiles", "Auditoría", "Mi cuenta", "Ayuda / flujo recomendado"],
    },
    "Coordinador": {
        "descripcion": "Revisión y cierre académico: puede controlar cursos, validar planeación, revisar evidencias, generar informes y consultar auditoría.",
        "modulos": ["Inicio", MODULO_CENTRO, MODULO_EXPEDIENTE, MODULO_PLANEADOR, "FD-GC71 - Planeación", "FD-GC72 - Informe académico", MODULO_EVIDENCIAS, MODULO_VALIDADOR, MODULO_DIAGNOSTICO, "Auditoría", "Mi cuenta", "Ayuda / flujo recomendado"],
    },
    "Docente": {
        "descripcion": "Operación docente: planeación, evaluación, evidencias, informes y descarga de soportes de sus cursos.",
        "modulos": ["Inicio", MODULO_CENTRO, MODULO_EXPEDIENTE, MODULO_PLANEADOR, "FD-GC71 - Planeación", "FD-GC72 - Informe académico", MODULO_EVIDENCIAS, MODULO_VALIDADOR, "Mi cuenta", "Ayuda / flujo recomendado"],
    },
    "Consulta": {
        "descripcion": "Consulta controlada: lectura de ruta operativa, cuenta propia y ayuda.",
        "modulos": ["Inicio", MODULO_CENTRO, "Mi cuenta", "Ayuda / flujo recomendado"],
    },
}


def safe_json_loads(value, default=None):
    if default is None:
        default = {}
    if value in (None, "", float("nan")):
        return default
    try:
        return json.loads(value)
    except Exception:
        return default


def df_to_payload(df: pd.DataFrame) -> List[Dict[str, str]]:
    if df is None or df.empty:
        return []
    cleaned = df.copy().fillna("")
    return cleaned.astype(str).to_dict(orient="records")


def payload_to_df(payload, columnas: Optional[List[str]] = None) -> pd.DataFrame:
    if not payload:
        return pd.DataFrame(columns=columnas or [])
    df = pd.DataFrame(payload)
    if columnas:
        for c in columnas:
            if c not in df.columns:
                df[c] = ""
        df = df[columnas]
    return df


def init_db():
    """Inicializa la base en modo local o productivo.

    - SQLite: desarrollo / demo / uso local.
    - PostgreSQL: Streamlit Cloud o servidor institucional con persistencia real.
    """
    DATA_DIR.mkdir(exist_ok=True)
    EVIDENCE_DIR.mkdir(exist_ok=True)
    EXPORT_DIR.mkdir(exist_ok=True)
    BACKUP_DIR.mkdir(exist_ok=True)
    conn = conexion_db()

    if usar_postgres():
        ddl = {
            "usuarios": """
                CREATE TABLE IF NOT EXISTS usuarios (
                    id SERIAL PRIMARY KEY,
                    usuario TEXT UNIQUE NOT NULL,
                    nombre_completo TEXT NOT NULL,
                    email TEXT,
                    rol TEXT NOT NULL,
                    salt TEXT NOT NULL,
                    password_hash TEXT NOT NULL,
                    activo INTEGER NOT NULL DEFAULT 1,
                    debe_cambiar_clave INTEGER NOT NULL DEFAULT 1,
                    creado_en TEXT NOT NULL,
                    actualizado_en TEXT,
                    ultimo_login TEXT
                )
            """,
            "auditoria": """
                CREATE TABLE IF NOT EXISTS auditoria (
                    id SERIAL PRIMARY KEY,
                    fecha TEXT NOT NULL,
                    usuario TEXT,
                    rol TEXT,
                    accion TEXT NOT NULL,
                    detalle TEXT
                )
            """,
            "cursos": """
                CREATE TABLE IF NOT EXISTS cursos (
                    id SERIAL PRIMARY KEY,
                    codigo TEXT,
                    grupo TEXT,
                    asignatura TEXT NOT NULL,
                    programa TEXT,
                    periodo TEXT,
                    profesor TEXT,
                    email_profesor TEXT,
                    creditos TEXT,
                    htp DOUBLE PRECISION DEFAULT 0,
                    hti DOUBLE PRECISION DEFAULT 0,
                    fecha_inicio TEXT,
                    fecha_fin TEXT,
                    estado TEXT DEFAULT 'Planeación',
                    avance_contenido DOUBLE PRECISION DEFAULT 0,
                    avance_evaluado DOUBLE PRECISION DEFAULT 0,
                    propietario_usuario TEXT,
                    creado_por TEXT,
                    creado_en TEXT NOT NULL,
                    actualizado_en TEXT,
                    payload_json TEXT DEFAULT '{}'
                )
            """,
            "evidencias": """
                CREATE TABLE IF NOT EXISTS evidencias (
                    id SERIAL PRIMARY KEY,
                    curso_id INTEGER REFERENCES cursos(id) ON DELETE SET NULL,
                    tipo TEXT,
                    nombre_original TEXT NOT NULL,
                    nombre_archivo TEXT NOT NULL,
                    mime TEXT,
                    tamano INTEGER DEFAULT 0,
                    descripcion TEXT,
                    subido_por TEXT,
                    subido_en TEXT NOT NULL,
                    contenido_b64 TEXT
                )
            """,
            "artefactos": """
                CREATE TABLE IF NOT EXISTS artefactos (
                    id SERIAL PRIMARY KEY,
                    curso_id INTEGER REFERENCES cursos(id) ON DELETE SET NULL,
                    tipo TEXT NOT NULL,
                    nombre_archivo TEXT NOT NULL,
                    descripcion TEXT,
                    generado_por TEXT,
                    generado_en TEXT NOT NULL
                )
            """,
        }
    else:
        ddl = {
            "usuarios": """
                CREATE TABLE IF NOT EXISTS usuarios (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    usuario TEXT UNIQUE NOT NULL,
                    nombre_completo TEXT NOT NULL,
                    email TEXT,
                    rol TEXT NOT NULL,
                    salt TEXT NOT NULL,
                    password_hash TEXT NOT NULL,
                    activo INTEGER NOT NULL DEFAULT 1,
                    debe_cambiar_clave INTEGER NOT NULL DEFAULT 1,
                    creado_en TEXT NOT NULL,
                    actualizado_en TEXT,
                    ultimo_login TEXT
                )
            """,
            "auditoria": """
                CREATE TABLE IF NOT EXISTS auditoria (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    fecha TEXT NOT NULL,
                    usuario TEXT,
                    rol TEXT,
                    accion TEXT NOT NULL,
                    detalle TEXT
                )
            """,
            "cursos": """
                CREATE TABLE IF NOT EXISTS cursos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    codigo TEXT,
                    grupo TEXT,
                    asignatura TEXT NOT NULL,
                    programa TEXT,
                    periodo TEXT,
                    profesor TEXT,
                    email_profesor TEXT,
                    creditos TEXT,
                    htp REAL DEFAULT 0,
                    hti REAL DEFAULT 0,
                    fecha_inicio TEXT,
                    fecha_fin TEXT,
                    estado TEXT DEFAULT 'Planeación',
                    avance_contenido REAL DEFAULT 0,
                    avance_evaluado REAL DEFAULT 0,
                    propietario_usuario TEXT,
                    creado_por TEXT,
                    creado_en TEXT NOT NULL,
                    actualizado_en TEXT,
                    payload_json TEXT DEFAULT '{}'
                )
            """,
            "evidencias": """
                CREATE TABLE IF NOT EXISTS evidencias (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    curso_id INTEGER,
                    tipo TEXT,
                    nombre_original TEXT NOT NULL,
                    nombre_archivo TEXT NOT NULL,
                    mime TEXT,
                    tamano INTEGER DEFAULT 0,
                    descripcion TEXT,
                    subido_por TEXT,
                    subido_en TEXT NOT NULL,
                    contenido_b64 TEXT,
                    FOREIGN KEY(curso_id) REFERENCES cursos(id)
                )
            """,
            "artefactos": """
                CREATE TABLE IF NOT EXISTS artefactos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    curso_id INTEGER,
                    tipo TEXT NOT NULL,
                    nombre_archivo TEXT NOT NULL,
                    descripcion TEXT,
                    generado_por TEXT,
                    generado_en TEXT NOT NULL,
                    FOREIGN KEY(curso_id) REFERENCES cursos(id)
                )
            """,
        }

    for sql in ddl.values():
        conn.execute(sql)
    conn.execute("CREATE INDEX IF NOT EXISTS idx_cursos_estado ON cursos(estado)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_cursos_owner ON cursos(propietario_usuario)")
    conn.execute("CREATE INDEX IF NOT EXISTS idx_evidencias_curso ON evidencias(curso_id)")
    conn.commit()
    conn.close()

    # Migraciones defensivas para instalaciones existentes.
    add_column_if_missing("evidencias", "contenido_b64", "TEXT")
    add_column_if_missing("usuarios", "intentos_fallidos", "INTEGER DEFAULT 0")
    add_column_if_missing("usuarios", "bloqueado_hasta", "TEXT")

    row = db_execute("SELECT COUNT(*) AS n FROM usuarios", fetchone=True)
    count = row["n"] if row is not None else 0
    if count == 0:
        admin_user, admin_password, admin_name, admin_email = initial_admin_config()
        salt, digest = hash_password(admin_password)
        db_execute(
            """
            INSERT INTO usuarios(usuario, nombre_completo, email, rol, salt, password_hash, activo, debe_cambiar_clave, creado_en, actualizado_en)
            VALUES (?, ?, ?, ?, ?, ?, 1, 1, ?, ?)
            """,
            (admin_user, admin_name, admin_email, "Administrador", salt, digest, ahora_iso(), ahora_iso()),
        )


def cursos_visibles_query(user: Dict[str, str]) -> Tuple[str, Tuple]:
    rol = user.get("rol", "")
    if rol in ("Administrador", "Coordinador"):
        return "SELECT * FROM cursos ORDER BY actualizado_en DESC, creado_en DESC", tuple()
    if rol == "Docente":
        return "SELECT * FROM cursos WHERE propietario_usuario=? OR creado_por=? ORDER BY actualizado_en DESC, creado_en DESC", (user.get("usuario", ""), user.get("usuario", ""))
    return "SELECT * FROM cursos ORDER BY actualizado_en DESC, creado_en DESC LIMIT 25", tuple()


def listar_cursos_visibles() -> pd.DataFrame:
    user = st.session_state.get("auth_user", {})
    q, params = cursos_visibles_query(user)
    df = read_sql_df(q, params=params)
    return df


def get_curso(curso_id: int) -> Optional[Dict]:
    conn = conexion_db()
    row = conn.execute("SELECT * FROM cursos WHERE id=?", (curso_id,)).fetchone()
    conn.close()
    return dict(row) if row else None


def upsert_curso(curso_id: Optional[int], datos: Dict[str, str], payload: Optional[Dict] = None) -> int:
    user = st.session_state.get("auth_user", {})
    now = ahora_iso()
    payload_json = json.dumps(payload or datos.get("payload", {}), ensure_ascii=False)
    values = {
        "codigo": datos.get("codigo", "").strip(),
        "grupo": datos.get("grupo", "").strip(),
        "asignatura": datos.get("asignatura", "").strip() or "Asignatura sin nombre",
        "programa": datos.get("programa", "").strip(),
        "periodo": datos.get("periodo", "").strip(),
        "profesor": datos.get("profesor", "").strip(),
        "email_profesor": datos.get("correo", datos.get("email_profesor", "")).strip(),
        "creditos": str(datos.get("creditos", "")).strip(),
        "htp": limpiar_numero(datos.get("htp", datos.get("horas_presenciales", 0))) or 0,
        "hti": limpiar_numero(datos.get("hti", datos.get("horas_independientes", 0))) or 0,
        "fecha_inicio": str(datos.get("fecha_inicio", "")),
        "fecha_fin": str(datos.get("fecha_fin", "")),
        "estado": datos.get("estado", "Planeación"),
        "avance_contenido": limpiar_numero(datos.get("avance_contenido", 0)) or 0,
        "avance_evaluado": limpiar_numero(datos.get("avance_evaluado", 0)) or 0,
        "propietario_usuario": datos.get("propietario_usuario", user.get("usuario", "")),
        "payload_json": payload_json,
    }
    conn = conexion_db()
    if curso_id:
        conn.execute(
            """
            UPDATE cursos SET codigo=?, grupo=?, asignatura=?, programa=?, periodo=?, profesor=?, email_profesor=?, creditos=?, htp=?, hti=?, fecha_inicio=?, fecha_fin=?, estado=?, avance_contenido=?, avance_evaluado=?, propietario_usuario=?, actualizado_en=?, payload_json=?
            WHERE id=?
            """,
            (values["codigo"], values["grupo"], values["asignatura"], values["programa"], values["periodo"], values["profesor"], values["email_profesor"], values["creditos"], values["htp"], values["hti"], values["fecha_inicio"], values["fecha_fin"], values["estado"], values["avance_contenido"], values["avance_evaluado"], values["propietario_usuario"], now, values["payload_json"], int(curso_id)),
        )
        new_id = int(curso_id)
        accion = "Actualizar expediente"
    else:
        params_insert = (values["codigo"], values["grupo"], values["asignatura"], values["programa"], values["periodo"], values["profesor"], values["email_profesor"], values["creditos"], values["htp"], values["hti"], values["fecha_inicio"], values["fecha_fin"], values["estado"], values["avance_contenido"], values["avance_evaluado"], values["propietario_usuario"], user.get("usuario", ""), now, now, values["payload_json"])
        if usar_postgres():
            cur = conn.execute(
                """
                INSERT INTO cursos(codigo, grupo, asignatura, programa, periodo, profesor, email_profesor, creditos, htp, hti, fecha_inicio, fecha_fin, estado, avance_contenido, avance_evaluado, propietario_usuario, creado_por, creado_en, actualizado_en, payload_json)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                RETURNING id
                """,
                params_insert,
            )
            new_id = int(cur.fetchone()["id"])
        else:
            cur = conn.execute(
                """
                INSERT INTO cursos(codigo, grupo, asignatura, programa, periodo, profesor, email_profesor, creditos, htp, hti, fecha_inicio, fecha_fin, estado, avance_contenido, avance_evaluado, propietario_usuario, creado_por, creado_en, actualizado_en, payload_json)
                VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """,
                params_insert,
            )
            new_id = int(cur.lastrowid)
        accion = "Crear expediente"
    conn.commit()
    conn.close()
    registrar_auditoria(accion, f"Curso ID={new_id} | {values['asignatura']} | Grupo {values['grupo']}")
    return new_id


def eliminar_curso(curso_id: int):
    conn = conexion_db()
    curso = conn.execute("SELECT asignatura, grupo FROM cursos WHERE id=?", (curso_id,)).fetchone()
    conn.execute("DELETE FROM cursos WHERE id=?", (curso_id,))
    conn.commit()
    conn.close()
    registrar_auditoria("Eliminar expediente", f"Curso ID={curso_id} | {dict(curso) if curso else ''}")


def registrar_artefacto(curso_id: Optional[int], tipo: str, nombre_archivo: str, descripcion: str = ""):
    conn = conexion_db()
    user = st.session_state.get("auth_user", {})
    conn.execute(
        "INSERT INTO artefactos(curso_id, tipo, nombre_archivo, descripcion, generado_por, generado_en) VALUES (?, ?, ?, ?, ?, ?)",
        (curso_id, tipo, nombre_archivo, descripcion, user.get("usuario", ""), ahora_iso()),
    )
    conn.commit()
    conn.close()


def seleccionar_curso_widget(label="Curso / expediente", key="curso_sel", incluir_nuevo=False) -> Optional[int]:
    df = listar_cursos_visibles()
    opciones = []
    if incluir_nuevo:
        opciones.append(None)
    opciones.extend(df["id"].tolist() if not df.empty else [])
    if not opciones:
        st.info("Aún no hay cursos registrados. Cree uno desde Expediente académico o desde Planeador superior.")
        return None
    def fmt(x):
        if x is None:
            return "➕ Crear nuevo expediente"
        fila = df.loc[df["id"] == x].iloc[0]
        return f"#{x} | {fila.get('asignatura','')} | Grupo {fila.get('grupo','')} | {fila.get('periodo','')} | {fila.get('estado','')}"
    return st.selectbox(label, opciones, format_func=fmt, key=key)


def build_ics_calendar(sesiones_df: pd.DataFrame, datos: Dict[str, str]) -> bytes:
    asignatura = datos.get("asignatura", "Clase")
    grupo = datos.get("grupo", "")
    lugar_default = datos.get("ambiente", "")
    stamp = datetime.utcnow().strftime("%Y%m%dT%H%M%SZ")
    lines = [
        "BEGIN:VCALENDAR",
        "VERSION:2.0",
        "PRODID:-//Gestor FDGC//Calendario Academico//ES",
        "CALSCALE:GREGORIAN",
        "METHOD:PUBLISH",
    ]
    for i, row in sesiones_df.reset_index(drop=True).iterrows():
        fecha = pd.to_datetime(row.get("Fecha", ""), errors="coerce")
        if pd.isna(fecha):
            continue
        horario = str(row.get("Horario", "")).strip()
        m = re.search(r"(\d{1,2}:\d{2})\s*[-–]\s*(\d{1,2}:\d{2})", horario)
        inicio = "08:00" if not m else m.group(1)
        fin = "10:00" if not m else m.group(2)
        dtstart = datetime.combine(fecha.date(), parse_time_value(inicio) or time(8,0)).strftime("%Y%m%dT%H%M%S")
        dtend = datetime.combine(fecha.date(), parse_time_value(fin) or time(10,0)).strftime("%Y%m%dT%H%M%S")
        uid = f"fdgc-{datos.get('codigo','curso')}-{datos.get('grupo','grupo')}-{i+1}@local"
        summary = f"{asignatura} {grupo} - Sesión {row.get('N° sesión', i+1)}"
        desc = f"Unidad: {row.get('Unidad','')}\\nContenido: {row.get('Contenido por desarrollar','')}\\nTrabajo presencial: {row.get('Descripción del trabajo presencial','')}\\nTrabajo independiente: {row.get('Descripción trabajo independiente','')}"
        location = str(row.get("Lugar / ambiente", "") or lugar_default)
        lines += [
            "BEGIN:VEVENT",
            f"UID:{uid}",
            f"DTSTAMP:{stamp}",
            f"DTSTART:{dtstart}",
            f"DTEND:{dtend}",
            f"SUMMARY:{summary}",
            f"DESCRIPTION:{desc}",
            f"LOCATION:{location}",
            "END:VEVENT",
        ]
    lines.append("END:VCALENDAR")
    return "\r\n".join(lines).encode("utf-8")


def validar_plan(modulos_df: pd.DataFrame, horarios_df: pd.DataFrame, sesiones_df: pd.DataFrame, evaluaciones_df: pd.DataFrame, datos: Dict[str, str]) -> Tuple[List[str], List[str]]:
    errores, alertas = [], []
    if not datos.get("asignatura"):
        errores.append("Falta la asignatura.")
    if not datos.get("grupo"):
        alertas.append("No se indicó grupo.")
    if limpiar_df(modulos_df, COLUMNAS_MODULOS).empty:
        errores.append("Debe registrar al menos una unidad/módulo.")
    if limpiar_df(horarios_df, COLUMNAS_HORARIOS).empty:
        errores.append("Debe registrar al menos un día y horario de clase.")
    ev = limpiar_df(evaluaciones_df, COLUMNAS_EVALUACIONES)
    total_eval = sum((limpiar_numero(v) or 0) for v in ev.get("Valor (%)", [])) if not ev.empty else 0
    if ev.empty:
        alertas.append("No hay evaluación concertada. El FD-GC71 saldrá sin plan evaluativo.")
    elif round(total_eval, 2) != 100:
        errores.append(f"La evaluación concertada suma {total_eval:.2f}%, debe sumar 100%.")
    if sesiones_df is not None and not sesiones_df.empty:
        fechas = pd.to_datetime(sesiones_df["Fecha"], errors="coerce")
        if fechas.isna().any():
            alertas.append("Hay sesiones sin fecha válida.")
        repetidas = sesiones_df.groupby(["Fecha", "Horario"]).size()
        if (repetidas > 1).any():
            alertas.append("Existen sesiones con la misma fecha y horario; revise si es intencional.")
    htp = limpiar_numero(datos.get("htp", 0)) or 0
    total_horas_mod = sum((limpiar_numero(v) or 0) for v in modulos_df.get("Horas presenciales", [])) if modulos_df is not None and not modulos_df.empty else 0
    if htp and total_horas_mod and abs(total_horas_mod - htp) > 2:
        alertas.append(f"Las horas presenciales por módulo suman {total_horas_mod:g}, pero la identificación reporta HTP={htp:g}.")
    return errores, alertas


def crear_paquete_curso_zip(datos: Dict[str, str], sesiones_df: pd.DataFrame, evaluaciones_df: pd.DataFrame, estudiantes_df: pd.DataFrame, representantes_df: pd.DataFrame, curso_id: Optional[int] = None) -> bytes:
    gc71 = crear_gc71_docx(datos, sesiones_df, evaluaciones_df, representantes_df)
    excel = crear_plantilla_evaluacion_xlsx(estudiantes_df, evaluaciones_df, datos)
    ics = build_ics_calendar(sesiones_df, datos)
    ses_csv = sesiones_df.to_csv(index=False).encode("utf-8-sig")
    ev_csv = evaluaciones_df.to_csv(index=False).encode("utf-8-sig")
    payload = {
        "datos": datos,
        "sesiones": df_to_payload(sesiones_df),
        "evaluaciones": df_to_payload(evaluaciones_df),
        "estudiantes": df_to_payload(estudiantes_df),
        "representantes": df_to_payload(representantes_df),
        "generado_en": ahora_iso(),
        "curso_id": curso_id,
    }
    buf = io.BytesIO()
    nombre_base = nombre_archivo_seguro(datos.get("asignatura", "curso"), date.today(), "FDGC_Paquete")
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
        z.writestr(f"{nombre_base}/01_FD-GC71_Guia_Didactica.docx", gc71)
        z.writestr(f"{nombre_base}/02_Plantilla_Evaluacion.xlsx", excel)
        z.writestr(f"{nombre_base}/03_Calendario_Clases.ics", ics)
        z.writestr(f"{nombre_base}/04_Sesiones.csv", ses_csv)
        z.writestr(f"{nombre_base}/05_Evaluaciones.csv", ev_csv)
        z.writestr(f"{nombre_base}/06_Payload_Reutilizable.json", json.dumps(payload, ensure_ascii=False, indent=2, default=str).encode("utf-8"))
        z.writestr(f"{nombre_base}/LEAME.txt", "Paquete generado desde Gestor FD-GC71/FD-GC72 Enterprise. Incluye guía, plantilla de notas, calendario, sesiones, evaluación y payload para trazabilidad.\n")
    return buf.getvalue()


def ui_inicio(st):
    user = st.session_state.get("auth_user", {})
    st.header("Panel de inicio")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Usuario", user.get("usuario", ""))
    c2.metric("Perfil", user.get("rol", ""))
    c3.metric("Estado", "Activo")
    cursos = listar_cursos_visibles()
    c4.metric("Cursos visibles", 0 if cursos.empty else len(cursos))
    st.info(ROLES_PERMISOS.get(user.get("rol"), {}).get("descripcion", ""))
    if get_app_env() == "production" and not usar_postgres():
        st.error("APP_ENV está en production pero no hay DATABASE_URL. Configure PostgreSQL antes de operar con datos reales.")
    elif usar_postgres():
        st.success("Persistencia productiva activa: PostgreSQL externo configurado.")
    if user.get("debe_cambiar_clave"):
        st.warning("Debe cambiar la contraseña inicial desde Mi cuenta antes de operar formalmente.")
    st.subheader("Ruta institucional recomendada")
    st.markdown(
        """
1. **Expediente académico:** registre el curso y deje trazabilidad de programa, grupo, periodo, docente y estado.
2. **Planeador superior:** defina unidades, intensidad, horario y evaluación; genere FD-GC71, calendario `.ics` y plantilla de notas en un solo paquete.
3. **Durante el curso:** cargue evidencias, use la plantilla de evaluación y conserve soportes por corte.
4. **FD-GC72:** al corte parcial y final, cargue listado/notas y genere el informe académico con análisis descriptivo.
5. **Validador institucional:** revise que evaluación sume 100%, que el cronograma sea consistente y que los soportes estén completos.
        """
    )
    st.caption("La idea es simple: menos copiar y pegar, más control. El sufrimiento académico no debe ser un requisito de grado.")


def ui_centro_control(st):
    st.header("Centro de control académico")
    df = listar_cursos_visibles()
    evid = read_sql_df("SELECT * FROM evidencias")
    art = read_sql_df("SELECT * FROM artefactos")
    aud = read_sql_df("SELECT * FROM auditoria ORDER BY id DESC LIMIT 200")
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Cursos", len(df))
    c2.metric("Evidencias", len(evid))
    c3.metric("Artefactos", len(art))
    c4.metric("Auditoría", len(aud))
    if df.empty:
        st.info("Aún no hay cursos registrados.")
        return
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Cursos por estado")
        estado = df.groupby("estado", dropna=False).size().reset_index(name="Cursos")
        st.bar_chart(estado.set_index("estado"))
    with col2:
        st.subheader("Avance promedio")
        avance = df[["asignatura", "avance_contenido", "avance_evaluado"]].copy()
        avance["Curso"] = avance["asignatura"].str.slice(0, 28)
        st.dataframe(avance[["Curso", "avance_contenido", "avance_evaluado"]], use_container_width=True, hide_index=True)
    st.subheader("Semáforo académico")
    sem = df.copy()
    sem["Riesgo"] = sem.apply(lambda r: "Alto" if (float(r.get("avance_contenido") or 0) < 50 and str(r.get("estado")) not in ["Planeación"]) else ("Medio" if float(r.get("avance_evaluado") or 0) < 30 and str(r.get("estado")) in ["En ejecución", "Corte parcial"] else "Bajo"), axis=1)
    st.dataframe(sem[["id", "codigo", "grupo", "asignatura", "periodo", "profesor", "estado", "avance_contenido", "avance_evaluado", "Riesgo"]], use_container_width=True, hide_index=True)


def ui_expediente_academico(st):
    st.header("Expediente académico del curso")
    st.caption("Registro maestro para no depender de archivos sueltos con nombres tipo FINAL_final_ahora_si_v8.docx. Todos hemos estado ahí.")
    df = listar_cursos_visibles()
    tab_lista, tab_crear, tab_editar = st.tabs(["Listado", "Crear expediente", "Editar / cerrar"])
    with tab_lista:
        if df.empty:
            st.info("No hay expedientes registrados.")
        else:
            filtro = st.text_input("Buscar por asignatura, código, grupo, docente o periodo")
            view = df.copy()
            if filtro:
                f = normalizar_texto(filtro)
                mask = view.apply(lambda row: f in normalizar_texto(" ".join(str(x) for x in row.values)), axis=1)
                view = view[mask]
            st.dataframe(view[["id", "codigo", "grupo", "asignatura", "programa", "periodo", "profesor", "estado", "avance_contenido", "avance_evaluado", "actualizado_en"]], use_container_width=True, hide_index=True)
            st.download_button("Descargar inventario CSV", view.to_csv(index=False).encode("utf-8-sig"), "inventario_expedientes.csv", "text/csv", use_container_width=True)
    with tab_crear:
        with st.form("crear_expediente"):
            c1, c2, c3 = st.columns(3)
            with c1:
                codigo = st.text_input("Código")
                grupo = st.text_input("Grupo")
                asignatura = st.text_input("Asignatura")
                programa = st.text_input("Programa académico")
            with c2:
                periodo = st.text_input("Período académico")
                profesor = st.text_input("Profesor", value=st.session_state.get("auth_user", {}).get("nombre_completo", ""))
                correo = st.text_input("Correo docente", value=st.session_state.get("auth_user", {}).get("email", ""))
                creditos = st.text_input("Número de créditos")
            with c3:
                htp = st.number_input("HTP semanal / referencia", min_value=0.0, step=0.5)
                hti = st.number_input("HTI semanal / referencia", min_value=0.0, step=0.5)
                fecha_inicio = st.date_input("Fecha inicio", value=date.today())
                fecha_fin = st.date_input("Fecha fin", value=date.today() + timedelta(days=110))
                estado = st.selectbox("Estado", ["Planeación", "En ejecución", "Corte parcial", "Cierre final", "Archivado"])
            enviar = st.form_submit_button("Crear expediente", use_container_width=True)
        if enviar:
            datos = dict(codigo=codigo, grupo=grupo, asignatura=asignatura, programa=programa, periodo=periodo, profesor=profesor, correo=correo, creditos=creditos, htp=htp, hti=hti, fecha_inicio=fecha_inicio, fecha_fin=fecha_fin, estado=estado)
            new_id = upsert_curso(None, datos, payload={"origen": "Expediente académico"})
            st.success(f"Expediente creado: #{new_id}")
            st.rerun()
    with tab_editar:
        curso_id = seleccionar_curso_widget("Seleccione expediente", key="edit_expediente")
        if curso_id:
            curso = get_curso(int(curso_id))
            with st.form("editar_expediente"):
                c1, c2, c3 = st.columns(3)
                with c1:
                    codigo = st.text_input("Código", value=curso.get("codigo") or "")
                    grupo = st.text_input("Grupo", value=curso.get("grupo") or "")
                    asignatura = st.text_input("Asignatura", value=curso.get("asignatura") or "")
                    programa = st.text_input("Programa", value=curso.get("programa") or "")
                with c2:
                    periodo = st.text_input("Periodo", value=curso.get("periodo") or "")
                    profesor = st.text_input("Profesor", value=curso.get("profesor") or "")
                    correo = st.text_input("Correo", value=curso.get("email_profesor") or "")
                    creditos = st.text_input("Créditos", value=curso.get("creditos") or "")
                with c3:
                    estado = st.selectbox("Estado", ["Planeación", "En ejecución", "Corte parcial", "Cierre final", "Archivado"], index=max(0, ["Planeación", "En ejecución", "Corte parcial", "Cierre final", "Archivado"].index(curso.get("estado") or "Planeación") if (curso.get("estado") or "Planeación") in ["Planeación", "En ejecución", "Corte parcial", "Cierre final", "Archivado"] else 0))
                    avance_contenido = st.slider("% avance contenido", 0, 100, int(float(curso.get("avance_contenido") or 0)))
                    avance_evaluado = st.slider("% evaluado", 0, 100, int(float(curso.get("avance_evaluado") or 0)))
                    fecha_inicio = st.text_input("Fecha inicio", value=curso.get("fecha_inicio") or "")
                    fecha_fin = st.text_input("Fecha fin", value=curso.get("fecha_fin") or "")
                guardar = st.form_submit_button("Guardar cambios", use_container_width=True)
            if guardar:
                payload = safe_json_loads(curso.get("payload_json"), {})
                datos = dict(codigo=codigo, grupo=grupo, asignatura=asignatura, programa=programa, periodo=periodo, profesor=profesor, correo=correo, creditos=creditos, fecha_inicio=fecha_inicio, fecha_fin=fecha_fin, estado=estado, avance_contenido=avance_contenido, avance_evaluado=avance_evaluado)
                upsert_curso(int(curso_id), datos, payload=payload)
                st.success("Expediente actualizado.")
                st.rerun()
            if st.session_state.get("auth_user", {}).get("rol") == "Administrador":
                with st.expander("Zona peligrosa"):
                    st.warning("Eliminar borra el expediente. Las evidencias físicas quedan en carpeta, pero sin vínculo visible.")
                    if st.button("Eliminar expediente", type="secondary"):
                        eliminar_curso(int(curso_id))
                        st.success("Expediente eliminado.")
                        st.rerun()


def ui_planeador_superior(st):
    st.header("Planeador superior FD-GC71")
    st.caption("Genera cronograma por horario real, valida evaluación, produce FD-GC71, Excel de notas, calendario ICS, CSV y paquete trazable.")
    curso_id = seleccionar_curso_widget("Usar expediente existente o crear desde cero", key="curso_planeador", incluir_nuevo=True)
    curso = get_curso(int(curso_id)) if curso_id else {}
    saved_payload = safe_json_loads(curso.get("payload_json") if curso else None, {})

    tab_datos, tab_plan, tab_generar = st.tabs(["1. Datos base", "2. Planeación y evaluación", "3. Validar y descargar"])
    with tab_datos:
        c1, c2, c3 = st.columns(3)
        with c1:
            programa = st.text_input("Programa académico", value=curso.get("programa", saved_payload.get("datos", {}).get("programa", "")) if curso else "")
            asignatura = st.text_input("Asignatura", value=curso.get("asignatura", saved_payload.get("datos", {}).get("asignatura", "")) if curso else "")
            codigo = st.text_input("Código", value=curso.get("codigo", saved_payload.get("datos", {}).get("codigo", "")) if curso else "")
            grupo = st.text_input("Grupo", value=curso.get("grupo", saved_payload.get("datos", {}).get("grupo", "")) if curso else "")
        with c2:
            profesor = st.text_input("Profesor", value=(curso.get("profesor") if curso else st.session_state.get("auth_user", {}).get("nombre_completo", "")) or "")
            correo = st.text_input("Correo", value=(curso.get("email_profesor") if curso else st.session_state.get("auth_user", {}).get("email", "")) or "")
            periodo = st.text_input("Período académico", value=curso.get("periodo", saved_payload.get("datos", {}).get("periodo", "")) if curso else "")
            creditos = st.text_input("Número de créditos", value=str(curso.get("creditos", "")) if curso else "")
        with c3:
            fecha_inicio = st.date_input("Inicio de clases", value=pd.to_datetime(curso.get("fecha_inicio") or date.today()).date() if curso and curso.get("fecha_inicio") else date.today())
            fecha_fin = st.date_input("Fin de clases", value=pd.to_datetime(curso.get("fecha_fin") or (date.today()+timedelta(days=110))).date() if curso and curso.get("fecha_fin") else date.today()+timedelta(days=110))
            htp = st.number_input("HTP total/referencia", min_value=0.0, step=0.5, value=float(curso.get("htp") or 0) if curso else 0.0)
            hti = st.number_input("HTI total/referencia", min_value=0.0, step=0.5, value=float(curso.get("hti") or 0) if curso else 0.0)
            estado = st.selectbox("Estado", ["Planeación", "En ejecución", "Corte parcial", "Cierre final", "Archivado"], index=0)
        st.subheader("Textos académicos con preforma editable")
        textos = {}
        for k, label in [
            ("justificacion", "Justificación"), ("competencias", "Competencias"), ("resultados", "Resultados de aprendizaje"),
            ("objetivo_general", "Objetivo general"), ("objetivos_especificos", "Objetivos específicos"), ("metodologias", "Metodologías y estrategias didácticas"),
            ("ambientes", "Ambientes de aprendizaje"), ("medios", "Medios educativos"), ("referencias", "Referencias bibliográficas"),
        ]:
            base = saved_payload.get("datos", {}).get(k, TEXTOS_PREDEFINIDOS_GC71.get(k, ""))
            textos[k] = st.text_area(label, value=base, height=90 if k not in ("objetivos_especificos", "referencias") else 120)

    with tab_plan:
        st.subheader("Unidades / módulos e intensidad")
        mod_default = payload_to_df(saved_payload.get("modulos"), COLUMNAS_MODULOS)
        if mod_default.empty:
            mod_default = pd.DataFrame([
                {"Unidad": "UNIDAD 1", "Contenido / tema central": "Fundamentos y contexto de la asignatura", "Horas presenciales": 8, "Sesiones": 4, "Trabajo presencial": "Clase orientadora, discusión guiada y taller aplicado.", "Trabajo independiente": "Lectura previa y desarrollo de actividad de preparación."},
                {"Unidad": "UNIDAD 2", "Contenido / tema central": "Métodos, herramientas y procedimientos", "Horas presenciales": 12, "Sesiones": 6, "Trabajo presencial": "Ejercicios prácticos y análisis de casos.", "Trabajo independiente": "Desarrollo de ejercicios y consulta bibliográfica."},
                {"Unidad": "UNIDAD 3", "Contenido / tema central": "Aplicación, integración y cierre", "Horas presenciales": 12, "Sesiones": 6, "Trabajo presencial": "Proyecto integrador y socialización de resultados.", "Trabajo independiente": "Preparación de entregables y retroalimentación final."},
            ])
        modulos_df = st.data_editor(mod_default, num_rows="dynamic", use_container_width=True, key="super_modulos")
        st.subheader("Horario de clase")
        hor_default = payload_to_df(saved_payload.get("horarios"), COLUMNAS_HORARIOS)
        if hor_default.empty:
            hor_default = pd.DataFrame([{ "Día": "Lunes", "Hora inicio": "18:00", "Hora fin": "20:00", "Lugar / ambiente": "Aula / plataforma institucional" }])
        horarios_df = st.data_editor(hor_default, num_rows="dynamic", use_container_width=True, key="super_horarios")
        criterio = st.radio("Criterio para repartir sesiones", ["Horas presenciales", "Sesiones"], horizontal=True)
        excluir = st.text_area("Fechas sin clase a excluir (una por línea, formato AAAA-MM-DD)", value="")
        fechas_excluidas = []
        for line in excluir.splitlines():
            try:
                fechas_excluidas.append(pd.to_datetime(line.strip()).date())
            except Exception:
                pass
        fechas_clase_df = generar_fechas_clase(fecha_inicio, fecha_fin, horarios_df, fechas_excluidas)
        sesiones_df = expandir_plan_sesiones(modulos_df, fechas_clase_df, criterio=criterio)
        st.subheader("Sesiones generadas")
        sesiones_edit = st.data_editor(sesiones_df, num_rows="dynamic", use_container_width=True, key="super_sesiones")
        st.subheader("Evaluación concertada")
        ev_default = payload_to_df(saved_payload.get("evaluaciones"), COLUMNAS_EVALUACIONES)
        if ev_default.empty:
            ev_default = pd.DataFrame([
                {"Tipo de evaluación": "Seguimiento", "Procedimiento de evaluación": "Talleres, actividades de clase y ejercicios aplicados", "Valor (%)": 30, "Fecha de realización": str(fecha_inicio + timedelta(days=30)), "Unidad relacionada": "UNIDAD 1", "Corte": "Primer corte"},
                {"Tipo de evaluación": "Parcial", "Procedimiento de evaluación": "Evaluación individual teórico-práctica", "Valor (%)": 30, "Fecha de realización": str(fecha_inicio + timedelta(days=60)), "Unidad relacionada": "UNIDAD 2", "Corte": "Segundo corte"},
                {"Tipo de evaluación": "Proyecto final", "Procedimiento de evaluación": "Entrega y sustentación de proyecto integrador", "Valor (%)": 40, "Fecha de realización": str(fecha_fin), "Unidad relacionada": "UNIDAD 3", "Corte": "Final"},
            ])
        evaluaciones_df = st.data_editor(ev_default, num_rows="dynamic", use_container_width=True, key="super_evaluaciones")
        st.subheader("Estudiantes / listado tradicional")
        uploaded_list = st.file_uploader("Cargar listado tradicional (.xls/.xlsx) opcional", type=["xls", "xlsx"], key="super_listado")
        if uploaded_list:
            estudiantes_df = leer_listado_estudiantes(uploaded_list)
            st.success(f"Listado leído: {len(estudiantes_df)} estudiantes.")
        else:
            estudiantes_df = payload_to_df(saved_payload.get("estudiantes"), ["Nombre completo", "Documento", "Correo", "Plan", "Observación", "Estado"])
            if estudiantes_df.empty:
                estudiantes_df = df_vacio(["Nombre completo", "Documento", "Correo", "Plan", "Observación", "Estado"], 5)
            estudiantes_df = st.data_editor(estudiantes_df, num_rows="dynamic", use_container_width=True, key="super_estudiantes")
        st.session_state["super_planeacion"] = dict(modulos=modulos_df, horarios=horarios_df, sesiones=sesiones_edit, evaluaciones=evaluaciones_df, estudiantes=estudiantes_df)

    with tab_generar:
        planeacion = st.session_state.get("super_planeacion", {})
        modulos_df = planeacion.get("modulos", pd.DataFrame(columns=COLUMNAS_MODULOS))
        horarios_df = planeacion.get("horarios", pd.DataFrame(columns=COLUMNAS_HORARIOS))
        sesiones_df = planeacion.get("sesiones", pd.DataFrame(columns=COLUMNAS_SESIONES))
        evaluaciones_df = planeacion.get("evaluaciones", pd.DataFrame(columns=COLUMNAS_EVALUACIONES))
        estudiantes_df = planeacion.get("estudiantes", pd.DataFrame())
        representantes_df = pd.DataFrame(columns=["Nombre de los estudiantes", "N° de cédula o carné estudiantil", "Firma"])
        datos = dict(programa=programa, asignatura=asignatura, codigo=codigo, grupo=grupo, profesor=profesor, correo=correo, periodo=periodo, creditos=creditos, htp=htp, hti=hti, fecha_inicio=fecha_inicio, fecha_fin=fecha_fin, estado=estado, **textos)
        errores, alertas = validar_plan(modulos_df, horarios_df, sesiones_df, evaluaciones_df, datos)
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Sesiones", 0 if sesiones_df is None else len(limpiar_df(sesiones_df, COLUMNAS_SESIONES)))
        c2.metric("Evaluación", f"{sum((limpiar_numero(v) or 0) for v in evaluaciones_df.get('Valor (%)', [])) if evaluaciones_df is not None and not evaluaciones_df.empty else 0:.0f}%")
        c3.metric("Estudiantes", 0 if estudiantes_df is None else len(limpiar_df(estudiantes_df, list(estudiantes_df.columns))))
        c4.metric("Errores", len(errores))
        if errores:
            st.error("Corrija antes de generar: " + " | ".join(errores))
        if alertas:
            st.warning("Alertas: " + " | ".join(alertas))
        if not errores:
            st.success("Plan listo para generar y guardar.")
            col_a, col_b, col_c, col_d = st.columns(4)
            gc71_bytes = crear_gc71_docx(datos, sesiones_df, evaluaciones_df, representantes_df)
            excel_bytes = crear_plantilla_evaluacion_xlsx(estudiantes_df, evaluaciones_df, datos)
            ics_bytes = build_ics_calendar(sesiones_df, datos)
            paquete_bytes = crear_paquete_curso_zip(datos, sesiones_df, evaluaciones_df, estudiantes_df, representantes_df, curso_id)
            base = nombre_archivo_seguro(asignatura, date.today(), "FDGC")
            col_a.download_button("FD-GC71 Word", gc71_bytes, f"{base}_FD-GC71.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
            col_b.download_button("Plantilla Excel", excel_bytes, f"{base}_Evaluacion.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
            col_c.download_button("Calendario ICS", ics_bytes, f"{base}_Calendario.ics", "text/calendar", use_container_width=True)
            col_d.download_button("Paquete completo ZIP", paquete_bytes, f"{base}_Paquete.zip", "application/zip", use_container_width=True)
            if st.button("Guardar/actualizar expediente con esta planeación", use_container_width=True, type="primary"):
                payload = {
                    "datos": {k: str(v) for k, v in datos.items()},
                    "modulos": df_to_payload(modulos_df),
                    "horarios": df_to_payload(horarios_df),
                    "sesiones": df_to_payload(sesiones_df),
                    "evaluaciones": df_to_payload(evaluaciones_df),
                    "estudiantes": df_to_payload(estudiantes_df),
                    "guardado_en": ahora_iso(),
                }
                new_id = upsert_curso(int(curso_id) if curso_id else None, datos, payload=payload)
                registrar_artefacto(new_id, "FD-GC71", f"{base}_FD-GC71.docx", "Generado desde Planeador superior")
                registrar_artefacto(new_id, "Excel evaluación", f"{base}_Evaluacion.xlsx", "Generado desde Planeador superior")
                registrar_auditoria("Guardar planeación superior", f"Curso ID={new_id}")
                st.success(f"Planeación guardada en expediente #{new_id}.")
                st.rerun()


def ui_evidencias(st):
    st.header("Evidencias y soportes")
    curso_id = seleccionar_curso_widget("Curso", key="curso_evidencias")
    if not curso_id:
        return
    curso = get_curso(int(curso_id))
    st.write(f"**Curso:** {curso.get('asignatura')} | **Grupo:** {curso.get('grupo')} | **Periodo:** {curso.get('periodo')}")
    tab_subir, tab_listar = st.tabs(["Subir soporte", "Repositorio"])
    with tab_subir:
        tipo = st.selectbox("Tipo de evidencia", ["Listado de clase", "Evaluación parcial", "Evaluación final", "Acta de concertación", "Asistencia", "Soporte metodológico", "Otro"])
        descripcion = st.text_area("Descripción / observación")
        files = st.file_uploader("Archivos", accept_multiple_files=True, type=["pdf", "doc", "docx", "xls", "xlsx", "csv", "png", "jpg", "jpeg", "zip"])
        if st.button("Guardar evidencias", type="primary", use_container_width=True):
            if not files:
                st.error("Seleccione al menos un archivo.")
            else:
                conn = conexion_db()
                user = st.session_state.get("auth_user", {})
                curso_dir = EVIDENCE_DIR / str(curso_id)
                curso_dir.mkdir(parents=True, exist_ok=True)
                for f in files:
                    raw = f.getvalue()
                    if len(raw) > max_evidence_bytes():
                        st.error(f"El archivo {f.name} pesa {len(raw)/1024/1024:.1f} MB y supera el límite configurado de {max_evidence_bytes()/1024/1024:.0f} MB.")
                        continue
                    ext = Path(f.name).suffix.lower()
                    internal = f"{uuid.uuid4().hex}{ext}"
                    # Caché local para descarga rápida. En Streamlit Cloud no se asume persistente.
                    (curso_dir / internal).write_bytes(raw)
                    contenido_b64 = base64.b64encode(raw).decode("ascii")
                    conn.execute(
                        "INSERT INTO evidencias(curso_id, tipo, nombre_original, nombre_archivo, mime, tamano, descripcion, subido_por, subido_en, contenido_b64) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)",
                        (int(curso_id), tipo, f.name, internal, getattr(f, "type", ""), len(raw), descripcion, user.get("usuario", ""), ahora_iso(), contenido_b64),
                    )
                conn.commit(); conn.close()
                registrar_auditoria("Subir evidencia", f"Curso ID={curso_id} | {len(files)} archivo(s) | Tipo={tipo}")
                st.success("Evidencias guardadas.")
                st.rerun()
    with tab_listar:
        df = read_sql_df("SELECT id, tipo, nombre_original, nombre_archivo, tamano, descripcion, subido_por, subido_en, contenido_b64 FROM evidencias WHERE curso_id=? ORDER BY id DESC", params=(int(curso_id),))
        if df.empty:
            st.info("Este curso aún no tiene evidencias.")
        else:
            cols_hide = [c for c in ["nombre_archivo", "contenido_b64"] if c in df.columns]
            st.dataframe(df.drop(columns=cols_hide), use_container_width=True, hide_index=True)
            for _, r in df.iterrows():
                path = EVIDENCE_DIR / str(curso_id) / str(r["nombre_archivo"])
                data = None
                if path.exists():
                    data = path.read_bytes()
                elif str(r.get("contenido_b64", "") or "").strip():
                    try:
                        data = base64.b64decode(str(r.get("contenido_b64")))
                    except Exception:
                        data = None
                if data is not None:
                    with st.expander(f"Descargar #{r['id']} - {r['nombre_original']}"):
                        st.write(r.get("descripcion", ""))
                        st.download_button("Descargar archivo", data, r["nombre_original"], use_container_width=True, key=f"download_evid_{r['id']}")
                else:
                    st.warning(f"No se encontró el contenido de la evidencia #{r['id']}: {r['nombre_original']}")


def ui_validador_institucional(st):
    st.header("Validador institucional")
    st.caption("Revisa coherencia básica antes de entregar: evaluación, listado, calificaciones y estructura. No reemplaza al comité; evita que el comité te mire feo.")
    tab_excel, tab_expediente = st.tabs(["Validar Excel de evaluación", "Validar expediente"])
    with tab_excel:
        f = st.file_uploader("Cargue plantilla de evaluación diligenciada", type=["xlsx", "xls"], key="validar_excel")
        corte = st.number_input("Nota mínima aprobatoria", min_value=0.0, max_value=5.0, value=3.0, step=0.1)
        if f:
            try:
                sheets = pd.read_excel(f, sheet_name=None)
                hallazgos = []
                if "Evaluaciones" not in sheets:
                    hallazgos.append(("Error", "No existe hoja Evaluaciones."))
                else:
                    ev = sheets["Evaluaciones"]
                    col_valor = encontrar_columna(ev, ["valor", "%"])
                    if col_valor:
                        total = pd.to_numeric(ev[col_valor], errors="coerce").fillna(0).sum()
                        if abs(total - 100) > 0.01:
                            hallazgos.append(("Error", f"La evaluación suma {total:.2f}%, debe sumar 100%."))
                        else:
                            hallazgos.append(("OK", "La evaluación suma 100%."))
                    else:
                        hallazgos.append(("Alerta", "No se identificó columna de valor porcentual."))
                if "Calificaciones" not in sheets:
                    hallazgos.append(("Error", "No existe hoja Calificaciones."))
                else:
                    cal = sheets["Calificaciones"]
                    if len(cal) == 0:
                        hallazgos.append(("Alerta", "La hoja Calificaciones no tiene estudiantes."))
                    notas_cols = [c for c in cal.columns if normalizar_texto(c).startswith("E") or "NOTA" in normalizar_texto(c)]
                    hallazgos.append(("Info", f"Filas de calificaciones: {len(cal)} | columnas de notas detectadas: {len(notas_cols)}."))
                    # Conteos aproximados
                    notas_num = pd.to_numeric(cal[notas_cols[-1]], errors="coerce") if notas_cols else pd.Series(dtype=float)
                    if not notas_num.empty:
                        hallazgos.append(("Info", f"Aprobados aproximados: {(notas_num >= corte).sum()} | Reprobados aproximados: {(notas_num < corte).sum()}."))
                st.dataframe(pd.DataFrame(hallazgos, columns=["Nivel", "Hallazgo"]), use_container_width=True, hide_index=True)
            except Exception as e:
                st.error(f"No se pudo validar: {e}")
    with tab_expediente:
        curso_id = seleccionar_curso_widget("Curso a validar", key="validar_curso")
        if curso_id:
            curso = get_curso(int(curso_id))
            payload = safe_json_loads(curso.get("payload_json"), {})
            modulos = payload_to_df(payload.get("modulos"), COLUMNAS_MODULOS)
            horarios = payload_to_df(payload.get("horarios"), COLUMNAS_HORARIOS)
            sesiones = payload_to_df(payload.get("sesiones"), COLUMNAS_SESIONES)
            evaluaciones = payload_to_df(payload.get("evaluaciones"), COLUMNAS_EVALUACIONES)
            datos = payload.get("datos", dict(curso))
            errores, alertas = validar_plan(modulos, horarios, sesiones, evaluaciones, datos)
            evid_row = db_execute("SELECT COUNT(*) AS n FROM evidencias WHERE curso_id=?", (int(curso_id),), fetchone=True)
            evid_count = evid_row["n"] if evid_row is not None else 0
            hallazgos = [("Error", e) for e in errores] + [("Alerta", a) for a in alertas]
            hallazgos.append(("Info", f"Evidencias asociadas: {evid_count}."))
            if not errores:
                hallazgos.append(("OK", "El expediente pasa las validaciones estructurales básicas."))
            st.dataframe(pd.DataFrame(hallazgos, columns=["Nivel", "Hallazgo"]), use_container_width=True, hide_index=True)


def ui_backup(st):
    st.header("Copias y restauración")
    st.caption("Respaldo portable: exporta tablas, evidencias y manifiesto. En producción esto no se improvisa el viernes a las 5:59 p. m.")
    if st.session_state.get("auth_user", {}).get("rol") != "Administrador":
        st.warning("Solo el perfil Administrador puede generar respaldos completos.")
        return

    def build_backup() -> bytes:
        buf = io.BytesIO()
        manifest = {
            "generado_en": ahora_iso(),
            "database": "PostgreSQL" if usar_postgres() else "SQLite",
            "version": APP_VERSION,
            "incluye": ["tablas_csv", "manifiesto", "evidencias_db", "cache_local_si_existe"],
        }
        tablas = ["usuarios", "auditoria", "cursos", "evidencias", "artefactos"]
        with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as z:
            for tabla in tablas:
                try:
                    df = read_sql_df(f"SELECT * FROM {tabla}")
                    z.writestr(f"tablas/{tabla}.csv", df.to_csv(index=False).encode("utf-8-sig"))
                except Exception as exc:
                    z.writestr(f"tablas/{tabla}_ERROR.txt", str(exc))
            # Archivo SQLite completo si aplica.
            if DB_PATH.exists() and not usar_postgres():
                z.write(DB_PATH, f"app_data/{DB_PATH.name}")
            # Evidencias desde DB para que funcione en Streamlit Cloud.
            try:
                evid = read_sql_df("SELECT id, curso_id, nombre_original, nombre_archivo, contenido_b64 FROM evidencias")
                for _, r in evid.iterrows():
                    data = None
                    if str(r.get("contenido_b64", "") or "").strip():
                        data = base64.b64decode(str(r.get("contenido_b64")))
                    else:
                        path = EVIDENCE_DIR / str(r.get("curso_id")) / str(r.get("nombre_archivo"))
                        if path.exists():
                            data = path.read_bytes()
                    if data is not None:
                        safe_name = re.sub(r"[^A-Za-z0-9_. -]", "_", str(r.get("nombre_original", "evidencia")))
                        z.writestr(f"evidencias/curso_{r.get('curso_id')}/id_{r.get('id')}_{safe_name}", data)
            except Exception as exc:
                z.writestr("evidencias_ERROR.txt", str(exc))
            # Cache local de exportaciones si existe.
            for folder in [EXPORT_DIR]:
                if folder.exists():
                    for file in folder.rglob("*"):
                        if file.is_file():
                            z.write(file, str(file.relative_to(APP_DIR)))
            z.writestr("manifest_backup.json", json.dumps(manifest, ensure_ascii=False, indent=2))
        return buf.getvalue()

    backup_bytes = build_backup()
    st.download_button("Descargar respaldo completo", backup_bytes, f"backup_fdgc_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip", "application/zip", use_container_width=True)
    st.divider()
    st.subheader("Estado técnico")
    st.json(health_status())
    st.write(f"Tamaño del respaldo actual: {len(backup_bytes)/1024:.1f} KB")


def ui_ayuda(st):
    st.header("Ayuda / flujo recomendado")
    st.markdown(
        """
### Flujo de operación recomendado

**Inicio de semestre**
- Crear expediente del curso.
- Definir unidades, intensidad, horario y evaluación concertada en Planeador superior.
- Descargar paquete completo: FD-GC71, Excel de evaluación, calendario ICS, CSV y JSON de trazabilidad.

**Mitad de curso**
- Cargar la plantilla de evaluación diligenciada.
- Usar FD-GC72 para calcular aprobados, reprobados, avance y análisis descriptivo.
- Guardar soportes en Evidencias.

**Final de curso**
- Actualizar evaluación final.
- Generar FD-GC72 definitivo.
- Validar expediente y descargar respaldo.

### Reglas de calidad
- La evaluación debe sumar exactamente 100%.
- Cada sesión debe tener fecha, contenido, trabajo presencial e independiente.
- La plantilla de evaluación debe conservar las hojas generadas por la app.
- El listado tradicional puede cargarse al inicio, mitad o final para alimentar los cálculos.
        """
    )


def ui_diagnostico_productivo(st):
    st.header("Diagnóstico productivo")
    st.caption("Lectura rápida del estado técnico del despliegue.")
    status = health_status()
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Versión", status.get("version", ""))
    c2.metric("Ambiente", status.get("app_env", ""))
    c3.metric("Base", status.get("database", ""))
    c4.metric("Plantillas", "OK" if status.get("templates_ok") else "Revisar")
    if get_app_env() == "production" and not usar_postgres():
        st.error("Modo producción sin PostgreSQL. Esto sirve para demo, no para operación institucional.")
    elif usar_postgres() and status.get("database_ok"):
        st.success("Base de datos externa operativa.")
    elif not status.get("database_ok"):
        st.error("No hay conexión sana a la base de datos.")
    st.json(status)
    st.subheader("Validación de repositorio")
    checks = [
        ("requirements.txt", (APP_DIR / "requirements.txt").exists()),
        ("runtime.txt", (APP_DIR / "runtime.txt").exists()),
        (".gitignore", (APP_DIR / ".gitignore").exists()),
        ("secrets.example.toml", (APP_DIR / ".streamlit" / "secrets.example.toml").exists()),
        ("FD-GC71.docx", TEMPLATE_GC71.exists()),
        ("FD-GC72.docx", TEMPLATE_GC72.exists()),
    ]
    st.dataframe(pd.DataFrame(checks, columns=["Elemento", "OK"]), use_container_width=True, hide_index=True)


def main():
    import streamlit as st
    globals()["st"] = st
    init_db()

    if "auth_user" not in st.session_state:
        pantalla_login(st)
        return

    user = st.session_state.get("auth_user", {})
    st.set_page_config(page_title="Gestor Académico Enterprise FD-GC71 / FD-GC72", layout="wide")
    st.title("Gestor académico Enterprise FD-GC71 / FD-GC72")
    st.caption("Planeación inteligente, expediente por curso, evaluación, informes, evidencias, auditoría, perfiles y respaldo.")

    with st.sidebar:
        st.markdown(f"**{user.get('nombre_completo', '')}**")
        st.caption(f"Perfil: {user.get('rol', '')}")
        if st.button("Cerrar sesión", use_container_width=True):
            registrar_auditoria("Logout", "Cierre de sesión")
            st.session_state.pop("auth_user", None)
            st.rerun()
        st.divider()
        modulos = ROLES_PERMISOS.get(user.get("rol"), {}).get("modulos", ["Inicio", "Mi cuenta", "Ayuda / flujo recomendado"])
        pagina = st.radio("Módulo", modulos)

    if not tiene_permiso(pagina):
        st.error("Este perfil no tiene permisos para abrir este módulo.")
        return

    if pagina == "Inicio":
        ui_inicio(st)
    elif pagina == MODULO_CENTRO:
        ui_centro_control(st)
    elif pagina == MODULO_EXPEDIENTE:
        ui_expediente_academico(st)
    elif pagina == MODULO_PLANEADOR:
        ui_planeador_superior(st)
    elif pagina.startswith("FD-GC71"):
        ui_gc71(st)
    elif pagina.startswith("FD-GC72"):
        ui_gc72(st)
    elif pagina == MODULO_EVIDENCIAS:
        ui_evidencias(st)
    elif pagina == MODULO_VALIDADOR:
        ui_validador_institucional(st)
    elif pagina == MODULO_BACKUP:
        ui_backup(st)
    elif pagina == MODULO_DIAGNOSTICO:
        ui_diagnostico_productivo(st)
    elif pagina == "Usuarios y perfiles":
        ui_admin_usuarios(st)
    elif pagina == "Auditoría":
        ui_auditoria(st)
    elif pagina == "Mi cuenta":
        ui_mi_cuenta(st)
    else:
        ui_ayuda(st)



# =============================================================================
# VERSIÓN 5.0 / SIGUIENTE NIVEL PRODUCTIVO
# Gobierno académico: aprobaciones, versionamiento, observaciones, parámetros,
# motor de calidad, reportes ejecutivos y trazabilidad con hash documental.
# =============================================================================

APP_VERSION = "5.0.0-next-level-cloud-production"

MODULO_FLUJO = "Aprobaciones y versionamiento"
MODULO_MOTOR = "Motor académico"
MODULO_REPORTES = "Reportes ejecutivos"
MODULO_PARAMETROS = "Parámetros institucionales"

ESTADOS_CURSO = [
    "Borrador",
    "En revisión",
    "Observado",
    "Ajustado",
    "Aprobado",
    "Cerrado",
]

TRANSICIONES_ESTADO = {
    "Borrador": ["En revisión"],
    "En revisión": ["Observado", "Aprobado"],
    "Observado": ["Ajustado"],
    "Ajustado": ["En revisión", "Aprobado"],
    "Aprobado": ["Cerrado", "Observado"],
    "Cerrado": [],
    "Planeación": ["En revisión"],
}

ROLES_PERMISOS = {
    "Administrador": {
        "descripcion": "Gobierno total: usuarios, expedientes, aprobaciones, parámetros, reportes, auditoría y respaldo.",
        "modulos": [
            "Inicio", MODULO_CENTRO, MODULO_EXPEDIENTE, MODULO_PLANEADOR,
            MODULO_FLUJO, MODULO_MOTOR, MODULO_REPORTES, MODULO_PARAMETROS,
            "FD-GC71 - Planeación", "FD-GC72 - Informe académico", MODULO_EVIDENCIAS,
            MODULO_VALIDADOR, MODULO_BACKUP, MODULO_DIAGNOSTICO,
            "Usuarios y perfiles", "Auditoría", "Mi cuenta", "Ayuda / flujo recomendado",
        ],
    },
    "Coordinador": {
        "descripcion": "Revisión académica: aprueba, observa, consulta reportes, valida consistencia y cierra expedientes.",
        "modulos": [
            "Inicio", MODULO_CENTRO, MODULO_EXPEDIENTE, MODULO_PLANEADOR,
            MODULO_FLUJO, MODULO_MOTOR, MODULO_REPORTES,
            "FD-GC71 - Planeación", "FD-GC72 - Informe académico", MODULO_EVIDENCIAS,
            MODULO_VALIDADOR, MODULO_DIAGNOSTICO, "Auditoría", "Mi cuenta", "Ayuda / flujo recomendado",
        ],
    },
    "Docente": {
        "descripcion": "Operación docente: planea, ajusta, carga evidencias, genera formatos y responde observaciones.",
        "modulos": [
            "Inicio", MODULO_CENTRO, MODULO_EXPEDIENTE, MODULO_PLANEADOR,
            MODULO_FLUJO, MODULO_MOTOR,
            "FD-GC71 - Planeación", "FD-GC72 - Informe académico", MODULO_EVIDENCIAS,
            MODULO_VALIDADOR, "Mi cuenta", "Ayuda / flujo recomendado",
        ],
    },
    "Consulta": {
        "descripcion": "Consulta controlada: tablero, reportes generales, cuenta propia y ayuda.",
        "modulos": ["Inicio", MODULO_CENTRO, MODULO_REPORTES, "Mi cuenta", "Ayuda / flujo recomendado"],
    },
}

_base_init_db_v4 = init_db
_base_upsert_curso_v4 = upsert_curso
_base_health_status_v4 = health_status
_base_crear_paquete_curso_zip_v4 = crear_paquete_curso_zip


def _dict_row(row) -> Dict[str, Any]:
    if row is None:
        return {}
    if isinstance(row, dict):
        return row
    return dict(row)


def app_bool(value) -> bool:
    if isinstance(value, bool):
        return value
    if value is None:
        return False
    return str(value).strip().lower() in {"1", "true", "sí", "si", "yes", "activo", "ok"}


def json_estable(value: Any) -> str:
    return json.dumps(value or {}, ensure_ascii=False, sort_keys=True, default=str, separators=(",", ":"))


def hash_documental(value: Any) -> str:
    return hashlib.sha256(json_estable(value).encode("utf-8")).hexdigest()


def init_db():
    """Inicializa la versión cloud productiva y agrega gobierno académico superior."""
    _base_init_db_v4()
    conn = conexion_db()
    try:
        if usar_postgres():
            tablas = [
                """
                CREATE TABLE IF NOT EXISTS curso_versiones (
                    id SERIAL PRIMARY KEY,
                    curso_id INTEGER REFERENCES cursos(id) ON DELETE CASCADE,
                    version_num INTEGER NOT NULL,
                    estado TEXT,
                    accion TEXT NOT NULL,
                    nota TEXT,
                    hash_payload TEXT NOT NULL,
                    payload_json TEXT DEFAULT '{}',
                    creado_por TEXT,
                    creado_en TEXT NOT NULL
                )
                """,
                """
                CREATE TABLE IF NOT EXISTS observaciones_curso (
                    id SERIAL PRIMARY KEY,
                    curso_id INTEGER REFERENCES cursos(id) ON DELETE CASCADE,
                    prioridad TEXT DEFAULT 'Media',
                    categoria TEXT DEFAULT 'Académica',
                    descripcion TEXT NOT NULL,
                    estado TEXT DEFAULT 'Abierta',
                    creado_por TEXT,
                    creado_en TEXT NOT NULL,
                    resuelto_por TEXT,
                    resuelto_en TEXT,
                    respuesta TEXT
                )
                """,
                """
                CREATE TABLE IF NOT EXISTS tareas_alertas (
                    id SERIAL PRIMARY KEY,
                    curso_id INTEGER REFERENCES cursos(id) ON DELETE CASCADE,
                    tipo TEXT,
                    mensaje TEXT NOT NULL,
                    fecha_limite TEXT,
                    estado TEXT DEFAULT 'Pendiente',
                    creado_por TEXT,
                    creado_en TEXT NOT NULL
                )
                """,
                """
                CREATE TABLE IF NOT EXISTS parametros_app (
                    clave TEXT PRIMARY KEY,
                    valor TEXT,
                    descripcion TEXT,
                    actualizado_por TEXT,
                    actualizado_en TEXT
                )
                """,
            ]
        else:
            tablas = [
                """
                CREATE TABLE IF NOT EXISTS curso_versiones (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    curso_id INTEGER,
                    version_num INTEGER NOT NULL,
                    estado TEXT,
                    accion TEXT NOT NULL,
                    nota TEXT,
                    hash_payload TEXT NOT NULL,
                    payload_json TEXT DEFAULT '{}',
                    creado_por TEXT,
                    creado_en TEXT NOT NULL,
                    FOREIGN KEY(curso_id) REFERENCES cursos(id)
                )
                """,
                """
                CREATE TABLE IF NOT EXISTS observaciones_curso (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    curso_id INTEGER,
                    prioridad TEXT DEFAULT 'Media',
                    categoria TEXT DEFAULT 'Académica',
                    descripcion TEXT NOT NULL,
                    estado TEXT DEFAULT 'Abierta',
                    creado_por TEXT,
                    creado_en TEXT NOT NULL,
                    resuelto_por TEXT,
                    resuelto_en TEXT,
                    respuesta TEXT,
                    FOREIGN KEY(curso_id) REFERENCES cursos(id)
                )
                """,
                """
                CREATE TABLE IF NOT EXISTS tareas_alertas (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    curso_id INTEGER,
                    tipo TEXT,
                    mensaje TEXT NOT NULL,
                    fecha_limite TEXT,
                    estado TEXT DEFAULT 'Pendiente',
                    creado_por TEXT,
                    creado_en TEXT NOT NULL,
                    FOREIGN KEY(curso_id) REFERENCES cursos(id)
                )
                """,
                """
                CREATE TABLE IF NOT EXISTS parametros_app (
                    clave TEXT PRIMARY KEY,
                    valor TEXT,
                    descripcion TEXT,
                    actualizado_por TEXT,
                    actualizado_en TEXT
                )
                """,
            ]
        for ddl in tablas:
            conn.execute(ddl)
        conn.execute("CREATE INDEX IF NOT EXISTS idx_versiones_curso ON curso_versiones(curso_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_observaciones_curso ON observaciones_curso(curso_id)")
        conn.execute("CREATE INDEX IF NOT EXISTS idx_tareas_curso ON tareas_alertas(curso_id)")
        conn.commit()
    finally:
        conn.close()
    seed_parametros_institucionales()


def seed_parametros_institucionales():
    defaults = {
        "umbral_riesgo_reprobacion": ("35", "Porcentaje de reprobación que activa alerta académica."),
        "dias_alerta_cierre": ("15", "Días antes de fin de curso para alerta de cierre documental."),
        "texto_metodologia_base": (TEXTOS_PREDEFINIDOS_GC71["metodologias"], "Texto base editable para metodologías."),
        "texto_analisis_positivo_base": (PREFORMAS_GC72["aspectos_positivos"]["Avance adecuado"], "Texto base para aspectos positivos del informe."),
        "texto_plan_mejora_base": (PREFORMAS_GC72["propuestas"]["Seguimiento formativo"], "Texto base para plan de mejora."),
    }
    for clave, (valor, desc) in defaults.items():
        existente = db_execute("SELECT clave FROM parametros_app WHERE clave=?", (clave,), fetchone=True)
        if existente is None:
            db_execute(
                "INSERT INTO parametros_app(clave, valor, descripcion, actualizado_por, actualizado_en) VALUES (?, ?, ?, ?, ?)",
                (clave, valor, desc, "sistema", ahora_iso()),
            )


def obtener_parametro(clave: str, default: str = "") -> str:
    row = db_execute("SELECT valor FROM parametros_app WHERE clave=?", (clave,), fetchone=True)
    if row is None:
        return default
    return str(_dict_row(row).get("valor", default) or default)


def actualizar_parametro(clave: str, valor: str, descripcion: str = ""):
    user = st.session_state.get("auth_user", {}) if "st" in globals() else {}
    existe = db_execute("SELECT clave FROM parametros_app WHERE clave=?", (clave,), fetchone=True)
    if existe is None:
        db_execute(
            "INSERT INTO parametros_app(clave, valor, descripcion, actualizado_por, actualizado_en) VALUES (?, ?, ?, ?, ?)",
            (clave, valor, descripcion, user.get("usuario", ""), ahora_iso()),
        )
    else:
        db_execute(
            "UPDATE parametros_app SET valor=?, descripcion=?, actualizado_por=?, actualizado_en=? WHERE clave=?",
            (valor, descripcion, user.get("usuario", ""), ahora_iso(), clave),
        )
    registrar_auditoria("Actualizar parámetro", clave)


def siguiente_version(curso_id: int) -> int:
    row = db_execute("SELECT COALESCE(MAX(version_num), 0) AS n FROM curso_versiones WHERE curso_id=?", (int(curso_id),), fetchone=True)
    return int(_dict_row(row).get("n", 0) or 0) + 1


def guardar_version_curso(curso_id: int, accion: str, nota: str = "", payload: Optional[Dict] = None, estado: Optional[str] = None):
    curso = get_curso(int(curso_id)) or {}
    payload = payload if payload is not None else safe_json_loads(curso.get("payload_json"), {})
    snapshot = {
        "curso": {k: v for k, v in curso.items() if k != "payload_json"},
        "payload": payload,
        "accion": accion,
        "estado": estado or curso.get("estado", ""),
    }
    digest = hash_documental(snapshot)
    user = st.session_state.get("auth_user", {}) if "st" in globals() else {}
    db_execute(
        """
        INSERT INTO curso_versiones(curso_id, version_num, estado, accion, nota, hash_payload, payload_json, creado_por, creado_en)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (int(curso_id), siguiente_version(int(curso_id)), estado or curso.get("estado", ""), accion, nota, digest, json.dumps(snapshot, ensure_ascii=False, default=str), user.get("usuario", "sistema"), ahora_iso()),
    )
    return digest


def upsert_curso(curso_id: Optional[int], datos: Dict[str, str], payload: Optional[Dict] = None) -> int:
    nuevo = curso_id is None
    cid = _base_upsert_curso_v4(curso_id, datos, payload)
    accion = "Creación de expediente" if nuevo else "Actualización de expediente"
    try:
        guardar_version_curso(cid, accion, "Versión automática al guardar", payload=payload, estado=datos.get("estado") or datos.get("Estado"))
    except Exception as exc:
        registrar_auditoria("Error versionamiento", f"Curso {cid}: {exc}")
    return cid


def cambiar_estado_curso(curso_id: int, nuevo_estado: str, nota: str = ""):
    curso = get_curso(int(curso_id)) or {}
    anterior = curso.get("estado", "")
    db_execute(
        "UPDATE cursos SET estado=?, actualizado_en=? WHERE id=?",
        (nuevo_estado, ahora_iso(), int(curso_id)),
    )
    guardar_version_curso(int(curso_id), f"Cambio de estado: {anterior} → {nuevo_estado}", nota, estado=nuevo_estado)
    registrar_auditoria("Cambio de estado", f"Curso ID={curso_id} | {anterior} → {nuevo_estado}")


def crear_observacion_curso(curso_id: int, prioridad: str, categoria: str, descripcion: str):
    user = st.session_state.get("auth_user", {}) if "st" in globals() else {}
    db_execute(
        """
        INSERT INTO observaciones_curso(curso_id, prioridad, categoria, descripcion, estado, creado_por, creado_en)
        VALUES (?, ?, ?, ?, 'Abierta', ?, ?)
        """,
        (int(curso_id), prioridad, categoria, descripcion, user.get("usuario", ""), ahora_iso()),
    )
    registrar_auditoria("Crear observación", f"Curso ID={curso_id} | {prioridad} | {categoria}")


def responder_observacion(obs_id: int, respuesta: str):
    user = st.session_state.get("auth_user", {}) if "st" in globals() else {}
    db_execute(
        "UPDATE observaciones_curso SET estado='Resuelta', respuesta=?, resuelto_por=?, resuelto_en=? WHERE id=?",
        (respuesta, user.get("usuario", ""), ahora_iso(), int(obs_id)),
    )
    registrar_auditoria("Resolver observación", f"Observación ID={obs_id}")


def observaciones_curso(curso_id: int) -> pd.DataFrame:
    return read_sql_df(
        "SELECT id, prioridad, categoria, descripcion, estado, creado_por, creado_en, resuelto_por, resuelto_en, respuesta FROM observaciones_curso WHERE curso_id=? ORDER BY id DESC",
        params=(int(curso_id),),
    )


def versiones_curso(curso_id: int) -> pd.DataFrame:
    return read_sql_df(
        "SELECT version_num, estado, accion, nota, hash_payload, creado_por, creado_en FROM curso_versiones WHERE curso_id=? ORDER BY version_num DESC",
        params=(int(curso_id),),
    )


def evidencias_count(curso_id: int) -> int:
    row = db_execute("SELECT COUNT(*) AS n FROM evidencias WHERE curso_id=?", (int(curso_id),), fetchone=True)
    return int(_dict_row(row).get("n", 0) or 0)


def score_calidad_expediente(curso_id: int) -> Tuple[int, List[Tuple[str, str, int]]]:
    """Retorna puntaje 0-100 y hallazgos: categoría, texto, peso perdido."""
    curso = get_curso(int(curso_id)) or {}
    payload = safe_json_loads(curso.get("payload_json"), {})
    modulos = payload_to_df(payload.get("modulos"), COLUMNAS_MODULOS)
    horarios = payload_to_df(payload.get("horarios"), COLUMNAS_HORARIOS)
    sesiones = payload_to_df(payload.get("sesiones"), COLUMNAS_SESIONES)
    evaluaciones = payload_to_df(payload.get("evaluaciones"), COLUMNAS_EVALUACIONES)
    datos = payload.get("datos", dict(curso))
    errores, alertas = validar_plan(modulos, horarios, sesiones, evaluaciones, datos)

    hallazgos: List[Tuple[str, str, int]] = []
    for e in errores:
        hallazgos.append(("Error", e, 12))
    for a in alertas:
        hallazgos.append(("Alerta", a, 5))
    if evidencias_count(int(curso_id)) == 0:
        hallazgos.append(("Alerta", "El expediente no tiene evidencias asociadas.", 8))
    abiertas = observaciones_curso(int(curso_id))
    abiertas_count = int((abiertas["estado"].astype(str) == "Abierta").sum()) if not abiertas.empty and "estado" in abiertas.columns else 0
    if abiertas_count:
        hallazgos.append(("Alerta", f"Tiene {abiertas_count} observación(es) abierta(s).", 7 * abiertas_count))
    estado = str(curso.get("estado", ""))
    if estado in {"Borrador", "Planeación"}:
        hallazgos.append(("Info", "El expediente aún está en etapa de construcción.", 3))
    if sesiones.empty:
        hallazgos.append(("Error", "No hay cronograma de sesiones guardado.", 15))
    if evaluaciones.empty:
        hallazgos.append(("Error", "No hay evaluación concertada guardada.", 15))

    penalizacion = min(100, sum(h[2] for h in hallazgos))
    return max(0, 100 - penalizacion), hallazgos


def riesgo_operativo_curso(row: pd.Series) -> Tuple[str, int, str]:
    puntos = 0
    razones = []
    estado = str(row.get("estado", "") or "")
    avance = limpiar_numero(row.get("avance_contenido")) or 0
    evaluado = limpiar_numero(row.get("avance_evaluado")) or 0
    if estado in ["Borrador", "Planeación", "Observado"]:
        puntos += 25; razones.append(f"estado {estado}")
    if avance < 50:
        puntos += 15; razones.append("avance bajo")
    if evaluado < 30:
        puntos += 10; razones.append("evaluación baja")
    try:
        fid = row.get("fecha_fin")
        if fid:
            ffin = datetime.fromisoformat(str(fid)[:10]).date()
            dias = (ffin - date.today()).days
            if dias < 0 and estado != "Cerrado":
                puntos += 35; razones.append("curso vencido sin cierre")
            elif dias <= int(obtener_parametro("dias_alerta_cierre", "15")) and estado not in ["Aprobado", "Cerrado"]:
                puntos += 20; razones.append("cierre cercano")
    except Exception:
        pass
    if puntos >= 60:
        return "Alto", puntos, ", ".join(razones)
    if puntos >= 30:
        return "Medio", puntos, ", ".join(razones)
    return "Bajo", puntos, ", ".join(razones) or "sin alertas críticas"


def matriz_riesgo_cursos() -> pd.DataFrame:
    df = listar_cursos_visibles()
    if df.empty:
        return df
    riesgos = df.apply(riesgo_operativo_curso, axis=1, result_type="expand")
    riesgos.columns = ["Riesgo", "Puntaje riesgo", "Razones"]
    out = pd.concat([df.reset_index(drop=True), riesgos], axis=1)
    if "id" in out.columns:
        out["Evidencias"] = out["id"].apply(lambda x: evidencias_count(int(x)))
        out["Score calidad"] = out["id"].apply(lambda x: score_calidad_expediente(int(x))[0])
    return out


def generar_sugerencias_academicas(curso_id: int) -> Dict[str, str]:
    curso = get_curso(int(curso_id)) or {}
    payload = safe_json_loads(curso.get("payload_json"), {})
    datos = payload.get("datos", dict(curso))
    modulos = payload_to_df(payload.get("modulos"), COLUMNAS_MODULOS)
    sesiones = payload_to_df(payload.get("sesiones"), COLUMNAS_SESIONES)
    evaluaciones = payload_to_df(payload.get("evaluaciones"), COLUMNAS_EVALUACIONES)

    temas = []
    if not modulos.empty:
        temas = [str(x).strip() for x in modulos.get("Contenido / tema central", []) if str(x).strip()]
    if not temas and not sesiones.empty:
        temas = [str(x).strip() for x in sesiones.get("Contenido por desarrollar", []) if str(x).strip()]
    temas_txt = "; ".join(temas[:8]) or "los contenidos definidos para la asignatura"
    asignatura = datos.get("asignatura") or curso.get("asignatura") or "la asignatura"
    programa = datos.get("programa") or curso.get("programa") or "el programa académico"
    total_eval = 0
    if not evaluaciones.empty and "Valor (%)" in evaluaciones.columns:
        total_eval = sum(limpiar_numero(v) or 0 for v in evaluaciones["Valor (%)"])

    justificacion = (
        f"La asignatura {asignatura} aporta al proceso formativo de {programa} mediante la integración de {temas_txt}. "
        "Su desarrollo permite articular fundamentos conceptuales, aplicación práctica, análisis de situaciones del contexto y producción de evidencias académicas verificables. "
        "La secuencia propuesta favorece el aprendizaje progresivo y la toma de decisiones sustentada en criterios técnicos, éticos y profesionales."
    )
    metodologia = (
        obtener_parametro("texto_metodologia_base", TEXTOS_PREDEFINIDOS_GC71["metodologias"]) + " "
        "Se recomienda mantener momentos de diagnóstico, desarrollo guiado, práctica aplicada, retroalimentación y cierre por evidencia, de forma que cada unidad deje productos verificables asociados a los resultados de aprendizaje."
    )
    plan_mejora = (
        "Fortalecer el seguimiento por cortes mediante alertas tempranas sobre inasistencia, bajo desempeño, actividades sin entregar y brechas conceptuales. "
        f"La evaluación concertada suma actualmente {total_eval:.1f}%, por lo que debe conservar coherencia entre actividad, fecha, unidad y evidencia esperada. "
        "El cierre del curso debe realizarse con listado, calificaciones, evidencias de socialización, informe académico y respaldo del expediente."
    )
    return {
        "Justificación sugerida": justificacion,
        "Metodología sugerida": metodologia,
        "Plan de mejora FD-GC72": plan_mejora,
    }


def crear_reporte_revision_markdown(curso_id: int) -> bytes:
    curso = get_curso(int(curso_id)) or {}
    score, hallazgos = score_calidad_expediente(int(curso_id))
    versiones = versiones_curso(int(curso_id))
    obs = observaciones_curso(int(curso_id))
    lines = [
        f"# Reporte de revisión académica - Curso {curso.get('id', curso_id)}",
        "",
        f"**Asignatura:** {curso.get('asignatura', '')}",
        f"**Grupo:** {curso.get('grupo', '')}",
        f"**Programa:** {curso.get('programa', '')}",
        f"**Periodo:** {curso.get('periodo', '')}",
        f"**Estado:** {curso.get('estado', '')}",
        f"**Score de calidad:** {score}/100",
        f"**Generado:** {ahora_iso()}",
        "",
        "## Hallazgos",
    ]
    if hallazgos:
        for nivel, texto, peso in hallazgos:
            lines.append(f"- **{nivel}** (-{peso}): {texto}")
    else:
        lines.append("- Sin hallazgos críticos.")
    lines += ["", "## Observaciones"]
    if obs.empty:
        lines.append("- Sin observaciones registradas.")
    else:
        for _, r in obs.iterrows():
            lines.append(f"- [{r.get('estado')}] {r.get('prioridad')} / {r.get('categoria')}: {r.get('descripcion')}")
    lines += ["", "## Versiones"]
    if versiones.empty:
        lines.append("- Sin versiones registradas.")
    else:
        for _, r in versiones.head(10).iterrows():
            lines.append(f"- v{r.get('version_num')} | {r.get('estado')} | {r.get('accion')} | hash `{str(r.get('hash_payload'))[:16]}...` | {r.get('creado_en')}")
    return "\n".join(lines).encode("utf-8")


def crear_reporte_ejecutivo_xlsx() -> bytes:
    df = matriz_riesgo_cursos()
    output = io.BytesIO()
    with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
        if df.empty:
            pd.DataFrame([{"Mensaje": "No hay cursos visibles para reportar."}]).to_excel(writer, sheet_name="Resumen", index=False)
        else:
            resumen = pd.DataFrame({
                "Indicador": ["Cursos", "Riesgo alto", "Riesgo medio", "Riesgo bajo", "Score calidad promedio"],
                "Valor": [
                    len(df),
                    int((df["Riesgo"] == "Alto").sum()),
                    int((df["Riesgo"] == "Medio").sum()),
                    int((df["Riesgo"] == "Bajo").sum()),
                    round(pd.to_numeric(df.get("Score calidad", pd.Series(dtype=float)), errors="coerce").mean(), 1),
                ],
            })
            resumen.to_excel(writer, sheet_name="Resumen", index=False)
            df.to_excel(writer, sheet_name="Cursos", index=False)
            workbook = writer.book
            fmt_header = workbook.add_format({"bold": True, "bg_color": "#D9EAF7", "border": 1})
            fmt_alert = workbook.add_format({"bg_color": "#FCE4D6", "border": 1})
            fmt_ok = workbook.add_format({"bg_color": "#E2F0D9", "border": 1})
            for ws_name in ["Resumen", "Cursos"]:
                ws = writer.sheets[ws_name]
                ws.freeze_panes(1, 0)
                ws.set_row(0, None, fmt_header)
                ws.autofilter(0, 0, max(1, len(df if ws_name == "Cursos" else resumen)), 12)
                ws.set_column(0, 12, 20)
            ws = writer.sheets["Cursos"]
            if "Riesgo" in df.columns:
                col = list(df.columns).index("Riesgo")
                ws.conditional_format(1, col, len(df), col, {"type": "text", "criteria": "containing", "value": "Alto", "format": fmt_alert})
                ws.conditional_format(1, col, len(df), col, {"type": "text", "criteria": "containing", "value": "Bajo", "format": fmt_ok})
    output.seek(0)
    return output.getvalue()


def crear_paquete_curso_zip(datos: Dict[str, str], sesiones_df: pd.DataFrame, evaluaciones_df: pd.DataFrame, estudiantes_df: pd.DataFrame, representantes_df: pd.DataFrame, curso_id: Optional[int] = None) -> bytes:
    paquete = _base_crear_paquete_curso_zip_v4(datos, sesiones_df, evaluaciones_df, estudiantes_df, representantes_df, curso_id)
    # Enriquecer el ZIP original con manifiesto de hash y reporte de revisión si existe curso_id.
    buf_in = io.BytesIO(paquete)
    buf_out = io.BytesIO()
    with zipfile.ZipFile(buf_in, "r") as zin, zipfile.ZipFile(buf_out, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            zout.writestr(item, zin.read(item.filename))
        manifiesto_superior = {
            "version_app": APP_VERSION,
            "generado_en": ahora_iso(),
            "hash_expediente": hash_documental({
                "datos": datos,
                "sesiones": df_to_payload(sesiones_df),
                "evaluaciones": df_to_payload(evaluaciones_df),
                "estudiantes": df_to_payload(estudiantes_df),
            }),
            "curso_id": curso_id,
            "reglas": [
                "La evaluación debe sumar 100%.",
                "Las sesiones deben tener fecha, contenido, trabajo presencial e independiente.",
                "El expediente debe conservar evidencias y versiones.",
            ],
        }
        zout.writestr("00_MANIFIESTO_SUPERIOR.json", json.dumps(manifiesto_superior, ensure_ascii=False, indent=2, default=str))
        if curso_id:
            try:
                zout.writestr("reporte_revision_academica.md", crear_reporte_revision_markdown(int(curso_id)))
            except Exception as exc:
                zout.writestr("reporte_revision_ERROR.txt", str(exc))
    buf_out.seek(0)
    return buf_out.getvalue()


def health_status() -> Dict[str, Any]:
    status = _base_health_status_v4()
    status["version"] = APP_VERSION
    extra = {}
    for table in ["curso_versiones", "observaciones_curso", "tareas_alertas", "parametros_app"]:
        try:
            row = db_execute(f"SELECT COUNT(*) AS n FROM {table}", fetchone=True)
            extra[table] = int(_dict_row(row).get("n", 0) or 0)
        except Exception as exc:
            extra[table] = f"ERROR: {str(exc)[:120]}"
    status["gobierno_academico"] = extra
    status["cloud_ready_level"] = "next-level"
    return status


def ui_centro_control(st):
    st.header("Centro de control académico")
    st.caption("Lectura ejecutiva de expedientes, riesgos, estados y calidad documental.")
    df = matriz_riesgo_cursos()
    if df.empty:
        st.info("No hay cursos visibles todavía. Cree un expediente desde Planeador superior.")
        return
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Cursos", len(df))
    c2.metric("Riesgo alto", int((df["Riesgo"] == "Alto").sum()))
    c3.metric("Observados", int((df.get("estado", pd.Series(dtype=str)).astype(str) == "Observado").sum()))
    c4.metric("Aprobados/cerrados", int(df.get("estado", pd.Series(dtype=str)).astype(str).isin(["Aprobado", "Cerrado"]).sum()))
    c5.metric("Score promedio", f"{pd.to_numeric(df.get('Score calidad', pd.Series(dtype=float)), errors='coerce').mean():.1f}")

    st.subheader("Mapa de riesgo")
    columnas = [c for c in ["id", "codigo", "grupo", "asignatura", "programa", "periodo", "profesor", "estado", "Riesgo", "Puntaje riesgo", "Score calidad", "Evidencias", "Razones"] if c in df.columns]
    st.dataframe(df[columnas].sort_values(["Riesgo", "Puntaje riesgo"], ascending=[True, False]), use_container_width=True, hide_index=True)

    st.subheader("Distribución por estado y riesgo")
    col1, col2 = st.columns(2)
    with col1:
        estado_counts = df["estado"].fillna("Sin estado").value_counts().rename_axis("Estado").reset_index(name="Cursos")
        st.bar_chart(estado_counts.set_index("Estado"))
    with col2:
        riesgo_counts = df["Riesgo"].fillna("Sin riesgo").value_counts().rename_axis("Riesgo").reset_index(name="Cursos")
        st.bar_chart(riesgo_counts.set_index("Riesgo"))


def ui_flujo_aprobaciones(st):
    st.header("Aprobaciones y versionamiento")
    st.caption("Control formal del expediente: envío, observaciones, aprobación, cierre, historial y hash documental.")
    curso_id = seleccionar_curso_widget("Expediente", key="flujo_curso")
    if not curso_id:
        return
    curso = get_curso(int(curso_id)) or {}
    estado = str(curso.get("estado", "Borrador") or "Borrador")
    user = st.session_state.get("auth_user", {})
    st.info(f"**Curso:** {curso.get('asignatura')} | **Grupo:** {curso.get('grupo')} | **Estado actual:** {estado}")
    score, hallazgos = score_calidad_expediente(int(curso_id))
    c1, c2, c3 = st.columns(3)
    c1.metric("Score calidad", f"{score}/100")
    c2.metric("Versiones", len(versiones_curso(int(curso_id))))
    c3.metric("Observaciones abiertas", int((observaciones_curso(int(curso_id)).get("estado", pd.Series(dtype=str)).astype(str) == "Abierta").sum()) if not observaciones_curso(int(curso_id)).empty else 0)

    tab_estado, tab_obs, tab_versiones, tab_reporte = st.tabs(["Flujo", "Observaciones", "Historial", "Reporte"])
    with tab_estado:
        posibles = TRANSICIONES_ESTADO.get(estado, TRANSICIONES_ESTADO.get("Borrador", []))
        if not posibles:
            st.success("El expediente no tiene transiciones pendientes.")
        else:
            st.write("Transiciones disponibles:")
            nota = st.text_area("Nota del cambio", placeholder="Explique brevemente el motivo del cambio de estado.")
            cols = st.columns(max(1, len(posibles)))
            for i, nuevo in enumerate(posibles):
                disabled = False
                if nuevo in ["Aprobado", "Cerrado"] and user.get("rol") not in ["Administrador", "Coordinador"]:
                    disabled = True
                if nuevo == "Aprobado" and score < 70:
                    st.warning("El score es inferior a 70. Puede aprobarse solo con criterio del coordinador, pero queda trazado.")
                with cols[i]:
                    if st.button(f"Mover a {nuevo}", disabled=disabled, use_container_width=True, key=f"estado_{nuevo}"):
                        cambiar_estado_curso(int(curso_id), nuevo, nota)
                        st.success(f"Estado actualizado a {nuevo}.")
                        st.rerun()
        if hallazgos:
            st.subheader("Hallazgos antes de aprobar")
            st.dataframe(pd.DataFrame(hallazgos, columns=["Nivel", "Hallazgo", "Penalización"]), use_container_width=True, hide_index=True)
        else:
            st.success("Sin hallazgos estructurales críticos.")
    with tab_obs:
        with st.form("crear_obs"):
            c1, c2 = st.columns(2)
            with c1:
                prioridad = st.selectbox("Prioridad", ["Alta", "Media", "Baja"], index=1)
            with c2:
                categoria = st.selectbox("Categoría", ["Académica", "Evaluación", "Cronograma", "Evidencia", "Forma", "Otra"])
            descripcion = st.text_area("Observación")
            enviar = st.form_submit_button("Registrar observación", use_container_width=True)
        if enviar:
            if descripcion.strip():
                crear_observacion_curso(int(curso_id), prioridad, categoria, descripcion)
                if estado not in ["Observado", "Cerrado"]:
                    cambiar_estado_curso(int(curso_id), "Observado", "Observación registrada")
                st.success("Observación registrada.")
                st.rerun()
            else:
                st.error("La observación no puede estar vacía.")
        obs = observaciones_curso(int(curso_id))
        if obs.empty:
            st.info("Sin observaciones.")
        else:
            st.dataframe(obs, use_container_width=True, hide_index=True)
            abiertas = obs[obs["estado"].astype(str) == "Abierta"] if "estado" in obs.columns else pd.DataFrame()
            if not abiertas.empty:
                obs_id = st.selectbox("Resolver observación", abiertas["id"].tolist(), format_func=lambda x: f"#{x} - {abiertas.loc[abiertas['id']==x, 'descripcion'].iloc[0][:70]}")
                respuesta = st.text_area("Respuesta / ajuste realizado")
                if st.button("Marcar como resuelta", use_container_width=True):
                    responder_observacion(int(obs_id), respuesta)
                    st.success("Observación resuelta.")
                    st.rerun()
    with tab_versiones:
        versiones = versiones_curso(int(curso_id))
        if versiones.empty:
            st.info("Aún no hay versiones registradas.")
        else:
            st.dataframe(versiones, use_container_width=True, hide_index=True)
            csv = versiones.to_csv(index=False).encode("utf-8-sig")
            st.download_button("Descargar historial CSV", csv, f"historial_versiones_curso_{curso_id}.csv", "text/csv", use_container_width=True)
    with tab_reporte:
        md = crear_reporte_revision_markdown(int(curso_id))
        st.download_button("Descargar reporte de revisión (.md)", md, f"reporte_revision_curso_{curso_id}.md", "text/markdown", use_container_width=True)
        st.code(md.decode("utf-8")[:4000], language="markdown")


def ui_motor_academico(st):
    st.header("Motor académico")
    st.caption("Asistente paramétrico para consistencia, redacción académica y cierre documental. No inventa: parte del expediente guardado.")
    curso_id = seleccionar_curso_widget("Curso", key="motor_curso")
    if not curso_id:
        return
    curso = get_curso(int(curso_id)) or {}
    score, hallazgos = score_calidad_expediente(int(curso_id))
    c1, c2, c3 = st.columns(3)
    c1.metric("Score de expediente", f"{score}/100")
    c2.metric("Estado", curso.get("estado", ""))
    c3.metric("Evidencias", evidencias_count(int(curso_id)))

    tab_calidad, tab_redaccion, tab_alertas = st.tabs(["Calidad", "Redacción asistida", "Alertas sugeridas"])
    with tab_calidad:
        if hallazgos:
            st.dataframe(pd.DataFrame(hallazgos, columns=["Nivel", "Hallazgo", "Penalización"]), use_container_width=True, hide_index=True)
        else:
            st.success("El expediente pasa las validaciones superiores.")
        if st.button("Registrar versión de control", use_container_width=True):
            digest = guardar_version_curso(int(curso_id), "Versión manual de control", "Generada desde Motor académico")
            st.success(f"Versión registrada. Hash: {digest[:24]}...")
    with tab_redaccion:
        sugerencias = generar_sugerencias_academicas(int(curso_id))
        for titulo, texto in sugerencias.items():
            st.text_area(titulo, texto, height=160, key=f"sug_{curso_id}_{titulo}")
        st.info("Copie estos textos al FD-GC71 o FD-GC72 si aplican. Son base técnica, no poesía académica con perfume.")
    with tab_alertas:
        df = matriz_riesgo_cursos()
        fila = df[df["id"].astype(str) == str(curso_id)] if not df.empty and "id" in df.columns else pd.DataFrame()
        if fila.empty:
            st.info("No hay alertas calculadas.")
        else:
            st.write(f"**Riesgo:** {fila.iloc[0].get('Riesgo')} | **Razones:** {fila.iloc[0].get('Razones')}")
            mensaje = st.text_area("Crear alerta interna", value=f"Revisar expediente {curso.get('asignatura', '')}: {fila.iloc[0].get('Razones')}")
            fecha_limite = st.date_input("Fecha límite", value=date.today() + timedelta(days=7), format="DD/MM/YYYY")
            if st.button("Guardar alerta", use_container_width=True):
                user = st.session_state.get("auth_user", {})
                db_execute(
                    "INSERT INTO tareas_alertas(curso_id, tipo, mensaje, fecha_limite, estado, creado_por, creado_en) VALUES (?, ?, ?, ?, 'Pendiente', ?, ?)",
                    (int(curso_id), "Revisión", mensaje, fecha_limite.isoformat(), user.get("usuario", ""), ahora_iso()),
                )
                registrar_auditoria("Crear alerta", f"Curso ID={curso_id}")
                st.success("Alerta guardada.")
        tareas = read_sql_df("SELECT id, tipo, mensaje, fecha_limite, estado, creado_por, creado_en FROM tareas_alertas WHERE curso_id=? ORDER BY id DESC", params=(int(curso_id),))
        if not tareas.empty:
            st.dataframe(tareas, use_container_width=True, hide_index=True)


def ui_reportes_ejecutivos(st):
    st.header("Reportes ejecutivos")
    st.caption("Corte institucional por curso, programa, periodo, estado, riesgo y calidad documental.")
    df = matriz_riesgo_cursos()
    if df.empty:
        st.info("No hay información para reportar.")
        return
    c1, c2, c3 = st.columns(3)
    with c1:
        periodos = ["Todos"] + sorted([str(x) for x in df.get("periodo", pd.Series(dtype=str)).dropna().unique()])
        periodo_sel = st.selectbox("Periodo", periodos)
    with c2:
        estados = ["Todos"] + sorted([str(x) for x in df.get("estado", pd.Series(dtype=str)).dropna().unique()])
        estado_sel = st.selectbox("Estado", estados)
    with c3:
        riesgos = ["Todos"] + sorted([str(x) for x in df.get("Riesgo", pd.Series(dtype=str)).dropna().unique()])
        riesgo_sel = st.selectbox("Riesgo", riesgos)
    f = df.copy()
    if periodo_sel != "Todos":
        f = f[f["periodo"].astype(str) == periodo_sel]
    if estado_sel != "Todos":
        f = f[f["estado"].astype(str) == estado_sel]
    if riesgo_sel != "Todos":
        f = f[f["Riesgo"].astype(str) == riesgo_sel]
    st.dataframe(f, use_container_width=True, hide_index=True)
    col1, col2 = st.columns(2)
    with col1:
        csv = f.to_csv(index=False).encode("utf-8-sig")
        st.download_button("Descargar reporte CSV", csv, "reporte_ejecutivo_fdgc.csv", "text/csv", use_container_width=True)
    with col2:
        xlsx = crear_reporte_ejecutivo_xlsx()
        st.download_button("Descargar reporte Excel", xlsx, "reporte_ejecutivo_fdgc.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)


def ui_parametros(st):
    st.header("Parámetros institucionales")
    st.caption("Textos base, umbrales y reglas editables sin tocar código. Esto es lo que vuelve mantenible el sistema.")
    if st.session_state.get("auth_user", {}).get("rol") != "Administrador":
        st.warning("Solo el Administrador puede modificar parámetros.")
        df = read_sql_df("SELECT clave, valor, descripcion, actualizado_por, actualizado_en FROM parametros_app ORDER BY clave")
        st.dataframe(df, use_container_width=True, hide_index=True)
        return
    df = read_sql_df("SELECT clave, valor, descripcion, actualizado_por, actualizado_en FROM parametros_app ORDER BY clave")
    st.dataframe(df, use_container_width=True, hide_index=True)
    st.subheader("Editar parámetro")
    claves = df["clave"].tolist() if not df.empty else []
    clave_sel = st.selectbox("Parámetro existente", ["Nuevo"] + claves)
    if clave_sel == "Nuevo":
        clave = st.text_input("Clave", value="nuevo_parametro")
        valor = st.text_area("Valor", height=120)
        descripcion = st.text_input("Descripción")
    else:
        fila = df[df["clave"] == clave_sel].iloc[0]
        clave = clave_sel
        valor = st.text_area("Valor", value=str(fila.get("valor", "")), height=160)
        descripcion = st.text_input("Descripción", value=str(fila.get("descripcion", "")))
    if st.button("Guardar parámetro", use_container_width=True):
        actualizar_parametro(clave, valor, descripcion)
        st.success("Parámetro actualizado.")
        st.rerun()


def ui_inicio(st):
    user = st.session_state.get("auth_user", {})
    st.header("Panel de inicio")
    status = health_status()
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Usuario", user.get("usuario", ""))
    c2.metric("Perfil", user.get("rol", ""))
    c3.metric("Ambiente", status.get("app_env", ""))
    c4.metric("Versión", APP_VERSION)
    if get_app_env() == "production" and not usar_postgres():
        st.error("APP_ENV está en production pero no hay DATABASE_URL. Configure PostgreSQL antes de operar con datos reales.")
    elif usar_postgres():
        st.success("Persistencia productiva activa: PostgreSQL externo configurado.")
    st.subheader("Ruta operativa superior")
    st.markdown(
        """
1. **Crear expediente** del curso en Planeador superior.
2. **Generar FD-GC71**, plantilla de evaluación, calendario y paquete ZIP.
3. **Enviar a revisión** desde Aprobaciones y versionamiento.
4. **Resolver observaciones** y conservar historial con hash documental.
5. **Cargar evidencias** de socialización, listado y evaluaciones.
6. **Generar FD-GC72** a mitad y cierre del curso.
7. **Validar, aprobar, cerrar y respaldar** el expediente.
        """
    )
    if user.get("debe_cambiar_clave"):
        st.warning("Debe cambiar la contraseña inicial desde Mi cuenta antes de operar formalmente.")


def ui_diagnostico_productivo(st):
    st.header("Diagnóstico productivo")
    st.caption("Estado técnico del despliegue, persistencia, plantillas y gobierno académico.")
    status = health_status()
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Versión", status.get("version", ""))
    c2.metric("Ambiente", status.get("app_env", ""))
    c3.metric("Base", status.get("database", ""))
    c4.metric("Nivel", status.get("cloud_ready_level", ""))
    if get_app_env() == "production" and not usar_postgres():
        st.error("Modo producción sin PostgreSQL. Esto sirve para demo, no para operación institucional.")
    elif usar_postgres() and status.get("database_ok"):
        st.success("Base de datos externa operativa.")
    elif not status.get("database_ok"):
        st.error("No hay conexión sana a la base de datos.")
    st.json(status)
    checks = [
        ("requirements.txt", (APP_DIR / "requirements.txt").exists()),
        ("runtime.txt", (APP_DIR / "runtime.txt").exists()),
        (".gitignore", (APP_DIR / ".gitignore").exists()),
        ("secrets.example.toml", (APP_DIR / ".streamlit" / "secrets.example.toml").exists()),
        ("FD-GC71.docx", TEMPLATE_GC71.exists()),
        ("FD-GC72.docx", TEMPLATE_GC72.exists()),
        ("Módulo versionamiento", True),
        ("Módulo aprobaciones", True),
        ("Motor académico", True),
    ]
    st.dataframe(pd.DataFrame(checks, columns=["Elemento", "OK"]), use_container_width=True, hide_index=True)


def ui_ayuda(st):
    st.header("Ayuda / flujo recomendado")
    st.markdown(
        """
### Flujo institucional completo

**Planeación**  
Cree el expediente, defina identificación, módulos, horarios, evaluación concertada y descargue FD-GC71.

**Revisión**  
Envíe el expediente a revisión. Coordinación puede observar, aprobar o devolver. Cada movimiento genera versión y hash.

**Ejecución**  
Cargue listado tradicional, evidencias de socialización, evaluaciones parciales y finales.

**Informe**  
Genere FD-GC72 a mitad y al final usando listado + plantilla de evaluación. El sistema calcula métricas y deja texto base editable.

**Cierre**  
Valide expediente, cierre estado, descargue respaldo y reporte ejecutivo.

### Reglas duras
- Evaluación = 100%.
- Sesiones con fecha, contenido, trabajo presencial e independiente.
- Aprobación con observaciones abiertas: mala idea, aunque el sistema permita trazarla.
- Producción real exige PostgreSQL/Supabase y secretos configurados en Streamlit Cloud.
        """
    )


def main():
    import streamlit as st
    globals()["st"] = st
    st.set_page_config(page_title="Gestor Académico Next Level FD-GC71 / FD-GC72", layout="wide")
    init_db()

    if "auth_user" not in st.session_state:
        pantalla_login(st)
        return

    user = st.session_state.get("auth_user", {})
    st.title("Gestor académico Next Level FD-GC71 / FD-GC72")
    st.caption("Expediente, planeación, evaluación, informes, evidencias, aprobaciones, versionamiento, motor académico y reportes ejecutivos.")

    with st.sidebar:
        st.markdown(f"**{user.get('nombre_completo', '')}**")
        st.caption(f"Perfil: {user.get('rol', '')}")
        st.caption(f"Versión: {APP_VERSION}")
        if st.button("Cerrar sesión", use_container_width=True):
            registrar_auditoria("Logout", "Cierre de sesión")
            st.session_state.pop("auth_user", None)
            st.rerun()
        st.divider()
        modulos = ROLES_PERMISOS.get(user.get("rol"), {}).get("modulos", ["Inicio", "Mi cuenta", "Ayuda / flujo recomendado"])
        pagina = st.radio("Módulo", modulos)

    if not tiene_permiso(pagina):
        st.error("Este perfil no tiene permisos para abrir este módulo.")
        return

    if pagina == "Inicio":
        ui_inicio(st)
    elif pagina == MODULO_CENTRO:
        ui_centro_control(st)
    elif pagina == MODULO_EXPEDIENTE:
        ui_expediente_academico(st)
    elif pagina == MODULO_PLANEADOR:
        ui_planeador_superior(st)
    elif pagina == MODULO_FLUJO:
        ui_flujo_aprobaciones(st)
    elif pagina == MODULO_MOTOR:
        ui_motor_academico(st)
    elif pagina == MODULO_REPORTES:
        ui_reportes_ejecutivos(st)
    elif pagina == MODULO_PARAMETROS:
        ui_parametros(st)
    elif pagina.startswith("FD-GC71"):
        ui_gc71(st)
    elif pagina.startswith("FD-GC72"):
        ui_gc72(st)
    elif pagina == MODULO_EVIDENCIAS:
        ui_evidencias(st)
    elif pagina == MODULO_VALIDADOR:
        ui_validador_institucional(st)
    elif pagina == MODULO_BACKUP:
        ui_backup(st)
    elif pagina == MODULO_DIAGNOSTICO:
        ui_diagnostico_productivo(st)
    elif pagina == "Usuarios y perfiles":
        ui_admin_usuarios(st)
    elif pagina == "Auditoría":
        ui_auditoria(st)
    elif pagina == "Mi cuenta":
        ui_mi_cuenta(st)
    else:
        ui_ayuda(st)


# =============================================================================
# UX PREMIUM 5.0 - Diseño, navegación y experiencia de usuario
# =============================================================================
APP_VERSION = "5.0.0-ux-premium"

UX_MODULE_META = {
    "Inicio": {"icon": "🏠", "group": "Inicio", "desc": "Resumen operativo, alertas y accesos rápidos.", "step": 1},
    MODULO_CENTRO: {"icon": "📊", "group": "Dirección", "desc": "Indicadores, riesgo y salud académica.", "step": 2},
    MODULO_EXPEDIENTE: {"icon": "📚", "group": "Operación", "desc": "Carpeta viva del curso y trazabilidad.", "step": 3},
    MODULO_PLANEADOR: {"icon": "🧭", "group": "Operación", "desc": "Planeación automática, calendario, FD-GC71 y paquete del curso.", "step": 4},
    MODULO_FLUJO: {"icon": "✅", "group": "Gobierno", "desc": "Revisión, observaciones, aprobación y cierre.", "step": 5},
    MODULO_MOTOR: {"icon": "🧠", "group": "Gobierno", "desc": "Calidad, alertas y redacción académica asistida.", "step": 6},
    MODULO_REPORTES: {"icon": "📈", "group": "Dirección", "desc": "Cortes ejecutivos, filtros y exportaciones.", "step": 7},
    MODULO_PARAMETROS: {"icon": "⚙️", "group": "Administración", "desc": "Reglas, textos base y umbrales institucionales.", "step": 8},
    "FD-GC71 - Planeación": {"icon": "📝", "group": "Formatos", "desc": "Guía didáctica y concertación de evaluación.", "step": 4},
    "FD-GC72 - Informe académico": {"icon": "📄", "group": "Formatos", "desc": "Informe académico con métricas y análisis descriptivo.", "step": 6},
    MODULO_EVIDENCIAS: {"icon": "🗂️", "group": "Operación", "desc": "Soportes, socialización, listados y archivos del curso.", "step": 5},
    MODULO_VALIDADOR: {"icon": "🛡️", "group": "Gobierno", "desc": "Control de completitud y consistencia institucional.", "step": 6},
    MODULO_BACKUP: {"icon": "💾", "group": "Sistema", "desc": "Respaldo, exportación y continuidad operacional.", "step": 8},
    MODULO_DIAGNOSTICO: {"icon": "🩺", "group": "Sistema", "desc": "Estado del despliegue, base de datos y plantillas.", "step": 8},
    "Usuarios y perfiles": {"icon": "👥", "group": "Administración", "desc": "Gestión de usuarios, roles y accesos.", "step": 8},
    "Auditoría": {"icon": "🔎", "group": "Gobierno", "desc": "Registro de acciones críticas y trazabilidad.", "step": 8},
    "Mi cuenta": {"icon": "🙋", "group": "Cuenta", "desc": "Datos personales y cambio de contraseña.", "step": 1},
    "Ayuda / flujo recomendado": {"icon": "❔", "group": "Ayuda", "desc": "Guía operativa del sistema.", "step": 1},
}

UX_GROUP_ORDER = ["Inicio", "Dirección", "Operación", "Formatos", "Gobierno", "Administración", "Sistema", "Cuenta", "Ayuda"]


def ux_meta(modulo: str) -> Dict[str, Any]:
    return UX_MODULE_META.get(modulo, {"icon": "▫️", "group": "Otros", "desc": "Módulo del sistema.", "step": 1})


def ux_label(modulo: str) -> str:
    meta = ux_meta(modulo)
    return f"{meta.get('icon', '▫️')} {modulo}"


def ux_apply_theme(st):
    """Inyecta el sistema visual premium sin depender de librerías externas."""
    st.markdown(
        """
<style>
:root {
    --ux-bg: #f6f8fc;
    --ux-card: #ffffff;
    --ux-ink: #152238;
    --ux-muted: #607089;
    --ux-line: rgba(21,34,56,.11);
    --ux-primary: #1f4fd8;
    --ux-primary-2: #0f2f86;
    --ux-accent: #14b8a6;
    --ux-warning: #f59e0b;
    --ux-danger: #ef4444;
    --ux-success: #16a34a;
    --ux-shadow: 0 12px 32px rgba(15, 31, 72, .08);
    --ux-radius: 22px;
}
html, body, [data-testid="stAppViewContainer"] {
    background:
        radial-gradient(circle at top left, rgba(31,79,216,.10), transparent 35%),
        radial-gradient(circle at top right, rgba(20,184,166,.10), transparent 30%),
        var(--ux-bg) !important;
}
.block-container {
    padding-top: 1.15rem !important;
    padding-bottom: 4rem !important;
    max-width: 1480px !important;
}
[data-testid="stHeader"] { background: rgba(246,248,252,.72) !important; backdrop-filter: blur(10px); }
[data-testid="stSidebar"] > div:first-child {
    background: linear-gradient(180deg, #0f172a 0%, #15254c 52%, #0f172a 100%) !important;
    color: #fff !important;
    border-right: 1px solid rgba(255,255,255,.08);
}
[data-testid="stSidebar"] label, [data-testid="stSidebar"] p, [data-testid="stSidebar"] span,
[data-testid="stSidebar"] div, [data-testid="stSidebar"] small { color: rgba(255,255,255,.88) !important; }
[data-testid="stSidebar"] input, [data-testid="stSidebar"] textarea {
    background: rgba(255,255,255,.10) !important;
    color: #fff !important;
    border: 1px solid rgba(255,255,255,.18) !important;
}
[data-testid="stSidebar"] [data-baseweb="select"] > div {
    background: rgba(255,255,255,.11) !important;
    border-color: rgba(255,255,255,.22) !important;
    border-radius: 14px !important;
}
[data-testid="stMetric"] {
    background: rgba(255,255,255,.88);
    border: 1px solid var(--ux-line);
    padding: 1rem 1.05rem;
    border-radius: 18px;
    box-shadow: 0 10px 28px rgba(15,31,72,.06);
}
[data-testid="stMetricLabel"] { color: var(--ux-muted) !important; }
[data-testid="stMetricValue"] { color: var(--ux-ink) !important; font-weight: 800 !important; }
.stButton > button, .stDownloadButton > button, [data-testid="baseButton-secondary"], [data-testid="baseButton-primary"] {
    border-radius: 14px !important;
    border: 1px solid rgba(31,79,216,.18) !important;
    font-weight: 700 !important;
    min-height: 2.65rem;
    box-shadow: 0 8px 18px rgba(31,79,216,.08);
    transition: all .18s ease-in-out;
}
.stButton > button:hover, .stDownloadButton > button:hover {
    transform: translateY(-1px);
    box-shadow: 0 14px 28px rgba(31,79,216,.14);
}
[data-testid="stExpander"] {
    border-radius: 18px !important;
    border: 1px solid var(--ux-line) !important;
    background: rgba(255,255,255,.82) !important;
    box-shadow: 0 8px 20px rgba(15,31,72,.045);
}
.stTabs [data-baseweb="tab-list"] {
    gap: .45rem;
    background: rgba(255,255,255,.65);
    border: 1px solid var(--ux-line);
    padding: .35rem;
    border-radius: 16px;
}
.stTabs [data-baseweb="tab"] {
    border-radius: 12px !important;
    font-weight: 750;
}
.stTabs [aria-selected="true"] {
    background: linear-gradient(135deg, var(--ux-primary), var(--ux-primary-2)) !important;
    color: white !important;
}
div[data-testid="stDataFrame"], div[data-testid="stTable"] {
    border-radius: 18px !important;
    overflow: hidden !important;
    box-shadow: 0 10px 24px rgba(15,31,72,.05);
}
.ux-hero {
    border-radius: 28px;
    padding: 1.35rem 1.45rem;
    background:
      linear-gradient(135deg, rgba(31,79,216,.96), rgba(15,47,134,.96) 58%, rgba(20,184,166,.88));
    color: #fff;
    box-shadow: 0 20px 48px rgba(15,47,134,.22);
    margin-bottom: 1.05rem;
    border: 1px solid rgba(255,255,255,.18);
}
.ux-hero h1 { margin: 0 0 .35rem 0; font-size: 1.75rem; line-height: 1.15; letter-spacing: -.03em; }
.ux-hero p { margin: 0; color: rgba(255,255,255,.86); font-size: .98rem; }
.ux-chip-row { display:flex; flex-wrap:wrap; gap:.45rem; margin-top:.8rem; }
.ux-chip {
    display:inline-flex; align-items:center; gap:.35rem;
    border-radius: 999px; padding:.34rem .7rem;
    background: rgba(255,255,255,.16); color:#fff;
    border:1px solid rgba(255,255,255,.20); font-size:.82rem; font-weight:750;
}
.ux-card-grid { display:grid; grid-template-columns: repeat(auto-fit, minmax(240px, 1fr)); gap: .9rem; margin: .65rem 0 1.1rem; }
.ux-card {
    background: rgba(255,255,255,.92);
    border: 1px solid var(--ux-line);
    border-radius: var(--ux-radius);
    padding: 1rem;
    box-shadow: var(--ux-shadow);
    min-height: 126px;
}
.ux-card h3 { color: var(--ux-ink); margin:.1rem 0 .35rem 0; font-size:1.05rem; }
.ux-card p { color: var(--ux-muted); margin:0; font-size:.92rem; line-height:1.45; }
.ux-card .ux-card-icon { font-size:1.55rem; line-height:1; margin-bottom:.45rem; }
.ux-kbd {
    padding:.15rem .44rem; border-radius:7px; background:#edf2ff; border:1px solid #dbe5ff; color:#1f4fd8; font-weight:800;
}
.ux-badge {
    display:inline-flex; align-items:center; gap:.3rem; border-radius:999px; padding:.28rem .65rem; font-weight:800; font-size:.78rem;
    background:#edf2ff; color:#1f4fd8; border:1px solid #dbe5ff;
}
.ux-badge.ok { background:#ecfdf5; color:#047857; border-color:#bbf7d0; }
.ux-badge.warn { background:#fffbeb; color:#b45309; border-color:#fde68a; }
.ux-badge.danger { background:#fef2f2; color:#b91c1c; border-color:#fecaca; }
.ux-soft-box {
    background: rgba(255,255,255,.86);
    border: 1px solid var(--ux-line);
    border-radius: 20px;
    padding: 1rem 1.1rem;
    box-shadow: 0 10px 24px rgba(15,31,72,.045);
    margin-bottom: .85rem;
}
.ux-mini-title { font-size:.78rem; text-transform:uppercase; letter-spacing:.08em; color:var(--ux-muted); font-weight:900; margin-bottom:.35rem; }
.ux-path {
    display:flex; flex-wrap:wrap; gap:.45rem; align-items:center; margin: .55rem 0 1rem;
}
.ux-path span {
    background: rgba(255,255,255,.9); border:1px solid var(--ux-line); border-radius:999px; padding:.36rem .64rem; font-size:.82rem; font-weight:750; color:var(--ux-muted);
}
.ux-path span.active { color:#fff; background:linear-gradient(135deg,var(--ux-primary),var(--ux-primary-2)); border-color:rgba(31,79,216,.25); }
.ux-sidebar-brand {
    background: rgba(255,255,255,.10);
    border: 1px solid rgba(255,255,255,.16);
    border-radius: 20px;
    padding: .9rem;
    margin-bottom: .75rem;
}
.ux-sidebar-brand h2 { color:#fff; font-size: 1.05rem; margin:0 0 .25rem 0; }
.ux-sidebar-brand p { margin:0; color:rgba(255,255,255,.70) !important; font-size:.80rem; }
.ux-user-pill {
    border-radius: 16px; padding:.72rem .8rem; background:rgba(255,255,255,.09); border:1px solid rgba(255,255,255,.14); margin:.65rem 0;
}
.ux-user-pill strong { color:#fff; display:block; margin-bottom:.1rem; }
.ux-user-pill small { color:rgba(255,255,255,.70) !important; }
.ux-footer-note { color: var(--ux-muted); font-size: .82rem; margin-top:1.2rem; }
@media (max-width: 760px) {
    .ux-hero { padding: 1.05rem; border-radius: 20px; }
    .ux-hero h1 { font-size: 1.28rem; }
    .ux-card-grid { grid-template-columns: 1fr; }
}
</style>
        """,
        unsafe_allow_html=True,
    )


def ux_badge(label: str, kind: str = "") -> str:
    cls = f"ux-badge {kind}".strip()
    return f"<span class='{cls}'>{label}</span>"


def ux_card(icon: str, title: str, text: str) -> str:
    return f"""
<div class="ux-card">
  <div class="ux-card-icon">{icon}</div>
  <h3>{title}</h3>
  <p>{text}</p>
</div>
"""


def ux_render_hero(st, modulo: str, user: Dict[str, Any]):
    meta = ux_meta(modulo)
    env = get_app_env()
    db_label = "PostgreSQL" if usar_postgres() else "SQLite local"
    badge_db = "Persistencia externa" if usar_postgres() else "Modo local/demo"
    st.markdown(
        f"""
<div class="ux-hero">
  <h1>{meta.get('icon')} {modulo}</h1>
  <p>{meta.get('desc')} · Usuario: <strong>{user.get('nombre_completo') or user.get('usuario','')}</strong> · Perfil: <strong>{user.get('rol','')}</strong></p>
  <div class="ux-chip-row">
    <span class="ux-chip">Versión {APP_VERSION}</span>
    <span class="ux-chip">Ambiente {env}</span>
    <span class="ux-chip">Base {db_label}</span>
    <span class="ux-chip">{badge_db}</span>
  </div>
</div>
        """,
        unsafe_allow_html=True,
    )


def ux_render_path(st, active_module: str):
    steps = [
        ("Crear expediente", MODULO_EXPEDIENTE),
        ("Planear FD-GC71", MODULO_PLANEADOR),
        ("Cargar evidencias", MODULO_EVIDENCIAS),
        ("Revisar/aprobar", MODULO_FLUJO),
        ("Generar FD-GC72", "FD-GC72 - Informe académico"),
        ("Cerrar y reportar", MODULO_REPORTES),
    ]
    html = '<div class="ux-path">'
    for label, module in steps:
        cls = "active" if active_module == module or (active_module.startswith("FD-GC71") and module == MODULO_PLANEADOR) else ""
        html += f"<span class='{cls}'>{label}</span>"
    html += "</div>"
    st.markdown(html, unsafe_allow_html=True)


def ux_safe_int(value: Any, default: int = 0) -> int:
    try:
        return int(value or default)
    except Exception:
        return default


def ux_system_counts() -> Dict[str, Any]:
    counts = {"cursos": 0, "alto": 0, "observados": 0, "abiertas": 0, "score": 0.0, "evidencias": 0}
    try:
        df = matriz_riesgo_cursos()
        if not df.empty:
            counts["cursos"] = len(df)
            if "Riesgo" in df.columns:
                counts["alto"] = int((df["Riesgo"].astype(str) == "Alto").sum())
            if "estado" in df.columns:
                counts["observados"] = int((df["estado"].astype(str) == "Observado").sum())
            if "Score calidad" in df.columns:
                counts["score"] = float(pd.to_numeric(df["Score calidad"], errors="coerce").fillna(0).mean())
            if "Evidencias" in df.columns:
                counts["evidencias"] = int(pd.to_numeric(df["Evidencias"], errors="coerce").fillna(0).sum())
    except Exception:
        pass
    try:
        row = db_execute("SELECT COUNT(*) AS n FROM observaciones_curso WHERE estado='Abierta'", fetchone=True)
        counts["abiertas"] = ux_safe_int(_dict_row(row).get("n", 0))
    except Exception:
        pass
    return counts


def ux_quick_nav(module_name: str):
    st.session_state["ux_pagina"] = module_name
    st.rerun()


def ux_sidebar(st, user: Dict[str, Any]) -> str:
    modulos = ROLES_PERMISOS.get(user.get("rol"), {}).get("modulos", ["Inicio", "Mi cuenta", "Ayuda / flujo recomendado"])
    if "ux_pagina" not in st.session_state or st.session_state["ux_pagina"] not in modulos:
        st.session_state["ux_pagina"] = modulos[0]

    with st.sidebar:
        st.markdown(
            """
<div class="ux-sidebar-brand">
  <h2>🎓 Gestor Académico</h2>
  <p>FD-GC71 · FD-GC72 · expediente académico</p>
</div>
            """,
            unsafe_allow_html=True,
        )
        st.markdown(
            f"""
<div class="ux-user-pill">
  <strong>{user.get('nombre_completo') or user.get('usuario','')}</strong>
  <small>{user.get('rol','')} · {user.get('usuario','')}</small>
</div>
            """,
            unsafe_allow_html=True,
        )

        buscar = st.text_input("Buscar módulo", placeholder="Ej. planeador, reportes, auditoría", key="ux_buscar_modulo")
        filtrados = []
        q = str(buscar or "").lower().strip()
        for m in modulos:
            meta = ux_meta(m)
            text = f"{m} {meta.get('desc','')} {meta.get('group','')}".lower()
            if not q or q in text:
                filtrados.append(m)
        if not filtrados:
            st.warning("No hay módulos con ese filtro.")
            filtrados = modulos

        # Ordena por grupo y paso para que el menú cuente una historia, no una lista de mercado.
        def sort_key(m: str):
            meta = ux_meta(m)
            group_idx = UX_GROUP_ORDER.index(meta.get("group", "Ayuda")) if meta.get("group", "Ayuda") in UX_GROUP_ORDER else 99
            return (group_idx, int(meta.get("step", 99)), m)
        filtrados = sorted(filtrados, key=sort_key)
        current = st.session_state.get("ux_pagina", filtrados[0])
        if current not in filtrados:
            current = filtrados[0]
        pagina = st.selectbox(
            "Módulo activo",
            filtrados,
            index=filtrados.index(current),
            format_func=ux_label,
            key="ux_pagina_selectbox",
        )
        st.session_state["ux_pagina"] = pagina

        meta = ux_meta(pagina)
        st.caption(meta.get("desc", ""))
        progreso = min(max(float(meta.get("step", 1)) / 8.0, 0.05), 1.0)
        st.progress(progreso, text=f"Etapa operativa {meta.get('step', 1)} de 8")
        st.divider()
        counts = ux_system_counts()
        c1, c2 = st.columns(2)
        c1.metric("Cursos", counts["cursos"])
        c2.metric("Riesgo alto", counts["alto"])
        st.caption(f"Obs. abiertas: {counts['abiertas']} · Evidencias: {counts['evidencias']}")
        st.divider()
        if st.button("🚪 Cerrar sesión", use_container_width=True):
            registrar_auditoria("Logout", "Cierre de sesión")
            st.session_state.pop("auth_user", None)
            st.session_state.pop("ux_pagina", None)
            st.rerun()
    return pagina


# Conservar referencias a módulos clásicos y envolverlos con cabecera UX.
_base_ui_gc71_ux = ui_gc71
_base_ui_gc72_ux = ui_gc72
_base_ui_expediente_ux = ui_expediente_academico
_base_ui_planeador_ux = ui_planeador_superior
_base_ui_evidencias_ux = ui_evidencias
_base_ui_validador_ux = ui_validador_institucional
_base_ui_backup_ux = ui_backup
_base_ui_diagnostico_ux = ui_diagnostico_productivo
_base_ui_admin_usuarios_ux = ui_admin_usuarios
_base_ui_auditoria_ux = ui_auditoria
_base_ui_mi_cuenta_ux = ui_mi_cuenta
_base_ui_parametros_ux = ui_parametros


def pantalla_login(st):
    ux_apply_theme(st)
    c1, c2, c3 = st.columns([.9, 1.25, .9])
    with c2:
        st.markdown(
            """
<div class="ux-hero" style="margin-top:2.2rem; text-align:left;">
  <h1>🎓 Gestor Académico</h1>
  <p>Planeación, evaluación, expediente, aprobación y cierre institucional de cursos.</p>
  <div class="ux-chip-row">
    <span class="ux-chip">FD-GC71</span>
    <span class="ux-chip">FD-GC72</span>
    <span class="ux-chip">Versionamiento</span>
    <span class="ux-chip">Auditoría</span>
  </div>
</div>
            """,
            unsafe_allow_html=True,
        )
        with st.container(border=True):
            st.subheader("Ingreso seguro")
            st.caption("Use su usuario institucional. La sesión queda auditada.")
            with st.form("login_form", clear_on_submit=False):
                usuario = st.text_input("Usuario", value="admin", placeholder="usuario")
                password = st.text_input("Contraseña", type="password", placeholder="••••••••")
                entrar = st.form_submit_button("Entrar al sistema", use_container_width=True)
            if entrar:
                user = autenticar_usuario(usuario, password)
                if user and not user.get("error"):
                    st.session_state["auth_user"] = user
                    registrar_auditoria("Login", "Ingreso correcto")
                    st.rerun()
                elif user and user.get("error"):
                    st.error(user["error"])
                else:
                    st.error("Usuario o contraseña incorrectos.")
        with st.expander("Primera instalación / credenciales iniciales"):
            if get_app_env() == "production":
                st.info("Las credenciales iniciales se leen desde los secretos INITIAL_ADMIN_USER e INITIAL_ADMIN_PASSWORD.")
            else:
                admin_user, _, _, _ = initial_admin_config()
                st.write(f"Usuario inicial: `{admin_user}`")
                st.write("Contraseña inicial: definida en secrets o, en local, `Admin123*`.")
            st.warning("Cambie la contraseña inicial apenas ingrese. Una contraseña por defecto es una alfombra roja para problemas.")


def ui_inicio(st):
    user = st.session_state.get("auth_user", {})
    status = health_status()
    counts = ux_system_counts()

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Cursos", counts["cursos"])
    c2.metric("Riesgo alto", counts["alto"])
    c3.metric("Obs. abiertas", counts["abiertas"])
    c4.metric("Score prom.", f"{counts['score']:.1f}")

    if get_app_env() == "production" and not usar_postgres():
        st.error("APP_ENV está en production pero no hay DATABASE_URL. Configure PostgreSQL/Supabase antes de operar con datos reales.")
    elif usar_postgres():
        st.success("Persistencia productiva activa: PostgreSQL externo configurado.")
    else:
        st.info("Modo local/demo activo. Ideal para pruebas; producción real debe ir con PostgreSQL/Supabase.")

    if user.get("debe_cambiar_clave"):
        st.warning("Debe cambiar la contraseña inicial desde Mi cuenta antes de operar formalmente.")

    st.markdown("<div class='ux-mini-title'>Accesos rápidos</div>", unsafe_allow_html=True)
    q1, q2, q3, q4 = st.columns(4)
    with q1:
        if st.button("🧭 Crear / planear curso", use_container_width=True):
            ux_quick_nav(MODULO_PLANEADOR)
    with q2:
        if st.button("✅ Revisar expediente", use_container_width=True):
            ux_quick_nav(MODULO_FLUJO)
    with q3:
        if st.button("📄 Generar informe", use_container_width=True):
            ux_quick_nav("FD-GC72 - Informe académico")
    with q4:
        if st.button("📈 Ver reportes", use_container_width=True):
            ux_quick_nav(MODULO_REPORTES)

    st.markdown(
        "<div class='ux-card-grid'>" +
        ux_card("📚", "Expediente vivo", "Cada curso conserva planeación, evidencias, evaluaciones, observaciones, versiones y cierre.") +
        ux_card("🧭", "Planeación guiada", "El horario alimenta el cronograma; las unidades alimentan la plantilla de evaluación y el FD-GC71.") +
        ux_card("🛡️", "Control institucional", "Estados, hash documental, auditoría, observaciones y validación antes de aprobar o cerrar.") +
        ux_card("📊", "Decisión ejecutiva", "Riesgo, avance, cursos observados, evidencias y reportes listos para coordinación.") +
        "</div>",
        unsafe_allow_html=True,
    )

    tab_flujo, tab_alertas, tab_estado = st.tabs(["Ruta recomendada", "Alertas", "Estado técnico"])
    with tab_flujo:
        st.markdown(
            """
<div class="ux-soft-box">
  <div class="ux-mini-title">Flujo operativo</div>
  <p><span class="ux-kbd">1</span> Cree el expediente del curso. &nbsp;
  <span class="ux-kbd">2</span> Planee FD-GC71, horario, unidades y evaluación. &nbsp;
  <span class="ux-kbd">3</span> Cargue listado y evidencias. &nbsp;
  <span class="ux-kbd">4</span> Envíe a revisión y resuelva observaciones. &nbsp;
  <span class="ux-kbd">5</span> Genere FD-GC72 y cierre el expediente.</p>
</div>
            """,
            unsafe_allow_html=True,
        )
    with tab_alertas:
        try:
            tareas = read_sql_df("SELECT id, tipo, mensaje, fecha_limite, estado, creado_por FROM tareas_alertas WHERE estado='Pendiente' ORDER BY fecha_limite ASC LIMIT 20")
            if tareas.empty:
                st.success("No hay alertas pendientes. Raro, pero hermoso.")
            else:
                st.dataframe(tareas, use_container_width=True, hide_index=True)
        except Exception as exc:
            st.info(f"No fue posible leer alertas: {exc}")
    with tab_estado:
        st.json({
            "version": APP_VERSION,
            "ambiente": status.get("app_env"),
            "base": status.get("database"),
            "database_ok": status.get("database_ok"),
            "cloud_ready_level": status.get("cloud_ready_level"),
        })


def ui_centro_control(st):
    df = matriz_riesgo_cursos()
    if df.empty:
        st.info("No hay cursos visibles todavía. Cree un expediente desde Planeador superior.")
        st.markdown(
            "<div class='ux-card-grid'>" +
            ux_card("🧭", "Primer paso", "Abra Planeador superior y cree el expediente inicial del curso.") +
            ux_card("📥", "Datos mínimos", "Programa, asignatura, grupo, docente, periodo y fechas base.") +
            ux_card("📦", "Salida esperada", "FD-GC71, plantilla Excel, calendario y paquete ZIP del curso.") +
            "</div>",
            unsafe_allow_html=True,
        )
        if st.button("Ir al Planeador superior", use_container_width=True):
            ux_quick_nav(MODULO_PLANEADOR)
        return

    score_mean = pd.to_numeric(df.get("Score calidad", pd.Series(dtype=float)), errors="coerce").mean()
    c1, c2, c3, c4, c5 = st.columns(5)
    c1.metric("Cursos", len(df))
    c2.metric("Riesgo alto", int((df.get("Riesgo", pd.Series(dtype=str)).astype(str) == "Alto").sum()))
    c3.metric("Observados", int((df.get("estado", pd.Series(dtype=str)).astype(str) == "Observado").sum()))
    c4.metric("Aprobados/cerrados", int(df.get("estado", pd.Series(dtype=str)).astype(str).isin(["Aprobado", "Cerrado"]).sum()))
    c5.metric("Score promedio", f"{score_mean:.1f}" if not math.isnan(score_mean) else "0.0")

    tab_mapa, tab_graficos, tab_acciones = st.tabs(["Mapa de riesgo", "Gráficos", "Acciones sugeridas"])
    with tab_mapa:
        col_filter_1, col_filter_2, col_filter_3 = st.columns(3)
        with col_filter_1:
            estado_sel = st.selectbox("Estado", ["Todos"] + sorted(df.get("estado", pd.Series(dtype=str)).dropna().astype(str).unique().tolist()), key="ux_cc_estado")
        with col_filter_2:
            riesgo_sel = st.selectbox("Riesgo", ["Todos"] + sorted(df.get("Riesgo", pd.Series(dtype=str)).dropna().astype(str).unique().tolist()), key="ux_cc_riesgo")
        with col_filter_3:
            texto = st.text_input("Buscar curso", placeholder="Asignatura, docente, grupo...", key="ux_cc_buscar")
        f = df.copy()
        if estado_sel != "Todos" and "estado" in f.columns:
            f = f[f["estado"].astype(str) == estado_sel]
        if riesgo_sel != "Todos" and "Riesgo" in f.columns:
            f = f[f["Riesgo"].astype(str) == riesgo_sel]
        if texto:
            mask = pd.Series(False, index=f.index)
            for col in ["codigo", "grupo", "asignatura", "programa", "periodo", "profesor", "Razones"]:
                if col in f.columns:
                    mask = mask | f[col].astype(str).str.contains(texto, case=False, na=False)
            f = f[mask]
        columnas = [c for c in ["id", "codigo", "grupo", "asignatura", "programa", "periodo", "profesor", "estado", "Riesgo", "Puntaje riesgo", "Score calidad", "Evidencias", "Razones"] if c in f.columns]
        st.dataframe(f[columnas], use_container_width=True, hide_index=True)
    with tab_graficos:
        col1, col2 = st.columns(2)
        with col1:
            estado_counts = df["estado"].fillna("Sin estado").value_counts().rename_axis("Estado").reset_index(name="Cursos") if "estado" in df.columns else pd.DataFrame()
            if not estado_counts.empty:
                st.bar_chart(estado_counts.set_index("Estado"))
        with col2:
            riesgo_counts = df["Riesgo"].fillna("Sin riesgo").value_counts().rename_axis("Riesgo").reset_index(name="Cursos") if "Riesgo" in df.columns else pd.DataFrame()
            if not riesgo_counts.empty:
                st.bar_chart(riesgo_counts.set_index("Riesgo"))
    with tab_acciones:
        acciones = []
        if "Riesgo" in df.columns:
            alto = df[df["Riesgo"].astype(str) == "Alto"]
            for _, r in alto.head(10).iterrows():
                acciones.append({"Prioridad": "Alta", "Curso": r.get("asignatura", ""), "Acción": "Revisar hallazgos y evidencias antes de aprobación", "Razón": r.get("Razones", "")})
        if "estado" in df.columns:
            obs = df[df["estado"].astype(str) == "Observado"]
            for _, r in obs.head(10).iterrows():
                acciones.append({"Prioridad": "Media", "Curso": r.get("asignatura", ""), "Acción": "Resolver observaciones abiertas", "Razón": r.get("Razones", "")})
        if acciones:
            st.dataframe(pd.DataFrame(acciones), use_container_width=True, hide_index=True)
        else:
            st.success("No se detectan acciones críticas en el corte actual.")


def ui_reportes_ejecutivos(st):
    df = matriz_riesgo_cursos()
    if df.empty:
        st.info("No hay información para reportar.")
        return
    with st.expander("Filtros del reporte", expanded=True):
        c1, c2, c3, c4 = st.columns(4)
        with c1:
            periodos = ["Todos"] + sorted([str(x) for x in df.get("periodo", pd.Series(dtype=str)).dropna().unique()])
            periodo_sel = st.selectbox("Periodo", periodos)
        with c2:
            estados = ["Todos"] + sorted([str(x) for x in df.get("estado", pd.Series(dtype=str)).dropna().unique()])
            estado_sel = st.selectbox("Estado", estados)
        with c3:
            riesgos = ["Todos"] + sorted([str(x) for x in df.get("Riesgo", pd.Series(dtype=str)).dropna().unique()])
            riesgo_sel = st.selectbox("Riesgo", riesgos)
        with c4:
            buscar = st.text_input("Buscar", placeholder="curso, docente, programa")
    f = df.copy()
    if periodo_sel != "Todos":
        f = f[f["periodo"].astype(str) == periodo_sel]
    if estado_sel != "Todos":
        f = f[f["estado"].astype(str) == estado_sel]
    if riesgo_sel != "Todos":
        f = f[f["Riesgo"].astype(str) == riesgo_sel]
    if buscar:
        mask = pd.Series(False, index=f.index)
        for col in ["codigo", "grupo", "asignatura", "programa", "periodo", "profesor", "Razones"]:
            if col in f.columns:
                mask = mask | f[col].astype(str).str.contains(buscar, case=False, na=False)
        f = f[mask]

    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Cursos filtrados", len(f))
    c2.metric("Riesgo alto", int((f.get("Riesgo", pd.Series(dtype=str)).astype(str) == "Alto").sum()))
    c3.metric("Aprobados/cerrados", int(f.get("estado", pd.Series(dtype=str)).astype(str).isin(["Aprobado", "Cerrado"]).sum()))
    c4.metric("Score prom.", f"{pd.to_numeric(f.get('Score calidad', pd.Series(dtype=float)), errors='coerce').mean():.1f}" if not f.empty else "0.0")

    st.dataframe(f, use_container_width=True, hide_index=True)
    col1, col2 = st.columns(2)
    with col1:
        csv = f.to_csv(index=False).encode("utf-8-sig")
        st.download_button("⬇️ Descargar CSV filtrado", csv, "reporte_ejecutivo_fdgc.csv", "text/csv", use_container_width=True)
    with col2:
        xlsx = crear_reporte_ejecutivo_xlsx()
        st.download_button("⬇️ Descargar Excel completo", xlsx, "reporte_ejecutivo_fdgc.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)


def ui_ayuda(st):
    st.markdown(
        "<div class='ux-card-grid'>" +
        ux_card("🧭", "Planeación", "Cree expediente, defina módulos, intensidad, horarios, contenidos y evaluación concertada.") +
        ux_card("✅", "Revisión", "Envíe a revisión, reciba observaciones, responda ajustes y conserve versiones con hash.") +
        ux_card("🗂️", "Ejecución", "Cargue evidencias, listados, notas parciales/finales y soportes de socialización.") +
        ux_card("📄", "Informe", "Genere FD-GC72 a mitad y cierre del curso con análisis editable y métricas calculadas.") +
        "</div>",
        unsafe_allow_html=True,
    )
    with st.expander("Reglas de oro", expanded=True):
        st.markdown(
            """
- La evaluación concertada debe sumar **100%**.
- Ningún curso debería cerrarse sin evidencias mínimas.
- Las observaciones deben resolverse antes de aprobar o cerrar.
- En producción real use **PostgreSQL/Supabase**, no SQLite local.
- El respaldo del expediente es parte del proceso, no decoración de fin de año.
            """
        )
    with st.expander("Atajos recomendados"):
        st.write("Use la barra lateral para buscar módulos. Ejemplos: `planeador`, `reportes`, `auditoría`, `evidencias`.")


def ui_gc71(st):
    with st.expander("Qué hace este módulo", expanded=False):
        st.markdown("Genera la guía didáctica FD-GC71, cronograma automático según horario, evaluación concertada y plantilla Excel de evaluación.")
    _base_ui_gc71_ux(st)


def ui_gc72(st):
    with st.expander("Qué hace este módulo", expanded=False):
        st.markdown("Genera el informe académico FD-GC72 usando listado tradicional, notas de corte y análisis descriptivo editable.")
    _base_ui_gc72_ux(st)


def ui_expediente_academico(st):
    with st.expander("Experiencia recomendada", expanded=False):
        st.markdown("Seleccione el curso, revise su estado, evidencias, observaciones y trazabilidad. El expediente es la fuente única de verdad.")
    _base_ui_expediente_ux(st)


def ui_planeador_superior(st):
    with st.expander("Experiencia recomendada", expanded=False):
        st.markdown("Use este módulo como punto de partida: crea curso, define datos base y descarga el paquete operativo completo.")
    _base_ui_planeador_ux(st)


def ui_evidencias(st):
    with st.expander("Experiencia recomendada", expanded=False):
        st.markdown("Cargue soportes por curso: socialización, listados, parciales, finales, actas y archivos complementarios.")
    _base_ui_evidencias_ux(st)


def ui_validador_institucional(st):
    with st.expander("Experiencia recomendada", expanded=False):
        st.markdown("Ejecute validaciones antes de enviar a revisión, aprobar o cerrar. Aquí se detectan los huecos feos.")
    _base_ui_validador_ux(st)


def ui_backup(st):
    with st.expander("Experiencia recomendada", expanded=False):
        st.markdown("Descargue respaldos periódicos, especialmente antes de cierres de periodo o migraciones.")
    _base_ui_backup_ux(st)


def ui_diagnostico_productivo(st):
    _base_ui_diagnostico_ux(st)


def ui_admin_usuarios(st):
    with st.expander("Experiencia recomendada", expanded=False):
        st.markdown("Cree usuarios nominales y perfiles mínimos necesarios. Evite cuentas genéricas para operación real.")
    _base_ui_admin_usuarios_ux(st)


def ui_auditoria(st):
    with st.expander("Experiencia recomendada", expanded=False):
        st.markdown("Revise acciones críticas, ingresos, cambios de estado, versiones y movimientos administrativos.")
    _base_ui_auditoria_ux(st)


def ui_mi_cuenta(st):
    _base_ui_mi_cuenta_ux(st)


def ui_parametros(st):
    _base_ui_parametros_ux(st)


def main():
    import streamlit as st
    globals()["st"] = st
    st.set_page_config(
        page_title="Gestor Académico UX Premium FD-GC71 / FD-GC72",
        page_icon="🎓",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    ux_apply_theme(st)
    try:
        init_db()
    except Exception as exc:
        st.error("La aplicación inició, pero no pudo preparar la base de datos.")
        st.markdown("""
        Esto suele pasar en Streamlit Cloud cuando `DATABASE_URL` está mal escrita,
        contiene el ejemplo `USUARIO:CLAVE@HOST:5432/BASE`, la base de datos no acepta conexiones externas
        o el proveedor requiere usar el pooler/puerto correcto.
        """)
        st.code(str(exc), language="text")
        st.info("Corrija los Secrets en Streamlit Cloud y reinicie la app. Para probar sin base externa, quite `DATABASE_URL` y use modo local/demo.")
        st.stop()

    if "auth_user" not in st.session_state:
        pantalla_login(st)
        return

    user = st.session_state.get("auth_user", {})
    pagina = ux_sidebar(st, user)

    if not tiene_permiso(pagina):
        st.error("Este perfil no tiene permisos para abrir este módulo.")
        return

    ux_render_hero(st, pagina, user)
    if pagina not in ["Inicio", "Mi cuenta", "Ayuda / flujo recomendado", MODULO_DIAGNOSTICO]:
        ux_render_path(st, pagina)

    if pagina == "Inicio":
        ui_inicio(st)
    elif pagina == MODULO_CENTRO:
        ui_centro_control(st)
    elif pagina == MODULO_EXPEDIENTE:
        ui_expediente_academico(st)
    elif pagina == MODULO_PLANEADOR:
        ui_planeador_superior(st)
    elif pagina == MODULO_FLUJO:
        ui_flujo_aprobaciones(st)
    elif pagina == MODULO_MOTOR:
        ui_motor_academico(st)
    elif pagina == MODULO_REPORTES:
        ui_reportes_ejecutivos(st)
    elif pagina == MODULO_PARAMETROS:
        ui_parametros(st)
    elif pagina.startswith("FD-GC71"):
        ui_gc71(st)
    elif pagina.startswith("FD-GC72"):
        ui_gc72(st)
    elif pagina == MODULO_EVIDENCIAS:
        ui_evidencias(st)
    elif pagina == MODULO_VALIDADOR:
        ui_validador_institucional(st)
    elif pagina == MODULO_BACKUP:
        ui_backup(st)
    elif pagina == MODULO_DIAGNOSTICO:
        ui_diagnostico_productivo(st)
    elif pagina == "Usuarios y perfiles":
        ui_admin_usuarios(st)
    elif pagina == "Auditoría":
        ui_auditoria(st)
    elif pagina == "Mi cuenta":
        ui_mi_cuenta(st)
    else:
        ui_ayuda(st)

    st.markdown("<div class='ux-footer-note'>Gestor Académico UX Premium · flujo: planeación → concertación → ejecución → informe → cierre.</div>", unsafe_allow_html=True)



# =============================================================================
# SUITE INTELIGENTE V6: Excel inteligente, comparación de cortes, semáforos,
# exportación masiva, hash/QR y redacción asistida.
# =============================================================================
APP_VERSION = "6.0.0-suite-inteligente"
MODULO_CARGADOR = "Cargador inteligente de Excel"
MODULO_COMPARADOR = "Comparador de cortes"
MODULO_SEMAFORO = "Semáforo del expediente"
MODULO_EXPORTACION = "Exportación masiva"
MODULO_VERIFICACION = "Hash y QR documental"
MODULO_ASISTENTE = "Asistente académico"

# Ajusta la taxonomía visual sin tocar el núcleo anterior.
try:
    UX_MODULE_META.update({
        MODULO_CARGADOR: {"icon": "🧩", "group": "Automatización", "desc": "Normaliza listados y notas aunque el Excel venga creativo.", "step": 5},
        MODULO_COMPARADOR: {"icon": "🔁", "group": "Automatización", "desc": "Compara listado inicial, corte parcial y corte final para alimentar FD-GC72.", "step": 6},
        MODULO_SEMAFORO: {"icon": "🚦", "group": "Dirección", "desc": "Estado visual de completitud, riesgo, evidencias y observaciones por curso.", "step": 6},
        MODULO_EXPORTACION: {"icon": "📦", "group": "Dirección", "desc": "Descarga expedientes, paquetes e índices por período, programa o estado.", "step": 8},
        MODULO_VERIFICACION: {"icon": "🔐", "group": "Gobierno", "desc": "Genera huella SHA-256 y QR de verificación para expedientes y documentos.", "step": 8},
        MODULO_ASISTENTE: {"icon": "✍️", "group": "Automatización", "desc": "Redacción asistida editable para FD-GC71 y FD-GC72.", "step": 4},
    })
    if "Automatización" not in UX_GROUP_ORDER:
        UX_GROUP_ORDER.insert(3, "Automatización")
except Exception:
    pass

# Expande permisos por perfil. La función tiene_permiso lee ROLES_PERMISOS en tiempo de ejecución.
def _add_modulos_a_rol(rol: str, nuevos: List[str]):
    if rol not in ROLES_PERMISOS:
        return
    actuales = ROLES_PERMISOS[rol].setdefault("modulos", [])
    # Inserta después del planeador cuando sea posible para mantener historia operacional.
    insert_at = len(actuales)
    for anchor in [MODULO_PLANEADOR, MODULO_CENTRO, "Inicio"]:
        if anchor in actuales:
            insert_at = actuales.index(anchor) + 1
            break
    for m in reversed(nuevos):
        if m not in actuales:
            actuales.insert(insert_at, m)

for _rol in ["Administrador", "Coordinador"]:
    _add_modulos_a_rol(_rol, [MODULO_CARGADOR, MODULO_COMPARADOR, MODULO_SEMAFORO, MODULO_EXPORTACION, MODULO_VERIFICACION, MODULO_ASISTENTE])
for _rol in ["Docente"]:
    _add_modulos_a_rol(_rol, [MODULO_CARGADOR, MODULO_COMPARADOR, MODULO_SEMAFORO, MODULO_VERIFICACION, MODULO_ASISTENTE])
for _rol in ["Consulta"]:
    _add_modulos_a_rol(_rol, [MODULO_SEMAFORO, MODULO_EXPORTACION])


def _file_bytes(uploaded_file) -> bytes:
    if uploaded_file is None:
        return b""
    try:
        return uploaded_file.getvalue()
    except Exception:
        pos = uploaded_file.tell()
        data = uploaded_file.read()
        try:
            uploaded_file.seek(pos)
        except Exception:
            pass
        return data


def _unique_columns(cols: Iterable[Any]) -> List[str]:
    seen: Dict[str, int] = {}
    out: List[str] = []
    for i, c in enumerate(cols):
        name = str(c).strip()
        if not name or name.lower() in {"nan", "none", "unnamed"}:
            name = f"Columna_{i+1}"
        base = name
        k = seen.get(base, 0)
        if k:
            name = f"{base}_{k+1}"
        seen[base] = k + 1
        out.append(name)
    return out


def _smart_header_row(raw: pd.DataFrame) -> int:
    if raw is None or raw.empty:
        return 0
    keywords = [
        "DOCUMENTO", "CEDULA", "CÉDULA", "IDENTIFIC", "CARN", "CODIGO", "CÓDIGO",
        "NOMBRE", "APELLIDO", "ESTUDIANTE", "CORREO", "EMAIL", "NOTA", "CALIFIC",
        "ESTADO", "OBSERV", "GRUPO", "PROGRAMA", "ASIGNATURA"
    ]
    best_idx, best_score = 0, -1
    max_rows = min(len(raw), 25)
    for idx in range(max_rows):
        vals = [normalizar_texto(v) for v in raw.iloc[idx].tolist() if not pd.isna(v)]
        text = " | ".join(vals)
        non_empty = sum(1 for v in vals if str(v).strip())
        score = sum(1 for k in keywords if normalizar_texto(k) in text) + min(non_empty, 8) * 0.08
        # Castiga filas puramente numéricas.
        numeric_like = sum(1 for v in vals if re.fullmatch(r"[0-9., -]+", str(v).strip() or ""))
        score -= numeric_like * 0.05
        if score > best_score:
            best_idx, best_score = idx, score
    return int(best_idx)


def leer_excel_inteligente(uploaded_file) -> Dict[str, pd.DataFrame]:
    """Lee xls/xlsx/csv detectando fila de encabezado y limpiando columnas."""
    if uploaded_file is None:
        return {}
    name = getattr(uploaded_file, "name", "archivo").lower()
    data = _file_bytes(uploaded_file)
    if not data:
        return {}
    bio = io.BytesIO(data)
    tablas: Dict[str, pd.DataFrame] = {}
    if name.endswith(".csv"):
        raw = pd.read_csv(bio, header=None, dtype=object)
        header = _smart_header_row(raw)
        df = raw.iloc[header + 1:].copy()
        df.columns = _unique_columns(raw.iloc[header].tolist())
        df = df.dropna(how="all").reset_index(drop=True)
        tablas["CSV"] = df
        return tablas
    xls = pd.ExcelFile(bio)
    for sheet in xls.sheet_names:
        raw = pd.read_excel(xls, sheet_name=sheet, header=None, dtype=object)
        if raw.empty:
            continue
        header = _smart_header_row(raw)
        df = raw.iloc[header + 1:].copy()
        df.columns = _unique_columns(raw.iloc[header].tolist())
        df = df.dropna(how="all").reset_index(drop=True)
        # Elimina columnas totalmente vacías o con nombres basura repetidos.
        df = df.loc[:, [c for c in df.columns if not df[c].isna().all()]]
        if not df.empty:
            tablas[str(sheet)] = df
    return tablas


ROLE_PATTERNS: Dict[str, List[str]] = {
    "documento": ["DOCUMENTO", "CEDULA", "CÉDULA", "IDENTIFIC", "ID", "CARN", "CODIGO", "CÓDIGO"],
    "nombre": ["NOMBRE", "ESTUDIANTE", "ALUMNO", "APELLIDO", "NOMBRES"],
    "correo": ["CORREO", "EMAIL", "MAIL", "E-MAIL"],
    "programa": ["PROGRAMA", "PLAN", "CARRERA"],
    "grupo": ["GRUPO", "CURSO"],
    "estado": ["ESTADO", "SITUACION", "SITUACIÓN", "OBSERV", "CONDIC"],
    "nota_parcial": ["PARCIAL", "CORTE", "SEGUIMIENTO", "NOTA PAR", "PROMEDIO PAR"],
    "nota_final": ["FINAL", "DEFINITIVA", "NOTA FINAL", "PROMEDIO FINAL", "CALIFICACION FINAL", "CALIFICACIÓN FINAL"],
    "nota": ["NOTA", "CALIFIC", "DEFINITIVA", "PROMEDIO", "TOTAL"],
}


def _guess_column(df: pd.DataFrame, role: str) -> Optional[str]:
    if df is None or df.empty:
        return None
    patterns = [normalizar_texto(p) for p in ROLE_PATTERNS.get(role, [])]
    best_col, best_score = None, -1
    for col in df.columns:
        n = normalizar_texto(col)
        score = 0
        for p in patterns:
            if p == n:
                score += 4
            elif p and p in n:
                score += 2
        sample = " ".join(normalizar_texto(v) for v in df[col].head(20).tolist() if not pd.isna(v))
        if role == "correo" and "@" in sample:
            score += 3
        if role in {"nota", "nota_parcial", "nota_final"}:
            # No todo número es una nota: documentos, códigos y carnés también son numéricos.
            if re.search(r"DOCUMENTO|CEDULA|CÉDULA|IDENTIFIC|CARN|CODIGO|CÓDIGO|^ID$", n):
                score -= 4
            else:
                nums = pd.to_numeric(df[col].astype(str).str.replace(",", ".", regex=False), errors="coerce")
                ratio = float(nums.notna().mean()) if len(nums) else 0
                header_signal = any(p and p in n for p in patterns)
                if ratio > 0.45 and header_signal:
                    score += 2
                elif ratio > 0.60 and role == "nota":
                    score += 1
        if role == "estado" and re.search(r"ACTIVO|DESERT|RETI|CANCEL|MATRIC", sample):
            score += 2
        if score > best_score:
            best_col, best_score = col, score
    return best_col if best_score > 0 else None


def _mapping_automatico(df: pd.DataFrame) -> Dict[str, Optional[str]]:
    return {role: _guess_column(df, role) for role in ROLE_PATTERNS.keys()}


def _clean_document_key(value: Any) -> str:
    if pd.isna(value):
        return ""
    s = str(value).strip().replace(".0", "")
    s = re.sub(r"\s+", "", s)
    s = re.sub(r"[^A-Za-z0-9]", "", s)
    return s.upper()


def _clean_text(value: Any) -> str:
    if pd.isna(value):
        return ""
    return str(value).strip()


def _coerce_grade(value: Any) -> Optional[float]:
    if value is None or pd.isna(value):
        return None
    s = str(value).strip().replace("%", "").replace(",", ".")
    # Evita convertir documentos largos por accidente.
    try:
        v = float(s)
    except Exception:
        return None
    # Si viene en escala 0-100, la lleva a 0-5 solo cuando es claramente porcentaje.
    # Valores superiores a 100 suelen ser documentos/códigos mal mapeados, no notas.
    if v > 100:
        return None
    if v > 5 and v <= 100:
        v = v / 20.0
    if v < 0 or v > 5:
        return None
    return round(v, 2)


def _estado_normalizado(value: Any) -> str:
    t = normalizar_texto(value)
    if re.search(r"DESERT|RETI|CANCEL|ANUL|INACT|ABAND", t):
        return "Desertó"
    if re.search(r"APLAZ|SUSP", t):
        return "Suspendido"
    if re.search(r"MATRIC|ACTIV|INSCR|REGULAR|CURS", t):
        return "Activo"
    return _clean_text(value) or "Activo"


def normalizar_estudiantes_inteligente(df: pd.DataFrame, mapping: Dict[str, Optional[str]], nota_minima: float = 3.0) -> pd.DataFrame:
    if df is None or df.empty:
        return pd.DataFrame(columns=["Documento", "Documento llave", "Nombre completo", "Correo", "Programa", "Grupo", "Estado", "Nota parcial", "Nota final", "Resultado parcial", "Resultado final"])
    out = pd.DataFrame(index=df.index)
    def col(role: str) -> Optional[str]:
        c = mapping.get(role)
        return c if c and c in df.columns and c != "— No usar —" else None
    out["Documento"] = df[col("documento")].apply(_clean_text) if col("documento") else ""
    out["Documento llave"] = out["Documento"].apply(_clean_document_key)
    out["Nombre completo"] = df[col("nombre")].apply(_clean_text) if col("nombre") else ""
    out["Correo"] = df[col("correo")].apply(_clean_text) if col("correo") else ""
    out["Programa"] = df[col("programa")].apply(_clean_text) if col("programa") else ""
    out["Grupo"] = df[col("grupo")].apply(_clean_text) if col("grupo") else ""
    out["Estado"] = df[col("estado")].apply(_estado_normalizado) if col("estado") else "Activo"
    nota_parcial_col = col("nota_parcial") or col("nota")
    nota_final_col = col("nota_final") or col("nota")
    out["Nota parcial"] = df[nota_parcial_col].apply(_coerce_grade) if nota_parcial_col else None
    out["Nota final"] = df[nota_final_col].apply(_coerce_grade) if nota_final_col else None
    out["Resultado parcial"] = out["Nota parcial"].apply(lambda v: "Sin nota" if v is None or pd.isna(v) else ("Aprueba" if float(v) >= nota_minima else "Reprueba"))
    out["Resultado final"] = out["Nota final"].apply(lambda v: "Sin nota" if v is None or pd.isna(v) else ("Aprueba" if float(v) >= nota_minima else "Reprueba"))
    out["Calidad registro"] = "OK"
    out.loc[out["Documento llave"].eq(""), "Calidad registro"] = "Sin documento"
    out.loc[out["Nombre completo"].eq(""), "Calidad registro"] = out["Calidad registro"].where(out["Calidad registro"].ne("OK"), "Sin nombre")
    out = out.dropna(how="all").reset_index(drop=True)
    # Elimina filas vacías totales.
    mask = out[["Documento llave", "Nombre completo", "Correo"]].astype(str).agg("".join, axis=1).str.strip().ne("")
    return out[mask].reset_index(drop=True)


def diagnostico_tabla_estudiantes(df: pd.DataFrame) -> pd.DataFrame:
    rows = []
    total = len(df) if df is not None else 0
    rows.append({"Tipo": "Info", "Hallazgo": "Registros procesados", "Cantidad": total, "Acción sugerida": "Validar muestra visual antes de descargar."})
    if df is None or df.empty:
        rows.append({"Tipo": "Error", "Hallazgo": "No hay registros útiles", "Cantidad": 0, "Acción sugerida": "Revise hoja, encabezado o archivo cargado."})
        return pd.DataFrame(rows)
    sin_doc = int(df["Documento llave"].eq("").sum()) if "Documento llave" in df.columns else 0
    if sin_doc:
        rows.append({"Tipo": "Alerta", "Hallazgo": "Estudiantes sin documento", "Cantidad": sin_doc, "Acción sugerida": "Completar documento o usar nombre como llave temporal."})
    sin_nombre = int(df["Nombre completo"].astype(str).str.strip().eq("").sum()) if "Nombre completo" in df.columns else 0
    if sin_nombre:
        rows.append({"Tipo": "Alerta", "Hallazgo": "Estudiantes sin nombre", "Cantidad": sin_nombre, "Acción sugerida": "Completar identificación nominal."})
    duplicados = int(df[df["Documento llave"].ne("")]["Documento llave"].duplicated().sum()) if "Documento llave" in df.columns else 0
    if duplicados:
        rows.append({"Tipo": "Error", "Hallazgo": "Documentos duplicados", "Cantidad": duplicados, "Acción sugerida": "Resolver duplicados antes de consolidar."})
    for c in ["Nota parcial", "Nota final"]:
        if c in df.columns:
            notas = pd.to_numeric(df[c], errors="coerce")
            fuera = int(((notas < 0) | (notas > 5)).sum())
            if fuera:
                rows.append({"Tipo": "Error", "Hallazgo": f"{c} fuera de escala 0-5", "Cantidad": fuera, "Acción sugerida": "Corregir escala o formato de notas."})
    desertores = int(df.get("Estado", pd.Series(dtype=str)).astype(str).str.contains("Desert", case=False, na=False).sum())
    if desertores:
        rows.append({"Tipo": "Info", "Hallazgo": "Registros marcados como desertores/retiros", "Cantidad": desertores, "Acción sugerida": "Validar que coincida con listado oficial."})
    return pd.DataFrame(rows)


def dataframe_to_xlsx_bytes(sheets: Dict[str, pd.DataFrame]) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
        for name, df in sheets.items():
            sheet = re.sub(r"[\\/*?:\[\]]", "_", str(name))[:31] or "Hoja"
            (df if df is not None else pd.DataFrame()).to_excel(writer, sheet_name=sheet, index=False)
            ws = writer.sheets[sheet]
            wb = writer.book
            fmt_header = wb.add_format({"bold": True, "text_wrap": True, "valign": "top", "border": 1})
            fmt_alert = wb.add_format({"bg_color": "#FFF2CC"})
            for i, col in enumerate((df if df is not None else pd.DataFrame()).columns):
                ws.write(0, i, col, fmt_header)
                width = min(max(len(str(col)) + 4, 12), 42)
                ws.set_column(i, i, width)
            if df is not None and not df.empty:
                ws.autofilter(0, 0, len(df), max(0, len(df.columns)-1))
    return bio.getvalue()


def ui_cargador_inteligente(st):
    st.markdown("""
<div class="ux-soft-box">
  <div class="ux-mini-title">Automatización crítica</div>
  <strong>Normalizador de listados y notas.</strong><br>
  Carga un Excel tradicional, detecta encabezados, sugiere columnas y entrega una tabla limpia para FD-GC72.
</div>
""", unsafe_allow_html=True)
    nota_minima = st.number_input("Nota mínima aprobatoria", min_value=0.0, max_value=5.0, value=3.0, step=0.1, key="smart_nota_min")
    uploaded = st.file_uploader("Cargue listado, notas o reporte académico", type=["xls", "xlsx", "csv"], key="smart_excel")
    if not uploaded:
        st.info("Cargue un archivo para iniciar. Este módulo está diseñado para sobrevivir al Excel institucional promedio. Valiente, pero no ingenuo.")
        return
    try:
        tablas = leer_excel_inteligente(uploaded)
    except Exception as e:
        st.error(f"No se pudo leer el archivo: {e}")
        return
    if not tablas:
        st.warning("No se encontraron hojas con datos útiles.")
        return
    hoja = st.selectbox("Hoja detectada", list(tablas.keys()), key="smart_sheet")
    df = tablas[hoja]
    st.caption(f"Filas detectadas: {len(df)} · Columnas: {len(df.columns)}")
    st.dataframe(df.head(30), use_container_width=True, hide_index=True)
    auto = _mapping_automatico(df)
    opciones = ["— No usar —"] + list(df.columns)
    st.subheader("Mapeo de columnas")
    cols = st.columns(3)
    roles_visibles = [
        ("documento", "Documento / llave"), ("nombre", "Nombre completo"), ("correo", "Correo"),
        ("estado", "Estado / observación"), ("nota_parcial", "Nota parcial"), ("nota_final", "Nota final"),
        ("programa", "Programa"), ("grupo", "Grupo"), ("nota", "Nota genérica")
    ]
    mapping: Dict[str, Optional[str]] = {}
    for i, (role, label) in enumerate(roles_visibles):
        default = auto.get(role) if auto.get(role) in opciones else "— No usar —"
        with cols[i % 3]:
            mapping[role] = st.selectbox(label, opciones, index=opciones.index(default), key=f"map_{role}")
    normalizado = normalizar_estudiantes_inteligente(df, mapping, nota_minima)
    diag = diagnostico_tabla_estudiantes(normalizado)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Registros", len(normalizado))
    c2.metric("Sin documento", int(normalizado["Documento llave"].eq("").sum()) if not normalizado.empty else 0)
    c3.metric("Duplicados", int(normalizado[normalizado["Documento llave"].ne("")]["Documento llave"].duplicated().sum()) if not normalizado.empty else 0)
    c4.metric("Desertores", int(normalizado.get("Estado", pd.Series(dtype=str)).astype(str).str.contains("Desert", case=False, na=False).sum()) if not normalizado.empty else 0)
    tab1, tab2 = st.tabs(["Tabla normalizada", "Diagnóstico"])
    with tab1:
        st.dataframe(normalizado, use_container_width=True, hide_index=True)
    with tab2:
        st.dataframe(diag, use_container_width=True, hide_index=True)
    st.session_state["smart_loader_normalizado"] = normalizado
    st.session_state["smart_loader_mapping"] = mapping
    xlsx = dataframe_to_xlsx_bytes({"Normalizado": normalizado, "Diagnostico": diag, "Original_muestra": df.head(200)})
    st.download_button("⬇️ Descargar Excel normalizado", xlsx, "listado_normalizado_fdgc.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
    st.download_button("⬇️ Descargar CSV normalizado", normalizado.to_csv(index=False).encode("utf-8-sig"), "listado_normalizado_fdgc.csv", "text/csv", use_container_width=True)


def _normalizar_automaticamente_archivo(uploaded_file, nota_minima: float) -> Tuple[pd.DataFrame, str, Dict[str, Optional[str]]]:
    tablas = leer_excel_inteligente(uploaded_file)
    if not tablas:
        return pd.DataFrame(), "", {}
    # Escoge la hoja con más filas útiles.
    hoja = max(tablas.keys(), key=lambda k: len(tablas[k]))
    df = tablas[hoja]
    mapping = _mapping_automatico(df)
    normalizado = normalizar_estudiantes_inteligente(df, mapping, nota_minima)
    return normalizado, hoja, mapping


def _base_comparacion(df_ini: pd.DataFrame, df_par: pd.DataFrame, df_fin: pd.DataFrame, nota_minima: float) -> pd.DataFrame:
    def prep(df: pd.DataFrame, pref: str) -> pd.DataFrame:
        if df is None or df.empty:
            return pd.DataFrame(columns=["llave"])
        tmp = df.copy()
        tmp["llave"] = tmp["Documento llave"].where(tmp["Documento llave"].astype(str).str.strip().ne(""), tmp["Nombre completo"].apply(lambda x: normalizar_texto(x).replace(" ", "")))
        cols = ["llave", "Documento", "Nombre completo", "Estado", "Nota parcial", "Nota final"]
        tmp = tmp[[c for c in cols if c in tmp.columns]].drop_duplicates("llave", keep="last")
        return tmp.rename(columns={c: f"{pref}_{c}" for c in tmp.columns if c != "llave"})
    base = prep(df_ini, "Inicial")
    par = prep(df_par, "Parcial")
    fin = prep(df_fin, "Final")
    comp = base.merge(par, on="llave", how="outer").merge(fin, on="llave", how="outer")
    if comp.empty:
        return comp
    comp["Documento"] = comp.get("Inicial_Documento", pd.Series(index=comp.index, dtype=object)).combine_first(comp.get("Parcial_Documento", pd.Series(index=comp.index, dtype=object))).combine_first(comp.get("Final_Documento", pd.Series(index=comp.index, dtype=object)))
    comp["Nombre completo"] = comp.get("Inicial_Nombre completo", pd.Series(index=comp.index, dtype=object)).combine_first(comp.get("Parcial_Nombre completo", pd.Series(index=comp.index, dtype=object))).combine_first(comp.get("Final_Nombre completo", pd.Series(index=comp.index, dtype=object)))
    def situacion(r):
        in_ini = not pd.isna(r.get("Inicial_Nombre completo")) or not pd.isna(r.get("Inicial_Documento"))
        in_par = not pd.isna(r.get("Parcial_Nombre completo")) or not pd.isna(r.get("Parcial_Documento"))
        in_fin = not pd.isna(r.get("Final_Nombre completo")) or not pd.isna(r.get("Final_Documento"))
        estado_fin = str(r.get("Final_Estado", ""))
        estado_par = str(r.get("Parcial_Estado", ""))
        nota_par = _coerce_grade(r.get("Parcial_Nota parcial"))
        nota_fin = _coerce_grade(r.get("Final_Nota final"))
        if not in_ini and (in_par or in_fin):
            return "Nuevo en cortes"
        if in_ini and not in_par and not in_fin:
            return "No aparece en cortes"
        if "Desert" in estado_fin or "Desert" in estado_par:
            return "Desertó / retiro"
        if nota_fin is not None and nota_fin < nota_minima:
            return "Reprueba al final"
        if nota_par is not None and nota_par < nota_minima and (nota_fin is None or nota_fin < nota_minima):
            return "Riesgo académico"
        if nota_fin is not None and nota_fin >= nota_minima:
            return "Aprueba al final"
        return "En seguimiento"
    comp["Situación"] = comp.apply(situacion, axis=1)
    return comp[["Documento", "Nombre completo", "Situación"] + [c for c in comp.columns if c not in {"Documento", "Nombre completo", "Situación"}]]


def _resumen_comparacion(comp: pd.DataFrame, df_ini: pd.DataFrame, df_par: pd.DataFrame, df_fin: pd.DataFrame, nota_minima: float) -> pd.DataFrame:
    if comp is None:
        comp = pd.DataFrame()
    def count_result(df, col, pred):
        if df is None or df.empty or col not in df.columns:
            return 0
        vals = pd.to_numeric(df[col], errors="coerce")
        return int(pred(vals).sum())
    matriculados = len(df_ini) if df_ini is not None else 0
    desertores = int(comp["Situación"].astype(str).str.contains("Desert|No aparece", case=False, na=False).sum()) if not comp.empty else 0
    apr_par = count_result(df_par, "Nota parcial", lambda s: s >= nota_minima)
    rep_par = count_result(df_par, "Nota parcial", lambda s: s < nota_minima)
    apr_fin = count_result(df_fin, "Nota final", lambda s: s >= nota_minima)
    rep_fin = count_result(df_fin, "Nota final", lambda s: s < nota_minima)
    return pd.DataFrame([{
        "Estudiantes matriculados": matriculados,
        "Estudiantes que Desertaron a la fecha": desertores,
        "Aprueban evaluación parcial N°": apr_par,
        "Reprueban evaluación parcial N°": rep_par,
        "Aprueban a la fecha N°": apr_fin,
        "Reprueban a la fecha N°": rep_fin,
        "Nuevos en cortes": int(comp["Situación"].eq("Nuevo en cortes").sum()) if not comp.empty else 0,
        "Riesgo académico": int(comp["Situación"].astype(str).str.contains("Riesgo", case=False, na=False).sum()) if not comp.empty else 0,
    }])


def ui_comparador_cortes(st):
    st.markdown("""
<div class="ux-soft-box">
  <div class="ux-mini-title">FD-GC72 casi automático</div>
  Compara el listado inicial contra cortes parcial y final para identificar desertores, nuevos, aprobados, reprobados y riesgo académico.
</div>
""", unsafe_allow_html=True)
    nota_minima = st.number_input("Nota mínima aprobatoria", min_value=0.0, max_value=5.0, value=3.0, step=0.1, key="cmp_nota_min")
    c1, c2, c3 = st.columns(3)
    with c1:
        f_ini = st.file_uploader("Listado inicial", type=["xls", "xlsx", "csv"], key="cmp_ini")
    with c2:
        f_par = st.file_uploader("Corte parcial", type=["xls", "xlsx", "csv"], key="cmp_par")
    with c3:
        f_fin = st.file_uploader("Corte final", type=["xls", "xlsx", "csv"], key="cmp_fin")
    if not f_ini:
        st.info("Suba al menos el listado inicial. Con parcial y final el análisis queda mucho más potente.")
        return
    try:
        df_ini, hoja_ini, _ = _normalizar_automaticamente_archivo(f_ini, nota_minima)
        df_par, hoja_par, _ = _normalizar_automaticamente_archivo(f_par, nota_minima) if f_par else (pd.DataFrame(), "", {})
        df_fin, hoja_fin, _ = _normalizar_automaticamente_archivo(f_fin, nota_minima) if f_fin else (pd.DataFrame(), "", {})
    except Exception as e:
        st.error(f"No se pudo procesar alguno de los archivos: {e}")
        return
    comp = _base_comparacion(df_ini, df_par, df_fin, nota_minima)
    resumen = _resumen_comparacion(comp, df_ini, df_par, df_fin, nota_minima)
    st.caption(f"Hojas usadas automáticamente: inicial={hoja_ini or 'N/A'} · parcial={hoja_par or 'N/A'} · final={hoja_fin or 'N/A'}")
    r = resumen.iloc[0].to_dict() if not resumen.empty else {}
    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Matrícula inicial", int(r.get("Estudiantes matriculados", 0)))
    k2.metric("Desertores / no aparecen", int(r.get("Estudiantes que Desertaron a la fecha", 0)))
    k3.metric("Aprueban parcial", int(r.get("Aprueban evaluación parcial N°", 0)))
    k4.metric("Riesgo académico", int(r.get("Riesgo académico", 0)))
    tab_res, tab_comp, tab_norm = st.tabs(["Resumen FD-GC72", "Comparación", "Normalizados"])
    with tab_res:
        st.dataframe(resumen, use_container_width=True, hide_index=True)
        st.markdown("**Lectura sugerida para análisis descriptivo:**")
        st.text_area("Texto base editable", value=texto_base_comparacion(comp, resumen), height=180)
    with tab_comp:
        st.dataframe(comp, use_container_width=True, hide_index=True)
    with tab_norm:
        st.write("Inicial")
        st.dataframe(df_ini, use_container_width=True, hide_index=True)
        if not df_par.empty:
            st.write("Parcial")
            st.dataframe(df_par, use_container_width=True, hide_index=True)
        if not df_fin.empty:
            st.write("Final")
            st.dataframe(df_fin, use_container_width=True, hide_index=True)
    xlsx = dataframe_to_xlsx_bytes({"Resumen_FDGC72": resumen, "Comparacion": comp, "Inicial_normalizado": df_ini, "Parcial_normalizado": df_par, "Final_normalizado": df_fin})
    st.download_button("⬇️ Descargar comparación completa", xlsx, "comparador_cortes_fdgc72.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)


def texto_base_comparacion(comp: pd.DataFrame, resumen: pd.DataFrame) -> str:
    if resumen is None or resumen.empty:
        return ""
    r = resumen.iloc[0].to_dict()
    matricula = int(r.get("Estudiantes matriculados", 0) or 0)
    desertores = int(r.get("Estudiantes que Desertaron a la fecha", 0) or 0)
    apr_par = int(r.get("Aprueban evaluación parcial N°", 0) or 0)
    rep_par = int(r.get("Reprueban evaluación parcial N°", 0) or 0)
    riesgo = int(r.get("Riesgo académico", 0) or 0)
    return (
        f"Durante el corte analizado se registra una matrícula base de {matricula} estudiante(s). "
        f"El seguimiento comparativo identifica {desertores} estudiante(s) con retiro, deserción o ausencia en cortes posteriores. "
        f"En la evaluación parcial aprueban {apr_par} estudiante(s) y reprueban {rep_par}. "
        f"Se recomienda priorizar acompañamiento académico a {riesgo} estudiante(s) en condición de riesgo, reforzando actividades de retroalimentación, talleres aplicados y seguimiento individual."
    )


def _componentes_semaforo(curso_id: int) -> pd.DataFrame:
    curso = get_curso(int(curso_id)) or {}
    payload = safe_json_loads(curso.get("payload_json"), {})
    sesiones = payload_to_df(payload.get("sesiones"), COLUMNAS_SESIONES)
    evaluaciones = payload_to_df(payload.get("evaluaciones"), COLUMNAS_EVALUACIONES)
    modulos = payload_to_df(payload.get("modulos"), COLUMNAS_MODULOS)
    total_eval = sum(limpiar_numero(v) or 0 for v in evaluaciones.get("Valor (%)", [])) if not evaluaciones.empty else 0
    evid_count = evidencias_count(int(curso_id))
    obs = observaciones_curso(int(curso_id))
    obs_abiertas = int((obs.get("estado", pd.Series(dtype=str)).astype(str) == "Abierta").sum()) if not obs.empty else 0
    score, hallazgos = score_calidad_expediente(int(curso_id))
    datos_ok = all(str(curso.get(k, "")).strip() for k in ["asignatura", "programa", "periodo", "profesor"])
    rows = [
        {"Componente": "Datos base", "Avance": 100 if datos_ok else 55, "Estado": "Verde" if datos_ok else "Amarillo", "Lectura": "Identificación mínima completa" if datos_ok else "Faltan campos de identificación"},
        {"Componente": "Módulos/unidades", "Avance": 100 if not modulos.empty else 0, "Estado": "Verde" if not modulos.empty else "Rojo", "Lectura": f"{len(modulos)} unidad(es) registradas"},
        {"Componente": "Cronograma", "Avance": 100 if not sesiones.empty else 0, "Estado": "Verde" if not sesiones.empty else "Rojo", "Lectura": f"{len(sesiones)} sesión(es) generadas"},
        {"Componente": "Evaluación concertada", "Avance": min(100, int(total_eval)), "Estado": "Verde" if abs(total_eval - 100) < 0.01 else ("Amarillo" if total_eval > 0 else "Rojo"), "Lectura": f"Suma evaluación: {total_eval:.1f}%"},
        {"Componente": "Evidencias", "Avance": 100 if evid_count >= 3 else (65 if evid_count > 0 else 0), "Estado": "Verde" if evid_count >= 3 else ("Amarillo" if evid_count > 0 else "Rojo"), "Lectura": f"{evid_count} evidencia(s) cargadas"},
        {"Componente": "Observaciones", "Avance": 100 if obs_abiertas == 0 else 45, "Estado": "Verde" if obs_abiertas == 0 else "Rojo", "Lectura": f"{obs_abiertas} observación(es) abierta(s)"},
        {"Componente": "Score institucional", "Avance": int(score), "Estado": "Verde" if score >= 80 else ("Amarillo" if score >= 60 else "Rojo"), "Lectura": f"Score de calidad: {score}/100"},
    ]
    return pd.DataFrame(rows)


def ui_semaforo_expediente(st):
    dfc = listar_cursos_visibles()
    if dfc.empty:
        st.info("No hay cursos visibles para construir semáforo.")
        return
    opciones = {f"{r.get('id')} · {r.get('asignatura','')} · {r.get('grupo','')} · {r.get('periodo','')}": int(r.get('id')) for _, r in dfc.iterrows()}
    sel = st.selectbox("Curso", list(opciones.keys()), key="sem_curso")
    curso_id = opciones[sel]
    comp = _componentes_semaforo(curso_id)
    score = int(comp[comp["Componente"].eq("Score institucional")]["Avance"].iloc[0]) if not comp.empty else 0
    c1, c2, c3 = st.columns(3)
    c1.metric("Score expediente", f"{score}/100")
    c2.metric("Componentes en rojo", int((comp["Estado"] == "Rojo").sum()))
    c3.metric("Componentes en verde", int((comp["Estado"] == "Verde").sum()))
    st.dataframe(comp, use_container_width=True, hide_index=True)
    for _, row in comp.iterrows():
        st.progress(float(row["Avance"]) / 100.0, text=f"{row['Componente']}: {row['Avance']}% · {row['Lectura']}")
    st.download_button("⬇️ Descargar semáforo CSV", comp.to_csv(index=False).encode("utf-8-sig"), f"semaforo_expediente_{curso_id}.csv", "text/csv", use_container_width=True)


def ui_exportacion_masiva(st):
    df = listar_cursos_visibles()
    if df.empty:
        st.info("No hay cursos visibles para exportar.")
        return
    st.markdown("<div class='ux-soft-box'><strong>Exportación institucional masiva.</strong><br>Filtre cursos y descargue paquetes individuales dentro de un ZIP maestro.</div>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    with c1:
        periodo = st.selectbox("Periodo", ["Todos"] + sorted(df.get("periodo", pd.Series(dtype=str)).dropna().astype(str).unique().tolist()), key="exp_periodo")
    with c2:
        programa = st.selectbox("Programa", ["Todos"] + sorted(df.get("programa", pd.Series(dtype=str)).dropna().astype(str).unique().tolist()), key="exp_programa")
    with c3:
        estado = st.selectbox("Estado", ["Todos"] + sorted(df.get("estado", pd.Series(dtype=str)).dropna().astype(str).unique().tolist()), key="exp_estado")
    f = df.copy()
    if periodo != "Todos": f = f[f["periodo"].astype(str) == periodo]
    if programa != "Todos": f = f[f["programa"].astype(str) == programa]
    if estado != "Todos": f = f[f["estado"].astype(str) == estado]
    st.metric("Cursos a exportar", len(f))
    st.dataframe(f[[c for c in ["id", "codigo", "grupo", "asignatura", "programa", "periodo", "profesor", "estado"] if c in f.columns]], use_container_width=True, hide_index=True)
    if st.button("Preparar ZIP maestro", type="primary", use_container_width=True):
        bio = io.BytesIO()
        indice_rows = []
        with zipfile.ZipFile(bio, "w", zipfile.ZIP_DEFLATED) as z:
            for _, r in f.iterrows():
                curso_id = int(r.get("id"))
                curso = get_curso(curso_id) or {}
                payload = safe_json_loads(curso.get("payload_json"), {})
                datos = payload.get("datos", dict(curso))
                sesiones = payload_to_df(payload.get("sesiones"), COLUMNAS_SESIONES)
                evaluaciones = payload_to_df(payload.get("evaluaciones"), COLUMNAS_EVALUACIONES)
                estudiantes = payload_to_df(payload.get("estudiantes"), ["Nombre completo", "Documento", "Correo", "Plan", "Observación", "Estado"])
                representantes = payload_to_df(payload.get("representantes"), COLUMNAS_REPRESENTANTES) if "COLUMNAS_REPRESENTANTES" in globals() else pd.DataFrame()
                nombre = f"{curso_id}_{safe_filename(str(curso.get('asignatura','curso')))}"
                try:
                    paquete = crear_paquete_curso_zip(datos, sesiones, evaluaciones, estudiantes, representantes, curso_id)
                    z.writestr(f"expedientes/{nombre}.zip", paquete)
                    estado_exp = "OK"
                except Exception as e:
                    z.writestr(f"expedientes/{nombre}_ERROR.txt", f"No se pudo generar paquete: {e}")
                    estado_exp = f"ERROR: {e}"
                score, _ = score_calidad_expediente(curso_id)
                indice_rows.append({"id": curso_id, "asignatura": curso.get("asignatura", ""), "grupo": curso.get("grupo", ""), "programa": curso.get("programa", ""), "periodo": curso.get("periodo", ""), "estado": curso.get("estado", ""), "score": score, "exportacion": estado_exp})
            indice = pd.DataFrame(indice_rows)
            z.writestr("indice_expedientes.csv", indice.to_csv(index=False).encode("utf-8-sig"))
            z.writestr("manifiesto.json", json.dumps({"generado_en": ahora_iso(), "usuario": st.session_state.get("auth_user", {}).get("usuario"), "total_cursos": len(indice_rows), "version": APP_VERSION}, ensure_ascii=False, indent=2))
        st.download_button("⬇️ Descargar ZIP maestro", bio.getvalue(), f"exportacion_masiva_fdgc_{datetime.now().strftime('%Y%m%d_%H%M')}.zip", "application/zip", use_container_width=True)
        registrar_auditoria("Exportación masiva", f"Cursos exportados: {len(f)}")




def safe_filename(value: str, max_len: int = 80) -> str:
    name = re.sub(r"[^A-Za-z0-9_. -]", "_", str(value or "archivo")).strip(" ._")
    name = re.sub(r"_+", "_", name)
    return (name or "archivo")[:max_len]

def _hash_expediente(curso_id: int) -> Tuple[str, Dict[str, Any]]:
    curso = get_curso(int(curso_id)) or {}
    payload = safe_json_loads(curso.get("payload_json"), {})
    score, hallazgos = score_calidad_expediente(int(curso_id))
    evid = evidencias_count(int(curso_id))
    versiones = versiones_curso(int(curso_id)) if "versiones_curso" in globals() else pd.DataFrame()
    data = {
        "app_version": APP_VERSION,
        "curso_id": int(curso_id),
        "codigo": curso.get("codigo", ""),
        "grupo": curso.get("grupo", ""),
        "asignatura": curso.get("asignatura", ""),
        "programa": curso.get("programa", ""),
        "periodo": curso.get("periodo", ""),
        "estado": curso.get("estado", ""),
        "actualizado_en": curso.get("actualizado_en", ""),
        "score": score,
        "evidencias": evid,
        "versiones": len(versiones) if versiones is not None else 0,
        "payload_sha256": hashlib.sha256(json.dumps(payload, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest(),
    }
    digest = hashlib.sha256(json.dumps(data, ensure_ascii=False, sort_keys=True).encode("utf-8")).hexdigest()
    data["sha256_expediente"] = digest
    return digest, data


def ui_verificacion_documental(st):
    dfc = listar_cursos_visibles()
    if dfc.empty:
        st.info("No hay cursos visibles para verificar.")
        return
    opciones = {f"{r.get('id')} · {r.get('asignatura','')} · {r.get('grupo','')} · {r.get('periodo','')}": int(r.get('id')) for _, r in dfc.iterrows()}
    sel = st.selectbox("Curso", list(opciones.keys()), key="hash_curso")
    curso_id = opciones[sel]
    digest, data = _hash_expediente(curso_id)
    st.code(digest, language="text")
    st.json(data)
    payload_json = json.dumps(data, ensure_ascii=False, indent=2)
    st.download_button("⬇️ Descargar ficha de verificación JSON", payload_json.encode("utf-8"), f"verificacion_expediente_{curso_id}.json", "application/json", use_container_width=True)
    try:
        import qrcode
        img = qrcode.make(payload_json[:2500])
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        st.image(buf.getvalue(), caption="QR de verificación del expediente")
        st.download_button("⬇️ Descargar QR PNG", buf.getvalue(), f"qr_expediente_{curso_id}.png", "image/png", use_container_width=True)
    except Exception:
        st.info("Para generar QR instale la dependencia opcional `qrcode[pil]`. La ficha JSON y el hash ya permiten verificación documental.")


def ui_asistente_academico(st):
    st.markdown("<div class='ux-soft-box'><strong>Redacción asistida editable.</strong><br>No reemplaza criterio docente: entrega borradores limpios para acelerar FD-GC71 y FD-GC72.</div>", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    with c1:
        asignatura = st.text_input("Asignatura", value="Sistemas de Información Geográfica")
        programa = st.text_input("Programa académico", value="")
        modalidad = st.selectbox("Modalidad", ["Presencial", "Teórico-práctica", "Práctica", "Virtual", "Mixta"])
    with c2:
        enfoque = st.text_area("Enfoque o propósito del curso", value="Aplicación de conceptos, métodos y herramientas para resolver problemas del contexto profesional.", height=100)
        retos = st.text_area("Retos del grupo / contexto", value="Heterogeneidad en conocimientos previos, necesidad de fortalecer trabajo autónomo y aplicación práctica.", height=100)
    nivel = st.selectbox("Nivel de detalle", ["Ejecutivo", "Robusto", "Muy robusto"], index=1)
    long = {"Ejecutivo": 1, "Robusto": 2, "Muy robusto": 3}[nivel]
    if st.button("Generar borradores", type="primary", use_container_width=True):
        base = generar_textos_academicos(asignatura, programa, modalidad, enfoque, retos, long)
        st.session_state["asistente_textos"] = base
    textos = st.session_state.get("asistente_textos")
    if textos:
        tab_gc71, tab_gc72 = st.tabs(["FD-GC71", "FD-GC72"])
        with tab_gc71:
            for k in ["Justificación", "Competencias", "Resultados de aprendizaje", "Objetivo general", "Objetivos específicos", "Metodología", "Ambientes", "Medios educativos"]:
                st.text_area(k, value=textos.get(k, ""), height=120, key=f"asis_{k}")
        with tab_gc72:
            for k in ["Aspectos positivos", "Inconvenientes", "Propuestas metodológicas", "Plan de mejora"]:
                st.text_area(k, value=textos.get(k, ""), height=120, key=f"asis_{k}")
        st.download_button("⬇️ Descargar borradores TXT", "\n\n".join([f"## {k}\n{v}" for k, v in textos.items()]).encode("utf-8"), "borradores_academicos_fdgc.txt", "text/plain", use_container_width=True)


def generar_textos_academicos(asignatura: str, programa: str, modalidad: str, enfoque: str, retos: str, nivel: int = 2) -> Dict[str, str]:
    asignatura = asignatura.strip() or "la asignatura"
    programa_txt = f" del programa {programa.strip()}" if programa.strip() else ""
    extra = " Asimismo, se promueve la articulación entre los contenidos, los resultados de aprendizaje, la evaluación formativa y la aplicación contextualizada en escenarios propios del desempeño profesional." if nivel >= 2 else ""
    extra2 = " El desarrollo del curso deberá privilegiar evidencias verificables, criterios de evaluación explícitos y retroalimentación oportuna, de manera que el proceso académico sea trazable, pertinente y susceptible de mejora continua." if nivel >= 3 else ""
    return {
        "Justificación": f"La asignatura {asignatura}{programa_txt} aporta a la formación académica y profesional mediante {enfoque.strip()}. Su desarrollo permite integrar fundamentos conceptuales, ejercicios aplicados y análisis de situaciones reales, fortaleciendo la capacidad del estudiante para tomar decisiones informadas en contextos disciplinares y laborales.{extra}{extra2}",
        "Competencias": f"El curso contribuye al fortalecimiento de competencias de análisis, interpretación, solución de problemas, comunicación técnica y aplicación metodológica. De manera particular, busca que el estudiante relacione los conceptos de {asignatura} con situaciones del contexto profesional, argumente sus decisiones y produzca evidencias académicas coherentes con los criterios establecidos.",
        "Resultados de aprendizaje": f"Al finalizar el curso, el estudiante estará en capacidad de reconocer los fundamentos de {asignatura}, aplicar procedimientos pertinentes, interpretar resultados, sustentar decisiones y elaborar productos académicos acordes con los requerimientos de la asignatura y del programa.",
        "Objetivo general": f"Desarrollar en el estudiante capacidades conceptuales, metodológicas y aplicadas asociadas con {asignatura}, mediante estrategias de aprendizaje {modalidad.lower()} orientadas a la solución de problemas y a la producción de evidencias académicas verificables.",
        "Objetivos específicos": "1. Reconocer los conceptos fundamentales de la asignatura.\n2. Aplicar métodos y herramientas pertinentes en ejercicios académicos y prácticos.\n3. Interpretar resultados y sustentar decisiones con criterios técnicos.\n4. Elaborar productos evaluables que evidencien apropiación progresiva de los aprendizajes.",
        "Metodología": f"La metodología combinará explicación orientadora, ejercicios guiados, trabajo aplicado, análisis de casos, actividades colaborativas y retroalimentación permanente. Considerando {retos.strip()}, se priorizarán actividades progresivas que permitan diagnosticar conocimientos previos, acompañar el avance del estudiante y consolidar aprendizajes mediante productos verificables.",
        "Ambientes": "El desarrollo de la asignatura podrá apoyarse en aula de clase, sala de sistemas, plataforma académica institucional, recursos digitales, bases de datos, guías de trabajo, ejercicios prácticos y espacios de socialización de resultados, según la naturaleza de las actividades programadas.",
        "Medios educativos": "Se emplearán presentaciones, documentos guía, lecturas académicas, recursos audiovisuales, herramientas tecnológicas, talleres, rúbricas de evaluación, bases de datos de práctica y material complementario dispuesto por el docente para favorecer el aprendizaje autónomo y presencial.",
        "Aspectos positivos": "Se evidencia avance en la apropiación de los contenidos desarrollados, participación en las actividades propuestas y disposición para relacionar los temas de clase con situaciones aplicadas. Las actividades de seguimiento han permitido identificar fortalezas, orientar la retroalimentación y consolidar productos académicos progresivos.",
        "Inconvenientes": f"Se identifican retos asociados con {retos.strip()}. Estos aspectos requieren seguimiento para evitar rezagos en el cumplimiento de actividades, dificultades en la comprensión de contenidos o disminución en el rendimiento académico del grupo.",
        "Propuestas metodológicas": "Fortalecer la retroalimentación temprana, implementar ejercicios aplicados por unidad temática, socializar rúbricas antes de cada evaluación, habilitar espacios de refuerzo y promover entregas parciales que permitan corregir dificultades antes del cierre del corte académico.",
        "Plan de mejora": "Priorizar estudiantes con bajo desempeño, programar actividades de nivelación, reforzar instrucciones de trabajo independiente, verificar comprensión de criterios de evaluación y documentar evidencias de acompañamiento académico durante el periodo.",
    }


def ui_inicio(st):
    counts = ux_system_counts()
    st.markdown(
        "<div class='ux-card-grid'>" +
        ux_card("🧩", "Excel inteligente", "Normaliza listados y notas con mapeo asistido para no depender de plantillas perfectas.") +
        ux_card("🔁", "Comparador de cortes", "Cruza inicio, parcial y final para alimentar el FD-GC72 con trazabilidad.") +
        ux_card("🚦", "Semáforo académico", "Muestra avance, evidencias, observaciones, evaluación y score por expediente.") +
        ux_card("🔐", "Hash y QR", "Genera huellas verificables para blindar documentos y expedientes.") +
        "</div>", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Cursos", counts.get("cursos", 0))
    c2.metric("Riesgo alto", counts.get("alto", 0))
    c3.metric("Observaciones abiertas", counts.get("abiertas", 0))
    c4.metric("Score promedio", f"{counts.get('score', 0):.1f}")
    st.markdown("### Ruta recomendada")
    st.markdown("""
1. Cree o actualice el expediente del curso.  
2. Planee FD-GC71 con módulos, horario y evaluación.  
3. Cargue listado inicial, parcial y final en el comparador.  
4. Valide semáforo, evidencias, observaciones y score.  
5. Genere FD-GC72, cierre y exporte expediente con hash.  
""")


# Router final V6. Sobrescribe main para incluir módulos de siguiente nivel.
def main():
    import streamlit as st
    globals()["st"] = st
    st.set_page_config(
        page_title="Gestor Académico Inteligente FD-GC71 / FD-GC72",
        page_icon="🎓",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    ux_apply_theme(st)
    try:
        init_db()
    except Exception as exc:
        st.error("La aplicación inició, pero no pudo preparar la base de datos.")
        st.markdown("""
        Esto suele pasar en Streamlit Cloud cuando `DATABASE_URL` está mal escrita,
        contiene el ejemplo `USUARIO:CLAVE@HOST:5432/BASE`, la base de datos no acepta conexiones externas
        o el proveedor requiere usar el pooler/puerto correcto.
        """)
        st.code(str(exc), language="text")
        st.info("Corrija los Secrets en Streamlit Cloud y reinicie la app. Para probar sin base externa, quite `DATABASE_URL` y use modo local/demo.")
        st.stop()

    if "auth_user" not in st.session_state:
        pantalla_login(st)
        return

    user = st.session_state.get("auth_user", {})
    pagina = ux_sidebar(st, user)

    if not tiene_permiso(pagina):
        st.error("Este perfil no tiene permisos para abrir este módulo.")
        return

    ux_render_hero(st, pagina, user)
    if pagina not in ["Inicio", "Mi cuenta", "Ayuda / flujo recomendado", MODULO_DIAGNOSTICO]:
        ux_render_path(st, pagina)

    if pagina == "Inicio":
        ui_inicio(st)
    elif pagina == MODULO_CENTRO:
        ui_centro_control(st)
    elif pagina == MODULO_SEMAFORO:
        ui_semaforo_expediente(st)
    elif pagina == MODULO_EXPEDIENTE:
        ui_expediente_academico(st)
    elif pagina == MODULO_PLANEADOR:
        ui_planeador_superior(st)
    elif pagina == MODULO_ASISTENTE:
        ui_asistente_academico(st)
    elif pagina == MODULO_CARGADOR:
        ui_cargador_inteligente(st)
    elif pagina == MODULO_COMPARADOR:
        ui_comparador_cortes(st)
    elif pagina == MODULO_FLUJO:
        ui_flujo_aprobaciones(st)
    elif pagina == MODULO_MOTOR:
        ui_motor_academico(st)
    elif pagina == MODULO_REPORTES:
        ui_reportes_ejecutivos(st)
    elif pagina == MODULO_EXPORTACION:
        ui_exportacion_masiva(st)
    elif pagina == MODULO_VERIFICACION:
        ui_verificacion_documental(st)
    elif pagina == MODULO_PARAMETROS:
        ui_parametros(st)
    elif pagina.startswith("FD-GC71"):
        ui_gc71(st)
    elif pagina.startswith("FD-GC72"):
        ui_gc72(st)
    elif pagina == MODULO_EVIDENCIAS:
        ui_evidencias(st)
    elif pagina == MODULO_VALIDADOR:
        ui_validador_institucional(st)
    elif pagina == MODULO_BACKUP:
        ui_backup(st)
    elif pagina == MODULO_DIAGNOSTICO:
        ui_diagnostico_productivo(st)
    elif pagina == "Usuarios y perfiles":
        ui_admin_usuarios(st)
    elif pagina == "Auditoría":
        ui_auditoria(st)
    elif pagina == "Mi cuenta":
        ui_mi_cuenta(st)
    else:
        ui_ayuda(st)

    st.markdown("<div class='ux-footer-note'>Gestor Académico Inteligente · planeación → validación → cortes → informe → cierre → verificación.</div>", unsafe_allow_html=True)


# =============================================================================
# INSTITUCIONAL V7: banco de asignaturas, coherencia académica, aprobación
# bloqueante, informe ejecutivo y exportación institucional estructurada.
# =============================================================================
APP_VERSION = "7.0.1-hotfix-streamlit-cloud"
MODULO_BANCO = "Banco institucional de asignaturas"
MODULO_COHERENCIA = "Coherencia académica"
MODULO_APROBACION_BLOQUEANTE = "Aprobación bloqueante"
MODULO_INFORME_INSTITUCIONAL = "Informe ejecutivo institucional"
MODULO_EXPORTACION_INSTITUCIONAL = "Exportación institucional estructurada"
MODULO_AUDITORIA_EXPEDIENTE = "Auditoría de expediente"

try:
    UX_MODULE_META.update({
        MODULO_BANCO: {"icon": "🏛️", "group": "Administración", "desc": "Catálogo maestro de asignaturas, unidades, resultados y evaluaciones base.", "step": 2},
        MODULO_COHERENCIA: {"icon": "🧪", "group": "Gobierno", "desc": "Valida alineación entre resultados, contenidos, evaluación, horas y evidencias.", "step": 6},
        MODULO_APROBACION_BLOQUEANTE: {"icon": "🔒", "group": "Gobierno", "desc": "Flujo formal con bloqueo de versiones aprobadas y cierre del expediente.", "step": 7},
        MODULO_INFORME_INSTITUCIONAL: {"icon": "🏢", "group": "Dirección", "desc": "Informe ejecutivo para coordinación/comité con riesgos, cumplimiento y recomendaciones.", "step": 8},
        MODULO_EXPORTACION_INSTITUCIONAL: {"icon": "🗃️", "group": "Dirección", "desc": "ZIP maestro por período/programa con manifiesto, matriz, informes y expedientes.", "step": 8},
        MODULO_AUDITORIA_EXPEDIENTE: {"icon": "🧾", "group": "Gobierno", "desc": "Trazabilidad detallada por expediente, usuario, estado, observación y versión.", "step": 8},
    })
except Exception:
    pass


def _v7_add_modulos_a_rol(rol: str, nuevos: List[str]):
    if rol not in ROLES_PERMISOS:
        return
    actuales = ROLES_PERMISOS[rol].setdefault("modulos", [])
    for m in nuevos:
        if m not in actuales:
            actuales.append(m)

for _rol in ["Administrador"]:
    _v7_add_modulos_a_rol(_rol, [MODULO_BANCO, MODULO_COHERENCIA, MODULO_APROBACION_BLOQUEANTE, MODULO_AUDITORIA_EXPEDIENTE, MODULO_INFORME_INSTITUCIONAL, MODULO_EXPORTACION_INSTITUCIONAL])
for _rol in ["Coordinador"]:
    _v7_add_modulos_a_rol(_rol, [MODULO_BANCO, MODULO_COHERENCIA, MODULO_APROBACION_BLOQUEANTE, MODULO_AUDITORIA_EXPEDIENTE, MODULO_INFORME_INSTITUCIONAL, MODULO_EXPORTACION_INSTITUCIONAL])
for _rol in ["Docente"]:
    _v7_add_modulos_a_rol(_rol, [MODULO_COHERENCIA, MODULO_APROBACION_BLOQUEANTE, MODULO_AUDITORIA_EXPEDIENTE])
for _rol in ["Consulta"]:
    _v7_add_modulos_a_rol(_rol, [MODULO_INFORME_INSTITUCIONAL])

_base_init_db_v7 = init_db

def init_db():
    """Inicializa la plataforma institucional V7 sin romper instalaciones previas."""
    _base_init_db_v7()
    conn = conexion_db()
    try:
        if usar_postgres():
            conn.execute("""
                CREATE TABLE IF NOT EXISTS asignaturas_base (
                    id SERIAL PRIMARY KEY,
                    codigo TEXT,
                    nombre TEXT NOT NULL,
                    programa TEXT,
                    area_formacion TEXT,
                    creditos TEXT,
                    htp DOUBLE PRECISION DEFAULT 0,
                    hti DOUBLE PRECISION DEFAULT 0,
                    tipo_asignatura TEXT,
                    justificacion TEXT,
                    competencias TEXT,
                    resultados TEXT,
                    objetivos TEXT,
                    metodologia TEXT,
                    ambientes TEXT,
                    medios TEXT,
                    bibliografia TEXT,
                    unidades_json TEXT DEFAULT '[]',
                    evaluaciones_json TEXT DEFAULT '[]',
                    activo INTEGER DEFAULT 1,
                    creado_por TEXT,
                    creado_en TEXT NOT NULL,
                    actualizado_en TEXT
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS workflow_eventos (
                    id SERIAL PRIMARY KEY,
                    curso_id INTEGER REFERENCES cursos(id) ON DELETE CASCADE,
                    evento TEXT NOT NULL,
                    estado_anterior TEXT,
                    estado_nuevo TEXT,
                    resultado TEXT,
                    detalle TEXT,
                    hash_expediente TEXT,
                    usuario TEXT,
                    rol TEXT,
                    creado_en TEXT NOT NULL
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS curso_bloqueos (
                    curso_id INTEGER PRIMARY KEY REFERENCES cursos(id) ON DELETE CASCADE,
                    bloqueado INTEGER DEFAULT 0,
                    motivo TEXT,
                    hash_bloqueo TEXT,
                    bloqueado_por TEXT,
                    bloqueado_en TEXT
                )
            """)
        else:
            conn.execute("""
                CREATE TABLE IF NOT EXISTS asignaturas_base (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    codigo TEXT,
                    nombre TEXT NOT NULL,
                    programa TEXT,
                    area_formacion TEXT,
                    creditos TEXT,
                    htp REAL DEFAULT 0,
                    hti REAL DEFAULT 0,
                    tipo_asignatura TEXT,
                    justificacion TEXT,
                    competencias TEXT,
                    resultados TEXT,
                    objetivos TEXT,
                    metodologia TEXT,
                    ambientes TEXT,
                    medios TEXT,
                    bibliografia TEXT,
                    unidades_json TEXT DEFAULT '[]',
                    evaluaciones_json TEXT DEFAULT '[]',
                    activo INTEGER DEFAULT 1,
                    creado_por TEXT,
                    creado_en TEXT NOT NULL,
                    actualizado_en TEXT
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS workflow_eventos (
                    id INTEGER PRIMARY KEY AUTOINCREMENT,
                    curso_id INTEGER,
                    evento TEXT NOT NULL,
                    estado_anterior TEXT,
                    estado_nuevo TEXT,
                    resultado TEXT,
                    detalle TEXT,
                    hash_expediente TEXT,
                    usuario TEXT,
                    rol TEXT,
                    creado_en TEXT NOT NULL,
                    FOREIGN KEY(curso_id) REFERENCES cursos(id)
                )
            """)
            conn.execute("""
                CREATE TABLE IF NOT EXISTS curso_bloqueos (
                    curso_id INTEGER PRIMARY KEY,
                    bloqueado INTEGER DEFAULT 0,
                    motivo TEXT,
                    hash_bloqueo TEXT,
                    bloqueado_por TEXT,
                    bloqueado_en TEXT,
                    FOREIGN KEY(curso_id) REFERENCES cursos(id)
                )
            """)
        conn.commit()
    finally:
        conn.close()


def _text_lines(value: Any) -> List[str]:
    raw = str(value or "").replace("\r", "\n")
    parts = []
    for line in raw.split("\n"):
        clean = re.sub(r"^\s*[-•*\d\.\)]+\s*", "", line).strip()
        if clean:
            parts.append(clean)
    if not parts and raw.strip():
        # Divide textos largos por punto y coma si no vienen en lista.
        parts = [p.strip() for p in re.split(r"[;]+", raw) if p.strip()]
    return parts


def _norm_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "").lower()).strip()


def _json_df(value: Any, columns: Optional[List[str]] = None) -> pd.DataFrame:
    if isinstance(value, pd.DataFrame):
        return value.copy()
    data = safe_json_loads(value, []) if isinstance(value, str) else (value or [])
    try:
        df = pd.DataFrame(data)
    except Exception:
        df = pd.DataFrame(columns=columns or [])
    if columns:
        for c in columns:
            if c not in df.columns:
                df[c] = ""
        df = df[columns]
    return df


def listar_asignaturas_base() -> pd.DataFrame:
    try:
        return read_sql_df("SELECT * FROM asignaturas_base WHERE activo=1 ORDER BY programa, nombre")
    except Exception:
        return pd.DataFrame()


def guardar_asignatura_base(data: Dict[str, Any], asignatura_id: Optional[int] = None) -> int:
    user = st.session_state.get("auth_user", {}) if "st" in globals() else {}
    now = ahora_iso()
    unidades_json = json.dumps(data.get("unidades", []), ensure_ascii=False, default=str)
    evaluaciones_json = json.dumps(data.get("evaluaciones", []), ensure_ascii=False, default=str)
    params = (
        data.get("codigo", ""), data.get("nombre", ""), data.get("programa", ""), data.get("area_formacion", ""),
        data.get("creditos", ""), limpiar_numero(data.get("htp", 0)) or 0, limpiar_numero(data.get("hti", 0)) or 0,
        data.get("tipo_asignatura", ""), data.get("justificacion", ""), data.get("competencias", ""),
        data.get("resultados", ""), data.get("objetivos", ""), data.get("metodologia", ""), data.get("ambientes", ""),
        data.get("medios", ""), data.get("bibliografia", ""), unidades_json, evaluaciones_json,
    )
    conn = conexion_db()
    try:
        if asignatura_id:
            conn.execute("""
                UPDATE asignaturas_base SET codigo=?, nombre=?, programa=?, area_formacion=?, creditos=?, htp=?, hti=?, tipo_asignatura=?, justificacion=?, competencias=?, resultados=?, objetivos=?, metodologia=?, ambientes=?, medios=?, bibliografia=?, unidades_json=?, evaluaciones_json=?, actualizado_en=? WHERE id=?
            """, params + (now, int(asignatura_id)))
            new_id = int(asignatura_id)
        else:
            if usar_postgres():
                cur = conn.execute("""
                    INSERT INTO asignaturas_base(codigo, nombre, programa, area_formacion, creditos, htp, hti, tipo_asignatura, justificacion, competencias, resultados, objetivos, metodologia, ambientes, medios, bibliografia, unidades_json, evaluaciones_json, creado_por, creado_en, actualizado_en)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?) RETURNING id
                """, params + (user.get("usuario", ""), now, now))
                new_id = int(cur.fetchone()["id"])
            else:
                cur = conn.execute("""
                    INSERT INTO asignaturas_base(codigo, nombre, programa, area_formacion, creditos, htp, hti, tipo_asignatura, justificacion, competencias, resultados, objetivos, metodologia, ambientes, medios, bibliografia, unidades_json, evaluaciones_json, creado_por, creado_en, actualizado_en)
                    VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                """, params + (user.get("usuario", ""), now, now))
                new_id = int(cur.lastrowid)
        conn.commit()
        registrar_auditoria("Guardar asignatura base", f"Asignatura base ID={new_id} | {data.get('nombre','')}")
        return new_id
    finally:
        conn.close()


def get_asignatura_base(asignatura_id: int) -> Optional[Dict[str, Any]]:
    row = db_execute("SELECT * FROM asignaturas_base WHERE id=?", (int(asignatura_id),), fetchone=True)
    return _dict_row(row) if row else None


def registrar_workflow_evento(curso_id: int, evento: str, anterior: str = "", nuevo: str = "", resultado: str = "", detalle: str = ""):
    user = st.session_state.get("auth_user", {}) if "st" in globals() else {}
    try:
        digest, _ = _hash_expediente(int(curso_id)) if "_hash_expediente" in globals() else ("", {})
    except Exception:
        digest = ""
    db_execute(
        """
        INSERT INTO workflow_eventos(curso_id, evento, estado_anterior, estado_nuevo, resultado, detalle, hash_expediente, usuario, rol, creado_en)
        VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
        """,
        (int(curso_id), evento, anterior, nuevo, resultado, detalle, digest, user.get("usuario", ""), user.get("rol", ""), ahora_iso()),
    )


def esta_bloqueado(curso_id: int) -> Tuple[bool, Dict[str, Any]]:
    row = db_execute("SELECT * FROM curso_bloqueos WHERE curso_id=?", (int(curso_id),), fetchone=True)
    data = _dict_row(row)
    return bool(int(data.get("bloqueado", 0) or 0)), data


def bloquear_curso(curso_id: int, motivo: str):
    user = st.session_state.get("auth_user", {}) if "st" in globals() else {}
    digest, _ = _hash_expediente(int(curso_id)) if "_hash_expediente" in globals() else (hash_documental({"curso_id": curso_id, "motivo": motivo}), {})
    db_execute("DELETE FROM curso_bloqueos WHERE curso_id=?", (int(curso_id),))
    db_execute(
        "INSERT INTO curso_bloqueos(curso_id, bloqueado, motivo, hash_bloqueo, bloqueado_por, bloqueado_en) VALUES (?, 1, ?, ?, ?, ?)",
        (int(curso_id), motivo, digest, user.get("usuario", ""), ahora_iso()),
    )
    registrar_workflow_evento(int(curso_id), "Bloqueo documental", detalle=motivo, resultado="Bloqueado")


def desbloquear_curso(curso_id: int, motivo: str):
    db_execute("UPDATE curso_bloqueos SET bloqueado=0, motivo=?, bloqueado_en=? WHERE curso_id=?", (motivo, ahora_iso(), int(curso_id)))
    registrar_workflow_evento(int(curso_id), "Desbloqueo documental", detalle=motivo, resultado="Desbloqueado")


_base_upsert_curso_v7 = upsert_curso

def upsert_curso(curso_id: Optional[int], datos: Dict[str, str], payload: Optional[Dict] = None) -> int:
    if curso_id:
        bloqueado, info = esta_bloqueado(int(curso_id))
        user = st.session_state.get("auth_user", {}) if "st" in globals() else {}
        # El administrador puede guardar, pero queda auditado. Docente/coordinador deben desbloquear o crear nueva versión.
        if bloqueado and user.get("rol") not in ("Administrador",):
            raise RuntimeError(f"El expediente está bloqueado por aprobación/cierre. Motivo: {info.get('motivo','')}")
    return _base_upsert_curso_v7(curso_id, datos, payload)


def _curso_payload_partes(curso_id: int) -> Tuple[Dict[str, Any], Dict[str, Any], pd.DataFrame, pd.DataFrame, pd.DataFrame]:
    curso = get_curso(int(curso_id)) or {}
    payload = safe_json_loads(curso.get("payload_json"), {})
    datos = payload.get("datos", {}) or {}
    if not datos:
        datos = dict(curso)
    sesiones = payload_to_df(payload.get("sesiones"), COLUMNAS_SESIONES) if "payload_to_df" in globals() else _json_df(payload.get("sesiones"))
    evaluaciones = payload_to_df(payload.get("evaluaciones"), COLUMNAS_EVALUACIONES) if "payload_to_df" in globals() else _json_df(payload.get("evaluaciones"))
    modulos = payload_to_df(payload.get("modulos"), COLUMNAS_MODULOS) if "payload_to_df" in globals() else _json_df(payload.get("modulos"))
    return curso, datos, sesiones, evaluaciones, modulos


def analizar_coherencia_curso(curso_id: int) -> Dict[str, Any]:
    curso, datos, sesiones, evaluaciones, modulos = _curso_payload_partes(int(curso_id))
    hallazgos: List[Dict[str, Any]] = []
    checks: List[Tuple[str, int, int]] = []

    def add_check(componente: str, puntos: int, maxp: int, severidad: str, hallazgo: str, recomendacion: str):
        checks.append((componente, max(0, min(puntos, maxp)), maxp))
        if puntos < maxp:
            hallazgos.append({
                "Componente": componente,
                "Severidad": severidad,
                "Hallazgo": hallazgo,
                "Recomendación": recomendacion,
                "Puntos": puntos,
                "Máximo": maxp,
            })

    # Identificación
    campos_base = [curso.get("asignatura"), curso.get("programa"), curso.get("periodo"), curso.get("profesor"), curso.get("grupo")]
    completos = sum(1 for v in campos_base if str(v or "").strip())
    add_check("Identificación", int(completos / 5 * 10), 10, "Media", "Faltan campos de identificación del curso.", "Complete asignatura, programa, período, profesor y grupo.")

    # Resultados / competencias
    resultados = _text_lines(datos.get("resultados_aprendizaje") or datos.get("resultados") or datos.get("Resultados de aprendizaje") or datos.get("ra") or "")
    competencias = _text_lines(datos.get("competencias") or datos.get("Competencias") or "")
    pts_ra = 12 if resultados else (6 if competencias else 0)
    add_check("Resultados de aprendizaje", pts_ra, 12, "Alta", "No se identifican resultados de aprendizaje explícitos.", "Registre resultados de aprendizaje medibles y verificables en el FD-GC71.")

    # Cronograma / sesiones
    sesiones_validas = not sesiones.empty
    contenidos_vacios = 0
    if sesiones_validas:
        col_contenido = "Contenido por desarrollar" if "Contenido por desarrollar" in sesiones.columns else ("Contenido" if "Contenido" in sesiones.columns else sesiones.columns[0])
        contenidos_vacios = int(sesiones[col_contenido].astype(str).str.strip().eq("").sum()) if col_contenido in sesiones.columns else 0
    pts_ses = 18 if sesiones_validas and contenidos_vacios == 0 else (10 if sesiones_validas else 0)
    add_check("Cronograma", pts_ses, 18, "Alta", f"El cronograma tiene {contenidos_vacios} sesión(es) sin contenido o no existe.", "Revise que todas las sesiones tengan fecha, contenido y trabajo presencial/independiente.")

    # Fechas fuera de período
    fuera = 0
    if sesiones_validas and "Fecha" in sesiones.columns:
        fi = pd.to_datetime(curso.get("fecha_inicio") or datos.get("fecha_inicio"), errors="coerce")
        ff = pd.to_datetime(curso.get("fecha_fin") or datos.get("fecha_fin"), errors="coerce")
        fechas = pd.to_datetime(sesiones["Fecha"], errors="coerce")
        if not pd.isna(fi) and not pd.isna(ff):
            fuera = int(((fechas < fi) | (fechas > ff)).fillna(False).sum())
    add_check("Fechas", 8 if fuera == 0 else 3, 8, "Media", f"Hay {fuera} sesión(es) por fuera del período académico.", "Reprograme sesiones fuera del rango del curso.")

    # Evaluación
    total_eval = 0.0
    eval_sin_fecha = 0
    eval_sin_desc = 0
    max_eval = 0.0
    if not evaluaciones.empty:
        total_eval = float(pd.to_numeric(evaluaciones.get("Valor (%)", pd.Series(dtype=float)), errors="coerce").fillna(0).sum())
        max_eval = float(pd.to_numeric(evaluaciones.get("Valor (%)", pd.Series(dtype=float)), errors="coerce").fillna(0).max()) if len(evaluaciones) else 0
        if "Fecha de realización" in evaluaciones.columns:
            eval_sin_fecha = int(evaluaciones["Fecha de realización"].astype(str).str.strip().isin(["", "NaT", "None"]).sum())
        if "Procedimiento de evaluación" in evaluaciones.columns:
            eval_sin_desc = int(evaluaciones["Procedimiento de evaluación"].astype(str).str.strip().eq("").sum())
    eval_ok = abs(total_eval - 100) <= 0.01 and eval_sin_fecha == 0 and eval_sin_desc == 0 and max_eval <= 45
    pts_eval = 20 if eval_ok else (12 if total_eval > 0 else 0)
    add_check("Evaluación", pts_eval, 20, "Alta", f"Evaluación con suma {total_eval:.1f}%, {eval_sin_fecha} sin fecha, {eval_sin_desc} sin descripción, peso máximo {max_eval:.1f}%.", "La evaluación debe sumar 100%, tener fechas/procedimientos y evitar concentración excesiva en una sola actividad.")

    # Alineación RA-contenidos-evaluación (heurística por palabras clave)
    unidades = []
    if not modulos.empty:
        for _, r in modulos.iterrows():
            txt = f"{r.get('Unidad','')} {r.get('Contenido / tema central','')}"
            if txt.strip(): unidades.append(txt)
    elif sesiones_validas:
        col = "Contenido por desarrollar" if "Contenido por desarrollar" in sesiones.columns else "Contenido"
        if col in sesiones.columns:
            unidades = sesiones[col].astype(str).head(8).tolist()
    texto_contenidos = _norm_text(" ".join(unidades))
    texto_eval = _norm_text(" ".join(evaluaciones.astype(str).fillna("").agg(" ".join, axis=1).tolist())) if not evaluaciones.empty else ""
    alineados = 0
    matriz = []
    for idx, ra in enumerate(resultados, start=1):
        palabras = [p for p in re.findall(r"[a-záéíóúñ]{5,}", _norm_text(ra)) if p not in {"estudiante", "aprendizaje", "curso", "asignatura", "capacidad", "desarrollar"}]
        score_c = sum(1 for p in palabras[:8] if p in texto_contenidos)
        score_e = sum(1 for p in palabras[:8] if p in texto_eval)
        ok = score_c > 0 and score_e > 0
        alineados += 1 if ok else 0
        matriz.append({"Resultado": f"RA{idx}", "Descripción": ra, "Contenido relacionado": "Sí" if score_c > 0 else "No evidente", "Evaluación relacionada": "Sí" if score_e > 0 else "No evidente", "Lectura": "Alineado" if ok else "Requiere asociación explícita"})
    if resultados:
        pts_align = int((alineados / max(len(resultados), 1)) * 17)
    else:
        pts_align = 0
    add_check("Alineación académica", pts_align, 17, "Alta", f"{alineados} de {len(resultados)} resultado(s) tienen evidencia de relación con contenidos y evaluación.", "Relacione cada resultado de aprendizaje con unidades temáticas y actividades evaluativas.")

    # Evidencias y observaciones
    evid = evidencias_count(int(curso_id)) if "evidencias_count" in globals() else 0
    obs = observaciones_curso(int(curso_id)) if "observaciones_curso" in globals() else pd.DataFrame()
    obs_abiertas = int((obs.get("estado", pd.Series(dtype=str)).astype(str) == "Abierta").sum()) if not obs.empty else 0
    pts_evid = 10 if evid >= 2 else (5 if evid else 0)
    add_check("Evidencias", pts_evid, 10, "Media", f"El expediente tiene {evid} evidencia(s).", "Cargue evidencia de socialización, listados, soportes de evaluación y documentos del curso.")
    pts_obs = 5 if obs_abiertas == 0 else 0
    add_check("Observaciones", pts_obs, 5, "Alta", f"Hay {obs_abiertas} observación(es) abierta(s).", "Resuelva observaciones antes de solicitar aprobación o cierre.")

    total = sum(p for _, p, _ in checks)
    max_total = sum(m for _, _, m in checks) or 1
    score = round(total / max_total * 100, 1)
    riesgo = "Bajo" if score >= 85 else ("Medio" if score >= 70 else "Alto")
    estado_sugerido = "Aprobable" if score >= 85 and obs_abiertas == 0 else ("Aprobable con observaciones" if score >= 75 else "No aprobable")
    return {
        "score": score,
        "riesgo": riesgo,
        "estado_sugerido": estado_sugerido,
        "hallazgos": pd.DataFrame(hallazgos),
        "matriz": pd.DataFrame(matriz),
        "checks": pd.DataFrame([{"Componente": c, "Puntos": p, "Máximo": m, "Cumplimiento": round(p / m * 100, 1) if m else 0} for c, p, m in checks]),
        "total_eval": total_eval,
        "resultados": resultados,
    }


def matriz_institucional_cursos() -> pd.DataFrame:
    df = listar_cursos_visibles()
    if df.empty:
        return df
    rows = []
    for _, r in df.iterrows():
        curso_id = int(r.get("id"))
        try:
            ana = analizar_coherencia_curso(curso_id)
            score = ana["score"]
            riesgo = ana["riesgo"]
            sugerido = ana["estado_sugerido"]
            hallazgos_altos = int((ana["hallazgos"].get("Severidad", pd.Series(dtype=str)).astype(str) == "Alta").sum()) if not ana["hallazgos"].empty else 0
        except Exception:
            score, riesgo, sugerido, hallazgos_altos = 0, "Sin análisis", "Revisar", 0
        evid = evidencias_count(curso_id) if "evidencias_count" in globals() else 0
        obs = observaciones_curso(curso_id) if "observaciones_curso" in globals() else pd.DataFrame()
        obs_abiertas = int((obs.get("estado", pd.Series(dtype=str)).astype(str) == "Abierta").sum()) if not obs.empty else 0
        bloqueado, _ = esta_bloqueado(curso_id)
        rows.append({
            "ID": curso_id,
            "Programa": r.get("programa", ""),
            "Asignatura": r.get("asignatura", ""),
            "Grupo": r.get("grupo", ""),
            "Período": r.get("periodo", ""),
            "Docente": r.get("profesor", ""),
            "Estado": r.get("estado", ""),
            "Score coherencia": score,
            "Riesgo": riesgo,
            "Sugerencia": sugerido,
            "Hallazgos altos": hallazgos_altos,
            "Evidencias": evid,
            "Observaciones abiertas": obs_abiertas,
            "Bloqueado": "Sí" if bloqueado else "No",
        })
    return pd.DataFrame(rows)


def excel_bytes_sheets(sheets: Dict[str, pd.DataFrame]) -> bytes:
    bio = io.BytesIO()
    with pd.ExcelWriter(bio, engine="xlsxwriter") as writer:
        for name, df in sheets.items():
            safe = re.sub(r"[^A-Za-z0-9 _-]", "", str(name))[:31] or "Hoja"
            (df if isinstance(df, pd.DataFrame) else pd.DataFrame()).to_excel(writer, index=False, sheet_name=safe)
            ws = writer.sheets[safe]
            wb = writer.book
            header_fmt = wb.add_format({"bold": True, "bg_color": "#1F4FD8", "font_color": "white", "border": 1})
            body_fmt = wb.add_format({"text_wrap": True, "valign": "top"})
            for col_num, value in enumerate((df.columns if isinstance(df, pd.DataFrame) else [])):
                ws.write(0, col_num, value, header_fmt)
                width = min(max(len(str(value)) + 4, 14), 42)
                try:
                    sample = df.iloc[:100, col_num].astype(str).map(len).max() if not df.empty else 0
                    width = min(max(width, int(sample) + 2), 48)
                except Exception:
                    pass
                ws.set_column(col_num, col_num, width, body_fmt)
            if isinstance(df, pd.DataFrame) and not df.empty:
                ws.autofilter(0, 0, len(df), max(len(df.columns) - 1, 0))
                ws.freeze_panes(1, 0)
    return bio.getvalue()


def crear_informe_ejecutivo_institucional_docx(matriz: pd.DataFrame, resumen: Dict[str, Any]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    title = doc.add_paragraph()
    run = title.add_run("INFORME EJECUTIVO INSTITUCIONAL DE SEGUIMIENTO ACADÉMICO")
    run.bold = True
    run.font.size = Pt(15)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Generado en: {ahora_iso()} · Versión del sistema: {APP_VERSION}")
    doc.add_paragraph(f"Cursos analizados: {resumen.get('cursos', 0)} · Score promedio: {resumen.get('score_promedio', 0):.1f} · Riesgo alto: {resumen.get('riesgo_alto', 0)} · Observaciones abiertas: {resumen.get('obs_abiertas', 0)}")
    doc.add_heading("Lectura ejecutiva", level=1)
    lectura = resumen.get("lectura", "") or "El sistema consolida el estado documental y académico de los expedientes a partir de planeación, evaluación, evidencias, observaciones y coherencia académica."
    doc.add_paragraph(lectura)
    doc.add_heading("Cursos críticos y seguimiento", level=1)
    cols = ["ID", "Programa", "Asignatura", "Grupo", "Estado", "Score coherencia", "Riesgo", "Observaciones abiertas", "Sugerencia"]
    df = matriz[cols].copy() if not matriz.empty else pd.DataFrame(columns=cols)
    if len(df) > 20:
        df = df.sort_values(["Riesgo", "Score coherencia"], ascending=[False, True]).head(20)
    table = doc.add_table(rows=1, cols=len(cols))
    table.style = "Table Grid"
    for i, c in enumerate(cols):
        cell = table.rows[0].cells[i]
        cell.text = c
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(8)
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(cols):
            cells[i].text = str(row.get(c, ""))
            for p in cells[i].paragraphs:
                for rr in p.runs:
                    rr.font.size = Pt(7)
    doc.add_heading("Recomendaciones institucionales", level=1)
    for rec in resumen.get("recomendaciones", []):
        doc.add_paragraph(str(rec), style="List Bullet")
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()


def ui_banco_asignaturas(st):
    st.markdown("<div class='ux-soft-box'><strong>Banco institucional de asignaturas.</strong><br>Use este módulo como fuente maestra para no repetir justificaciones, resultados, unidades y evaluaciones desde cero.</div>", unsafe_allow_html=True)
    df = listar_asignaturas_base()
    tab_lista, tab_crear, tab_clonar = st.tabs(["Catálogo", "Crear / editar", "Crear curso desde banco"])
    with tab_lista:
        if df.empty:
            st.info("Aún no hay asignaturas base registradas.")
        else:
            st.dataframe(df[[c for c in ["id", "codigo", "nombre", "programa", "creditos", "htp", "hti", "tipo_asignatura", "actualizado_en"] if c in df.columns]], use_container_width=True, hide_index=True)
            st.download_button("⬇️ Descargar banco CSV", df.to_csv(index=False).encode("utf-8-sig"), "banco_asignaturas.csv", "text/csv", use_container_width=True)
    with tab_crear:
        opciones = {"Nueva asignatura": None}
        for _, r in df.iterrows():
            opciones[f"#{r.get('id')} · {r.get('nombre','')} · {r.get('programa','')}"] = int(r.get("id"))
        sel = st.selectbox("Registro", list(opciones.keys()), key="banco_sel")
        reg = get_asignatura_base(opciones[sel]) if opciones[sel] else {}
        unidades_base = _json_df(reg.get("unidades_json", "[]"), ["Unidad", "Contenido / tema central", "Horas presenciales", "Sesiones", "Trabajo presencial", "Trabajo independiente"])
        evaluaciones_base = _json_df(reg.get("evaluaciones_json", "[]"), ["Tipo de evaluación", "Procedimiento de evaluación", "Valor (%)", "Fecha de realización", "Unidad relacionada", "Corte"])
        with st.form("form_banco_asignatura"):
            c1, c2, c3 = st.columns(3)
            codigo = c1.text_input("Código", value=str(reg.get("codigo", "")))
            nombre = c2.text_input("Asignatura", value=str(reg.get("nombre", "")))
            programa = c3.text_input("Programa", value=str(reg.get("programa", "")))
            c4, c5, c6, c7 = st.columns(4)
            area = c4.text_input("Área de formación", value=str(reg.get("area_formacion", "")))
            creditos = c5.text_input("Créditos", value=str(reg.get("creditos", "")))
            htp = c6.number_input("HTP", min_value=0.0, value=float(reg.get("htp", 0) or 0), step=1.0)
            hti = c7.number_input("HTI", min_value=0.0, value=float(reg.get("hti", 0) or 0), step=1.0)
            tipo = st.selectbox("Tipo de asignatura", ["Teórica", "Teórico-práctica", "Práctica"], index=0 if not reg.get("tipo_asignatura") else max(0, ["Teórica", "Teórico-práctica", "Práctica"].index(reg.get("tipo_asignatura")) if reg.get("tipo_asignatura") in ["Teórica", "Teórico-práctica", "Práctica"] else 0))
            justificacion = st.text_area("Justificación", value=str(reg.get("justificacion", "")), height=90)
            competencias = st.text_area("Competencias", value=str(reg.get("competencias", "")), height=90)
            resultados = st.text_area("Resultados de aprendizaje", value=str(reg.get("resultados", "")), height=90)
            objetivos = st.text_area("Objetivos", value=str(reg.get("objetivos", "")), height=90)
            metodologia = st.text_area("Metodología", value=str(reg.get("metodologia", "")), height=90)
            ambientes = st.text_area("Ambientes", value=str(reg.get("ambientes", "")), height=75)
            medios = st.text_area("Medios educativos", value=str(reg.get("medios", "")), height=75)
            bibliografia = st.text_area("Bibliografía", value=str(reg.get("bibliografia", "")), height=90)
            unidades_edit = st.data_editor(unidades_base, num_rows="dynamic", use_container_width=True, key="banco_unidades")
            eval_edit = st.data_editor(evaluaciones_base, num_rows="dynamic", use_container_width=True, key="banco_evals")
            submitted = st.form_submit_button("Guardar asignatura base", type="primary", use_container_width=True)
        if submitted:
            data = {"codigo": codigo, "nombre": nombre, "programa": programa, "area_formacion": area, "creditos": creditos, "htp": htp, "hti": hti, "tipo_asignatura": tipo, "justificacion": justificacion, "competencias": competencias, "resultados": resultados, "objetivos": objetivos, "metodologia": metodologia, "ambientes": ambientes, "medios": medios, "bibliografia": bibliografia, "unidades": unidades_edit.fillna("").to_dict("records"), "evaluaciones": eval_edit.fillna("").to_dict("records")}
            guardar_asignatura_base(data, opciones[sel])
            st.success("Asignatura base guardada.")
            st.rerun()
    with tab_clonar:
        df = listar_asignaturas_base()
        if df.empty:
            st.info("Cree primero una asignatura base.")
        else:
            opciones2 = {f"#{r.get('id')} · {r.get('nombre','')} · {r.get('programa','')}": int(r.get("id")) for _, r in df.iterrows()}
            sel2 = st.selectbox("Asignatura base", list(opciones2.keys()), key="banco_clonar_sel")
            reg2 = get_asignatura_base(opciones2[sel2]) or {}
            c1, c2, c3 = st.columns(3)
            periodo = c1.text_input("Período académico", value="2026-2", key="banco_periodo")
            grupo = c2.text_input("Grupo", value="01", key="banco_grupo")
            profesor = c3.text_input("Profesor", value=st.session_state.get("auth_user", {}).get("nombre_completo", ""), key="banco_profesor")
            if st.button("Crear expediente desde asignatura base", type="primary", use_container_width=True):
                datos = {
                    "codigo": reg2.get("codigo", ""), "grupo": grupo, "asignatura": reg2.get("nombre", ""), "programa": reg2.get("programa", ""), "periodo": periodo, "profesor": profesor,
                    "creditos": reg2.get("creditos", ""), "htp": reg2.get("htp", 0), "hti": reg2.get("hti", 0), "estado": "Borrador",
                    "justificacion": reg2.get("justificacion", ""), "competencias": reg2.get("competencias", ""), "resultados_aprendizaje": reg2.get("resultados", ""), "objetivos": reg2.get("objetivos", ""), "metodologia": reg2.get("metodologia", ""), "ambientes": reg2.get("ambientes", ""), "medios": reg2.get("medios", ""), "bibliografia": reg2.get("bibliografia", ""),
                }
                payload = {"datos": datos, "modulos": safe_json_loads(reg2.get("unidades_json"), []), "evaluaciones": safe_json_loads(reg2.get("evaluaciones_json"), []), "sesiones": []}
                cid = upsert_curso(None, datos, payload)
                registrar_workflow_evento(cid, "Creación desde banco", nuevo="Borrador", resultado="OK", detalle=f"Asignatura base {reg2.get('id')}")
                st.success(f"Expediente creado con ID {cid}. Ábralo en Planeador superior para generar cronograma.")


def ui_coherencia_academica(st):
    curso_id = seleccionar_curso_widget("Curso a analizar", key="coh_curso")
    if not curso_id:
        return
    ana = analizar_coherencia_curso(int(curso_id))
    c1, c2, c3 = st.columns(3)
    c1.metric("Score de coherencia", f"{ana['score']}/100")
    c2.metric("Riesgo", ana["riesgo"])
    c3.metric("Concepto", ana["estado_sugerido"])
    st.progress(float(ana["score"]) / 100.0, text=f"Coherencia académica: {ana['score']}%")
    tab1, tab2, tab3, tab4 = st.tabs(["Hallazgos", "Matriz RA-contenido-evaluación", "Checklist", "Exportar"])
    with tab1:
        if ana["hallazgos"].empty:
            st.success("Sin hallazgos relevantes. Expediente técnicamente consistente.")
        else:
            st.dataframe(ana["hallazgos"], use_container_width=True, hide_index=True)
    with tab2:
        if ana["matriz"].empty:
            st.warning("No hay resultados de aprendizaje explícitos para construir matriz de alineación.")
        else:
            st.dataframe(ana["matriz"], use_container_width=True, hide_index=True)
    with tab3:
        st.dataframe(ana["checks"], use_container_width=True, hide_index=True)
        for _, r in ana["checks"].iterrows():
            st.progress(float(r["Cumplimiento"]) / 100.0, text=f"{r['Componente']}: {r['Cumplimiento']}%")
    with tab4:
        xlsx = excel_bytes_sheets({"Hallazgos": ana["hallazgos"], "Matriz": ana["matriz"], "Checklist": ana["checks"]})
        st.download_button("⬇️ Descargar análisis de coherencia Excel", xlsx, f"coherencia_curso_{curso_id}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
        informe_txt = f"ANÁLISIS DE COHERENCIA ACADÉMICA\nCurso ID: {curso_id}\nScore: {ana['score']}/100\nRiesgo: {ana['riesgo']}\nConcepto: {ana['estado_sugerido']}\n\nHallazgos:\n" + (ana["hallazgos"].to_string(index=False) if not ana["hallazgos"].empty else "Sin hallazgos")
        st.download_button("⬇️ Descargar informe TXT", informe_txt.encode("utf-8"), f"coherencia_curso_{curso_id}.txt", "text/plain", use_container_width=True)


def ui_aprobacion_bloqueante(st):
    st.markdown("<div class='ux-soft-box'><strong>Flujo bloqueante.</strong><br>Los expedientes aprobados o cerrados quedan bloqueados. Cualquier ajuste posterior debe quedar auditado o generar nueva versión.</div>", unsafe_allow_html=True)
    curso_id = seleccionar_curso_widget("Expediente", key="bloq_curso")
    if not curso_id:
        return
    curso = get_curso(int(curso_id)) or {}
    ana = analizar_coherencia_curso(int(curso_id))
    bloqueado, info = esta_bloqueado(int(curso_id))
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Estado actual", curso.get("estado", ""))
    c2.metric("Score", f"{ana['score']}/100")
    c3.metric("Riesgo", ana["riesgo"])
    c4.metric("Bloqueado", "Sí" if bloqueado else "No")
    if bloqueado:
        st.warning(f"Expediente bloqueado: {info.get('motivo','')} · Hash: {info.get('hash_bloqueo','')[:16]}…")
    estado = str(curso.get("estado", "Borrador"))
    rol = st.session_state.get("auth_user", {}).get("rol", "")
    estados = ["Borrador", "En revisión", "Observado", "Ajustado", "Aprobado", "En ejecución", "Corte parcial", "Corte final", "Cerrado"]
    idx = estados.index(estado) if estado in estados else 0
    st.markdown("### Ruta de aprobación")
    st.markdown(" → ".join([f"**{e}**" if i == idx else e for i, e in enumerate(estados)]))
    nota = st.text_area("Nota / motivación del movimiento", height=90, key="bloq_nota")
    col1, col2, col3 = st.columns(3)
    puede_coord = rol in ("Administrador", "Coordinador")
    puede_doc = rol in ("Administrador", "Docente")
    with col1:
        if st.button("Enviar a revisión", disabled=not puede_doc or estado not in ["Borrador", "Observado", "Ajustado"], use_container_width=True):
            cambiar_estado_curso(int(curso_id), "En revisión", nota or "Enviado a revisión")
            registrar_workflow_evento(int(curso_id), "Envío a revisión", estado, "En revisión", "OK", nota)
            st.rerun()
    with col2:
        if st.button("Observar", disabled=not puede_coord or estado != "En revisión", use_container_width=True):
            cambiar_estado_curso(int(curso_id), "Observado", nota or "Observado por coordinación")
            if nota.strip():
                crear_observacion_curso(int(curso_id), "Alta", "Revisión académica", nota.strip())
            registrar_workflow_evento(int(curso_id), "Observación", estado, "Observado", "OK", nota)
            st.rerun()
    with col3:
        bloquear = ana["score"] >= 85 and ana["riesgo"] != "Alto" and estado == "En revisión"
        if st.button("Aprobar y bloquear", disabled=not puede_coord or not bloquear, type="primary", use_container_width=True):
            cambiar_estado_curso(int(curso_id), "Aprobado", nota or "Aprobado por coordinación")
            bloquear_curso(int(curso_id), nota or "Aprobación oficial del FD-GC71")
            registrar_workflow_evento(int(curso_id), "Aprobación bloqueante", estado, "Aprobado", "OK", nota)
            st.rerun()
        if not bloquear and estado == "En revisión":
            st.caption("Para aprobar: score ≥ 85, riesgo no alto y sin hallazgos críticos abiertos.")
    st.divider()
    c4, c5, c6 = st.columns(3)
    with c4:
        if st.button("Marcar en ejecución", disabled=not puede_coord or estado != "Aprobado", use_container_width=True):
            cambiar_estado_curso(int(curso_id), "En ejecución", nota or "Inicio de ejecución")
            registrar_workflow_evento(int(curso_id), "Inicio de ejecución", estado, "En ejecución", "OK", nota)
            st.rerun()
    with c5:
        if st.button("Cerrar expediente", disabled=not puede_coord or estado not in ["Corte final", "En ejecución", "Aprobado"], use_container_width=True):
            cambiar_estado_curso(int(curso_id), "Cerrado", nota or "Cierre de expediente")
            bloquear_curso(int(curso_id), nota or "Cierre oficial del expediente académico")
            registrar_workflow_evento(int(curso_id), "Cierre bloqueante", estado, "Cerrado", "OK", nota)
            st.rerun()
    with c6:
        if st.button("Desbloquear con justificación", disabled=not (rol == "Administrador" and bloqueado), use_container_width=True):
            desbloquear_curso(int(curso_id), nota or "Desbloqueo administrativo justificado")
            st.rerun()
    st.markdown("### Hallazgos que condicionan aprobación")
    st.dataframe(ana["hallazgos"], use_container_width=True, hide_index=True)


def ui_informe_institucional(st):
    matriz = matriz_institucional_cursos()
    if matriz.empty:
        st.info("No hay cursos para consolidar.")
        return
    st.markdown("<div class='ux-soft-box'><strong>Informe ejecutivo institucional.</strong><br>Consolidado para coordinación, comité o dirección académica. Prioriza cumplimiento, riesgo y trazabilidad.</div>", unsafe_allow_html=True)
    c1, c2, c3 = st.columns(3)
    periodos = ["Todos"] + sorted(matriz["Período"].dropna().astype(str).unique().tolist())
    programas = ["Todos"] + sorted(matriz["Programa"].dropna().astype(str).unique().tolist())
    riesgos = ["Todos"] + sorted(matriz["Riesgo"].dropna().astype(str).unique().tolist())
    periodo = c1.selectbox("Período", periodos, key="inf_periodo")
    programa = c2.selectbox("Programa", programas, key="inf_programa")
    riesgo = c3.selectbox("Riesgo", riesgos, key="inf_riesgo")
    f = matriz.copy()
    if periodo != "Todos": f = f[f["Período"].astype(str) == periodo]
    if programa != "Todos": f = f[f["Programa"].astype(str) == programa]
    if riesgo != "Todos": f = f[f["Riesgo"].astype(str) == riesgo]
    resumen = {
        "cursos": len(f),
        "score_promedio": float(pd.to_numeric(f["Score coherencia"], errors="coerce").fillna(0).mean()) if not f.empty else 0,
        "riesgo_alto": int((f["Riesgo"].astype(str) == "Alto").sum()) if not f.empty else 0,
        "obs_abiertas": int(pd.to_numeric(f["Observaciones abiertas"], errors="coerce").fillna(0).sum()) if not f.empty else 0,
        "recomendaciones": [
            "Priorizar revisión de expedientes con riesgo alto o score inferior a 70.",
            "Cerrar observaciones abiertas antes de aprobar o cerrar el expediente académico.",
            "Verificar que evaluación, resultados de aprendizaje y contenidos queden explícitamente alineados.",
            "Exigir evidencia de socialización y soportes de evaluación por cada curso.",
        ],
    }
    resumen["lectura"] = f"El consolidado incluye {resumen['cursos']} curso(s). El score promedio de coherencia es {resumen['score_promedio']:.1f}/100; se identifican {resumen['riesgo_alto']} curso(s) en riesgo alto y {resumen['obs_abiertas']} observación(es) abiertas."
    k1, k2, k3, k4 = st.columns(4)
    k1.metric("Cursos", resumen["cursos"])
    k2.metric("Score promedio", f"{resumen['score_promedio']:.1f}")
    k3.metric("Riesgo alto", resumen["riesgo_alto"])
    k4.metric("Obs. abiertas", resumen["obs_abiertas"])
    st.dataframe(f, use_container_width=True, hide_index=True)
    if not f.empty:
        st.bar_chart(f.set_index("Asignatura")[["Score coherencia"]])
    xlsx = excel_bytes_sheets({"Consolidado": f, "Resumen": pd.DataFrame([resumen]), "Recomendaciones": pd.DataFrame({"Recomendación": resumen["recomendaciones"]})})
    docx = crear_informe_ejecutivo_institucional_docx(f, resumen)
    c5, c6 = st.columns(2)
    c5.download_button("⬇️ Descargar matriz ejecutiva Excel", xlsx, "informe_ejecutivo_institucional.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
    c6.download_button("⬇️ Descargar informe ejecutivo Word", docx, "informe_ejecutivo_institucional.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)


def ui_exportacion_institucional(st):
    matriz = matriz_institucional_cursos()
    if matriz.empty:
        st.info("No hay cursos para exportar.")
        return
    st.markdown("<div class='ux-soft-box'><strong>Exportación institucional estructurada.</strong><br>Genera un ZIP maestro ordenado por período, programa y curso, con manifiestos y matrices de seguimiento.</div>", unsafe_allow_html=True)
    c1, c2 = st.columns(2)
    periodo = c1.selectbox("Período", ["Todos"] + sorted(matriz["Período"].dropna().astype(str).unique().tolist()), key="expinst_periodo")
    programa = c2.selectbox("Programa", ["Todos"] + sorted(matriz["Programa"].dropna().astype(str).unique().tolist()), key="expinst_programa")
    f = matriz.copy()
    if periodo != "Todos": f = f[f["Período"].astype(str) == periodo]
    if programa != "Todos": f = f[f["Programa"].astype(str) == programa]
    st.metric("Expedientes incluidos", len(f))
    st.dataframe(f, use_container_width=True, hide_index=True)
    if st.button("Generar ZIP institucional", type="primary", use_container_width=True):
        bio = io.BytesIO()
        resumen = {
            "cursos": len(f),
            "score_promedio": float(pd.to_numeric(f["Score coherencia"], errors="coerce").fillna(0).mean()) if not f.empty else 0,
            "riesgo_alto": int((f["Riesgo"].astype(str) == "Alto").sum()) if not f.empty else 0,
            "obs_abiertas": int(pd.to_numeric(f["Observaciones abiertas"], errors="coerce").fillna(0).sum()) if not f.empty else 0,
            "recomendaciones": ["Revisar expedientes críticos.", "Cerrar observaciones.", "Validar evidencia de socialización."],
            "lectura": "Exportación institucional generada desde el gestor académico.",
        }
        with zipfile.ZipFile(bio, "w", zipfile.ZIP_DEFLATED) as z:
            z.writestr("00_indice/Consolidado_institucional.csv", f.to_csv(index=False).encode("utf-8-sig"))
            z.writestr("00_indice/Consolidado_institucional.xlsx", excel_bytes_sheets({"Consolidado": f}))
            z.writestr("00_indice/Informe_ejecutivo.docx", crear_informe_ejecutivo_institucional_docx(f, resumen))
            manifiesto = {"generado_en": ahora_iso(), "version": APP_VERSION, "usuario": st.session_state.get("auth_user", {}).get("usuario"), "filtros": {"periodo": periodo, "programa": programa}, "total": len(f), "sha256_consolidado": hash_documental(f.to_dict("records"))}
            z.writestr("00_indice/manifiesto_institucional.json", json.dumps(manifiesto, ensure_ascii=False, indent=2))
            for _, r in f.iterrows():
                curso_id = int(r.get("ID"))
                curso, datos, sesiones, evaluaciones, modulos = _curso_payload_partes(curso_id)
                periodo_dir = safe_filename(r.get("Período", "Sin_periodo"))
                programa_dir = safe_filename(r.get("Programa", "Sin_programa"))
                curso_dir = safe_filename(f"{curso_id}_{r.get('Asignatura','Curso')}_{r.get('Grupo','')}")
                base = f"{periodo_dir}/{programa_dir}/{curso_dir}"
                try:
                    paquete = crear_paquete_curso_zip(datos, sesiones, evaluaciones, pd.DataFrame(), pd.DataFrame(), curso_id)
                    z.writestr(f"{base}/Expediente_curso.zip", paquete)
                except Exception as exc:
                    z.writestr(f"{base}/ERROR_generacion.txt", str(exc))
                ana = analizar_coherencia_curso(curso_id)
                z.writestr(f"{base}/coherencia.xlsx", excel_bytes_sheets({"Hallazgos": ana["hallazgos"], "Matriz": ana["matriz"], "Checklist": ana["checks"]}))
                z.writestr(f"{base}/manifiesto_expediente.json", json.dumps({"curso_id": curso_id, "curso": dict(curso), "score": ana["score"], "riesgo": ana["riesgo"], "hash": _hash_expediente(curso_id)[0] if "_hash_expediente" in globals() else ""}, ensure_ascii=False, indent=2, default=str))
        registrar_auditoria("Exportación institucional V7", f"Expedientes: {len(f)}")
        st.download_button("⬇️ Descargar ZIP institucional", bio.getvalue(), f"exportacion_institucional_{datetime.now().strftime('%Y%m%d_%H%M')}.zip", "application/zip", use_container_width=True)


def ui_auditoria_expediente(st):
    curso_id = seleccionar_curso_widget("Expediente a auditar", key="audexp_curso")
    if not curso_id:
        return
    tabs = st.tabs(["Versiones", "Observaciones", "Eventos workflow", "Auditoría general", "Bloqueo"])
    with tabs[0]:
        dfv = versiones_curso(int(curso_id)) if "versiones_curso" in globals() else pd.DataFrame()
        st.dataframe(dfv, use_container_width=True, hide_index=True)
    with tabs[1]:
        dfo = observaciones_curso(int(curso_id)) if "observaciones_curso" in globals() else pd.DataFrame()
        st.dataframe(dfo, use_container_width=True, hide_index=True)
    with tabs[2]:
        try:
            dfe = read_sql_df("SELECT * FROM workflow_eventos WHERE curso_id=? ORDER BY id DESC", params=(int(curso_id),))
        except Exception:
            dfe = pd.DataFrame()
        st.dataframe(dfe, use_container_width=True, hide_index=True)
    with tabs[3]:
        try:
            dfa = read_sql_df("SELECT * FROM auditoria WHERE detalle LIKE ? ORDER BY id DESC", params=(f"%Curso ID={curso_id}%",))
        except Exception:
            dfa = pd.DataFrame()
        st.dataframe(dfa, use_container_width=True, hide_index=True)
    with tabs[4]:
        bloqueado, info = esta_bloqueado(int(curso_id))
        st.json(info if info else {"bloqueado": False})
    # exportación unificada
    try:
        paquete = excel_bytes_sheets({
            "Versiones": dfv if 'dfv' in locals() else pd.DataFrame(),
            "Observaciones": dfo if 'dfo' in locals() else pd.DataFrame(),
            "Workflow": dfe if 'dfe' in locals() else pd.DataFrame(),
            "Auditoria": dfa if 'dfa' in locals() else pd.DataFrame(),
            "Bloqueo": pd.DataFrame([info]) if 'info' in locals() and info else pd.DataFrame(),
        })
        st.download_button("⬇️ Descargar auditoría del expediente", paquete, f"auditoria_expediente_{curso_id}.xlsx", "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet", use_container_width=True)
    except Exception:
        pass


def ui_inicio(st):
    counts = ux_system_counts()
    matriz = matriz_institucional_cursos()
    score_prom = float(pd.to_numeric(matriz.get("Score coherencia", pd.Series(dtype=float)), errors="coerce").fillna(0).mean()) if not matriz.empty else 0
    riesgo_alto = int((matriz.get("Riesgo", pd.Series(dtype=str)).astype(str) == "Alto").sum()) if not matriz.empty else 0
    st.markdown(
        "<div class='ux-card-grid'>" +
        ux_card("🏛️", "Banco institucional", "Asignaturas base con resultados, unidades, evaluación y bibliografía para no empezar desde cero.") +
        ux_card("🧪", "Coherencia académica", "El sistema cruza resultados, contenidos, evaluación, horas, evidencias y observaciones.") +
        ux_card("🔒", "Aprobación bloqueante", "Versiones aprobadas o cerradas quedan bloqueadas con hash y justificación.") +
        ux_card("🏢", "Informe institucional", "Matriz ejecutiva, Word de comité y exportación estructurada por período/programa.") +
        "</div>", unsafe_allow_html=True)
    c1, c2, c3, c4 = st.columns(4)
    c1.metric("Cursos", counts.get("cursos", 0))
    c2.metric("Riesgo alto", riesgo_alto)
    c3.metric("Observaciones abiertas", counts.get("abiertas", 0))
    c4.metric("Score coherencia", f"{score_prom:.1f}")
    st.markdown("### Ruta institucional recomendada")
    st.markdown("""
1. Configure el **Banco institucional de asignaturas**.  
2. Cree el expediente desde banco o desde el planeador.  
3. Genere FD-GC71, cronograma y evaluación concertada.  
4. Valide **coherencia académica** antes de enviar a revisión.  
5. Aplique **aprobación bloqueante** con versión, hash y observaciones.  
6. Cargue cortes, evidencias y genere FD-GC72.  
7. Exporte informe institucional y ZIP maestro por período/programa.
""")
    if not matriz.empty:
        st.markdown("### Semáforo ejecutivo")
        st.dataframe(matriz.head(12), use_container_width=True, hide_index=True)


# Router final V7. Sobrescribe main para incluir gobierno institucional completo.
def main():
    import streamlit as st
    globals()["st"] = st
    st.set_page_config(
        page_title="Gestor Académico Institucional FD-GC71 / FD-GC72",
        page_icon="🎓",
        layout="wide",
        initial_sidebar_state="expanded",
    )
    ux_apply_theme(st)
    try:
        init_db()
    except Exception as exc:
        st.error("La aplicación inició, pero no pudo preparar la base de datos.")
        st.markdown("""
        Esto suele pasar en Streamlit Cloud cuando `DATABASE_URL` está mal escrita,
        contiene el ejemplo `USUARIO:CLAVE@HOST:5432/BASE`, la base de datos no acepta conexiones externas
        o el proveedor requiere usar el pooler/puerto correcto.
        """)
        st.code(str(exc), language="text")
        st.info("Corrija los Secrets en Streamlit Cloud y reinicie la app. Para probar sin base externa, quite `DATABASE_URL` y use modo local/demo.")
        st.stop()

    if "auth_user" not in st.session_state:
        pantalla_login(st)
        return

    user = st.session_state.get("auth_user", {})
    pagina = ux_sidebar(st, user)

    if not tiene_permiso(pagina):
        st.error("Este perfil no tiene permisos para abrir este módulo.")
        return

    ux_render_hero(st, pagina, user)
    if pagina not in ["Inicio", "Mi cuenta", "Ayuda / flujo recomendado", MODULO_DIAGNOSTICO]:
        ux_render_path(st, pagina)

    if pagina == "Inicio":
        ui_inicio(st)
    elif pagina == MODULO_BANCO:
        ui_banco_asignaturas(st)
    elif pagina == MODULO_COHERENCIA:
        ui_coherencia_academica(st)
    elif pagina == MODULO_APROBACION_BLOQUEANTE:
        ui_aprobacion_bloqueante(st)
    elif pagina == MODULO_INFORME_INSTITUCIONAL:
        ui_informe_institucional(st)
    elif pagina == MODULO_EXPORTACION_INSTITUCIONAL:
        ui_exportacion_institucional(st)
    elif pagina == MODULO_AUDITORIA_EXPEDIENTE:
        ui_auditoria_expediente(st)
    elif pagina == MODULO_CENTRO:
        ui_centro_control(st)
    elif pagina == MODULO_SEMAFORO:
        ui_semaforo_expediente(st)
    elif pagina == MODULO_EXPEDIENTE:
        ui_expediente_academico(st)
    elif pagina == MODULO_PLANEADOR:
        ui_planeador_superior(st)
    elif pagina == MODULO_ASISTENTE:
        ui_asistente_academico(st)
    elif pagina == MODULO_CARGADOR:
        ui_cargador_inteligente(st)
    elif pagina == MODULO_COMPARADOR:
        ui_comparador_cortes(st)
    elif pagina == MODULO_FLUJO:
        ui_flujo_aprobaciones(st)
    elif pagina == MODULO_MOTOR:
        ui_motor_academico(st)
    elif pagina == MODULO_REPORTES:
        ui_reportes_ejecutivos(st)
    elif pagina == MODULO_EXPORTACION:
        ui_exportacion_masiva(st)
    elif pagina == MODULO_VERIFICACION:
        ui_verificacion_documental(st)
    elif pagina == MODULO_PARAMETROS:
        ui_parametros(st)
    elif pagina.startswith("FD-GC71"):
        ui_gc71(st)
    elif pagina.startswith("FD-GC72"):
        ui_gc72(st)
    elif pagina == MODULO_EVIDENCIAS:
        ui_evidencias(st)
    elif pagina == MODULO_VALIDADOR:
        ui_validador_institucional(st)
    elif pagina == MODULO_BACKUP:
        ui_backup(st)
    elif pagina == MODULO_DIAGNOSTICO:
        ui_diagnostico_productivo(st)
    elif pagina == "Usuarios y perfiles":
        ui_admin_usuarios(st)
    elif pagina == "Auditoría":
        ui_auditoria(st)
    elif pagina == "Mi cuenta":
        ui_mi_cuenta(st)
    else:
        ui_ayuda(st)

    st.markdown("<div class='ux-footer-note'>Gestor Académico Institucional · banco → planeación → coherencia → aprobación → cortes → informe → exportación.</div>", unsafe_allow_html=True)

if __name__ == "__main__":
    main()
